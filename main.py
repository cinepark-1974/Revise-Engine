# =================================================================
# 👖 BLUE JEANS REVISE ENGINE
# main.py — Streamlit App (3-Stage Pipeline)
# =================================================================
# © 2026 BLUE JEANS PICTURES. All rights reserved.
#
# v1.0 (2026-04-21)  초기 릴리스
# v2.6 (2026-04-30)  헐리우드 작법 A25~A28 추가
# v2.7 (2026-05-03)  ★ 자동 배치 분할 시스템 — 토큰 잘림 결함 해결
# v2.8 (2026-05-03)  ★ Beat-Aware Diagnose — 15-Beat 구조 인식 + 약점 비트 보강 ADD 자동 분배.
#                    71→100씬 확장 같은 대규모 작업 지원.
# v2.9 (2026-05-03)  ★ 비트 보강 확장 모드 (4번째 작업 모드) — Beat-Aware + LOCKED 강제 차단.
# v3.0 (2026-05-03)  ★ 라우팅 버그 수정 — 모드 전환 시 잔존 데이터 강제 초기화.
#                    Fast Path 0 조건 강화 (work_mode in {expansion} only).
#                    진단 직전 모드 미리보기 박스 + 라우팅 디버그 표시.
# v3.1 (2026-05-04)  ★ UI 한국어화 — "배치" → "N차 각색"으로 통일.
#                    각 씬에 핵심 수정 방향 한 줄 표시 (작가 친화 UI).
#                    ADD/REWRITE/DELETE 한국어 라벨 + 통계 표시.
# v3.2 (2026-05-04)  ★ Writer Engine v3.5.1 동기화 — 7개 항목 일괄 동기화.
#                    1. DOCX 줄간격 표준화 (3중 안전망 — _normalize_screenplay_blank_lines + prev_block_type + add_blank_line)
#                    2. A29: 시간 정밀 표기 금지 (Writer A19 강화)
#                    3. A30: Character Voice Differentiation
#                    4. A31: Midpoint Erosion 차단
#                    5. A32: POV Rotation Enforcer
#                    6. ROMCOM Obstacle Intensity 모듈 (만남 빈도 50% 상한 + 장벽 4유형)
#                    7. SPACE_DIVERSITY_CHECK 메타 마커 차단
# v3.2.1 (2026-05-04) ★ 핫픽스 — show_step_2_revise()의 TypeError 수정.
#                    proposed_direction/preservation_notes가 None일 때 슬라이싱 실패 방지.
#                    부분 수정 모드에서 발생한 버그 해결.
# v3.3 (2026-05-04)  ★ Round N+1 사이클 시스템 — 검증 보고서 자동 흡수 + JSON 출력 + A29 자동 후처리.
#                    1. parse_verification_docx — 검증 보고서 DOCX 파서
#                       (✗[N] 미반영 / △[Partial] / A29~A32 위반 / 원본 유지 씬 / 재수정 권고 자동 추출)
#                    2. export_verify_json — 검증 결과를 JSON으로 내보내기 (다음 라운드 자동 흡수용)
#                    3. auto_fix_a29_violations — 집필 결과의 "한 박자/찰나/0.3초" 자동 치환
#                    4. UI: 검증 보고서 흡수 expander + Round 번호 + A29 자동 후처리 옵션
#                    5. 8점 도달 사이클: Round 1 결과 → 검증 → JSON 다운로드 → Round 2 흡수 → 8점
# v3.3.1 (2026-05-04) ★ UX 개선 — 피드백 자료 섹션 라벨 명확화.
#                    [1차 수정용] Rewrite Engine JSON / [N차 수정용] Revise 검증 보고서 명시.
#                    피드백 자료 헤더 아래 워크플로우 안내 박스 추가.
#                    Rewrite JSON 캡션에 MOON 자산 분리 명시 (전략 흡수 / 원고만 제외).
# v3.3.2 (2026-05-04) ★ 파일명 백업 관리 강화 — 모든 출력 파일에 라운드·시·분·점수 명시.
#                    형식: 제목_종류_R(라운드)_score(점수)_YYYYMMDD_HHMM_Blue.확장자
#                    같은 날 여러 라운드 작업해도 파일 자동 구분.
#                    수정본 DOCX, 검증 DOCX, 검증 JSON, 진행 백업, 전체 JSON 모두 통일.
# v3.3.3 (2026-05-04) ★ 핫픽스 — 검증 보고서 JSON 빈 데이터 버그 수정.
#                    원인: export_verify_json이 verify_result 최상위에서 키를 찾았으나,
#                    실제 LLM 응답은 verify_result["verify_report"][...]에 데이터 저장됨.
#                    해결: verify_report 래핑 자동 인식 + 다중 키명 폴백 체인.
#                    이제 점수·판정·항목·위반·권고 모두 정상 추출.
# v3.3.4 (2026-05-04) ★ 핫픽스 — DOCX 빌더 두 가지 버그 수정.
#                    1. 캐릭터명 정규식 강화 — 미래(F), 수강생1, 참석자2 등이
#                       대사로 인식 안 되어 지문 스타일로 폴백되던 문제 해결.
#                       (숫자·괄호 마커 F/M/TEL/전화/VO/OS 추가 인식)
#                    2. 빈 씬 자동 제거 — LLM이 "통합" 처방 시 한쪽 씬을 비워서
#                       헤더만 남기는 패턴(예: S#16 본문 없이 헤더만) 자동 필터링.
# v3.3.5 (2026-05-04) ★ A33 신규 룰 + 자동 후처리 — Verbatim Repetition Prohibition.
#                    「오랜만에」 Round 2에서 발견된 LLM 출력 결함 3가지 방어:
#                    1. 같은 씬 통째로 중복 출력 (S#36-2, S#44 케이스) → 자동 제거
#                    2. 같은 씬 내 동일 대사 중복 (S#17 "알아요." 케이스) → 자동 제거
#                    3. 작품 전체 verbatim 반복 (수익률 34 케이스) → 검출·경고
#                    A33 룰 prompt.py 추가 + auto_fix_a33_violations 함수 +
#                    auto_fix_duplicate_scene_blocks 함수.
# Pipeline: DIAGNOSE → REVISE → VERIFY
# Models: Opus 4.6 (집필) / Sonnet 4.6 (분석)
# =================================================================

import os
import re
import json
import io
from datetime import datetime

import streamlit as st
import anthropic
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from prompt import (
    SYSTEM_PROMPT,
    GENRE_RULES,
    INTENSITY_RULES,
    build_diagnose_prompt,
    build_revise_prompt,
    build_verify_prompt,
    get_report_filename,
    get_period_keys_for_ui,
    get_period_labels_for_ui,
    parse_rewrite_engine_json,
    split_into_batches,
    merge_batch_results,
    # v2.0 신규
    build_tone_dna_extraction_prompt,
    build_diff_analysis_prompt,
    build_distribution_diagnostic_prompt,
    absorb_rewrite_engine_metadata,
    # v2.1 신규
    build_genre_dna_extraction_prompt,
    # v2.2 신규
    build_section_detection_prompt,
    build_section_detect_step1_prompt,
    derive_section_ranges_from_step1,
    build_boundary_smoothness_block,
    build_cascade_analysis_prompt,
    # v2.8 신규 — Beat-Aware Diagnose
    build_beat_mapping_prompt,
    distribute_added_scenes_across_beats,
    SAVE_THE_CAT_15_BEATS,
)

# =================================================================
# [0] 모델 설정 & 페이지 설정
# =================================================================
MODEL_WRITE   = "claude-opus-4-6"      # Stage 2: 실제 집필
MODEL_ANALYZE = "claude-sonnet-4-6"    # Stage 1 & 3: 분석

st.set_page_config(
    page_title="BLUE JEANS · Revise Engine",
    page_icon="👖",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# =================================================================
# [1] 디자인 시스템 (Writer/Rewrite Engine과 동일 톤)
# =================================================================
st.markdown("""
<style>
@import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');
@import url('https://cdn.jsdelivr.net/gh/projectnoonnu/2408-3@latest/Paperlogy.css');
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@700;900&display=swap');

:root {
    --navy: #191970; --y: #FFCB05; --bg: #F7F7F5;
    --card: #FFFFFF; --card-border: #E2E2E0; --t: #1A1A2E;
    --g: #2EC484; --r: #E14444; --dim: #8E8E99; --light-bg: #EEEEF6;
    --serif: 'Paperlogy', 'Noto Serif KR', 'Georgia', serif;
    --display: 'Playfair Display', 'Paperlogy', 'Georgia', serif;
    --body: 'Pretendard', -apple-system, sans-serif;
    --heading: 'Paperlogy', 'Pretendard', sans-serif;
}

html, body, [class*="css"] {
    font-family: var(--body); color: var(--t); -webkit-font-smoothing: antialiased;
}
.stApp, [data-testid="stAppViewContainer"], [data-testid="stMain"],
[data-testid="stMainBlockContainer"], [data-testid="stHeader"],
[data-testid="stBottom"] {
    background-color: var(--bg) !important; color: var(--t) !important;
}
.stMarkdown, .stText, .stCode { color: var(--t) !important; }
h1, h2, h3, h4, h5, h6 { color: var(--navy) !important; font-family: var(--heading) !important; }
p, span, label, div, li { color: inherit; }
section[data-testid="stSidebar"] { display: none; }

.stTextInput input, .stTextArea textarea,
[data-testid="stTextInput"] input, [data-testid="stTextArea"] textarea {
    background-color: var(--card) !important; color: var(--t) !important;
    border: 1.5px solid var(--card-border) !important; border-radius: 8px !important;
    font-family: var(--body) !important; font-size: 0.92rem !important;
    padding: 0.65rem 0.85rem !important;
}
.stTextInput input:focus, .stTextArea textarea:focus {
    border-color: var(--navy) !important;
    box-shadow: 0 0 0 2px rgba(25,25,112,0.08) !important;
}

[data-testid="stSelectbox"] > div > div,
[data-testid="stSelectbox"] [data-baseweb="select"] > div {
    background-color: var(--card) !important; color: var(--t) !important;
    border: 1.5px solid var(--card-border) !important; border-radius: 8px !important;
}

.stButton > button {
    background-color: var(--navy) !important; color: #FFFFFF !important;
    border: none !important; border-radius: 8px !important;
    font-family: var(--heading) !important; font-weight: 700 !important;
    padding: 0.7rem 1.4rem !important; letter-spacing: 0.02em !important;
    transition: all 0.2s ease !important;
}
.stButton > button:hover {
    background-color: var(--y) !important; color: var(--navy) !important;
    transform: translateY(-1px);
}

[data-testid="stDownloadButton"] button {
    background-color: var(--y) !important; color: var(--navy) !important;
    border: none !important; border-radius: 8px !important;
    font-family: var(--heading) !important; font-weight: 800 !important;
    padding: 0.75rem 1.4rem !important;
}
[data-testid="stDownloadButton"] button:hover {
    background-color: var(--navy) !important; color: var(--y) !important;
}

[data-testid="stFileUploader"] {
    background-color: var(--card) !important;
    border: 2px dashed var(--card-border) !important;
    border-radius: 10px !important; padding: 1rem !important;
}

.rev-hero {
    text-align: center;
    padding: 52px 0 28px;
    border-bottom: 1px solid var(--card-border);
    margin-bottom: 40px;
}
.rev-hero .brand {
    font-size: 0.85rem; font-weight: 700;
    color: var(--navy); letter-spacing: 0.15em;
    font-family: var(--heading);
}
.rev-hero .title {
    font-size: 2.6rem; font-weight: 900; color: var(--navy);
    font-family: var(--display); letter-spacing: -0.02em;
    position: relative; display: inline-block;
    line-height: 1; margin: 14px 0 0 0;
}
.rev-hero .title::after {
    content: ''; position: absolute; bottom: 2px; left: 0;
    width: 100%; height: 4px; background: var(--y); border-radius: 2px;
}
.rev-hero .tag {
    font-size: 0.7rem; color: var(--dim);
    letter-spacing: 0.15em;
    margin-top: 0.6rem;
}

.rev-stepbar {
    display: flex; justify-content: space-between;
    background: var(--card); border: 1px solid var(--card-border);
    border-radius: 12px; padding: 16px 24px; margin-bottom: 32px;
}
.rev-step {
    flex: 1; text-align: center;
    font-family: var(--heading); font-weight: 700;
    color: var(--dim); font-size: 0.88rem;
    position: relative;
}
.rev-step.active { color: var(--navy); }
.rev-step.done { color: var(--g); }
.rev-step .num {
    display: inline-block; width: 28px; height: 28px;
    line-height: 28px; border-radius: 50%;
    background: var(--light-bg); color: var(--dim);
    font-weight: 900; margin-right: 8px;
}
.rev-step.active .num { background: var(--navy); color: var(--y); }
.rev-step.done .num { background: var(--g); color: #FFFFFF; }

.rev-card {
    background: var(--card); border: 1px solid var(--card-border);
    border-radius: 12px; padding: 24px; margin-bottom: 20px;
}
.rev-card-title {
    font-family: var(--heading); font-weight: 800;
    color: var(--navy); font-size: 1.15rem;
    margin-bottom: 12px; letter-spacing: -0.01em;
}
.rev-caption {
    color: var(--dim); font-size: 0.85rem;
    margin-top: 4px; margin-bottom: 16px;
}

.rev-badge {
    display: inline-block; padding: 4px 10px;
    background: var(--light-bg); color: var(--navy);
    border-radius: 6px; font-size: 0.78rem; font-weight: 700;
    font-family: var(--heading); margin-right: 6px;
}
.rev-badge.y { background: var(--y); color: var(--navy); }
.rev-badge.g { background: var(--g); color: #FFFFFF; }
.rev-badge.r { background: var(--r); color: #FFFFFF; }

.rev-verdict {
    display: inline-block; padding: 10px 20px;
    border-radius: 8px; font-family: var(--heading);
    font-weight: 900; font-size: 1rem; letter-spacing: 0.05em;
}
.rev-verdict.approved { background: var(--g); color: #FFFFFF; }
.rev-verdict.needs    { background: var(--y); color: var(--navy); }
.rev-verdict.rejected { background: var(--r); color: #FFFFFF; }
</style>
""", unsafe_allow_html=True)

# =================================================================
# [2] 세션 상태 초기화
# =================================================================
INIT_STATE = {
    "step": 0,                  # 0:입력 / 1:DIAGNOSE / 2:REVISE / 3:VERIFY / 4:완료
    "title": "",
    "raw_text": "",             # 원본 시나리오 (DOCX에서 추출)
    "raw_filename": "",
    "instruction": "",          # 수정 지시문
    "locked": "",               # LOCKED 요소
    "profession_input": "",     # 주요 캐릭터 직업 (선택사항)
    "period_key": "(현대)",     # 시대 (사극·시대극)
    "historical_type": "정통",  # 역사영화 유형 (정통/팩션/퓨전)
    "fact_based": False,        # 실화 기반 작품 여부
    "genre": "드라마",
    "intensity": "BALANCED",
    "diagnose_result": None,    # Stage 1 JSON 결과
    "revise_batches": None,     # 배치 분할 결과 (list)
    "batch_results": {},        # {batch_index: revise_result, ...}
    "current_batch": 0,         # 현재 처리 중인 배치 (1부터)
    "batch_size": 6,            # 한 배치당 씬 개수 (REVISE 단계, 수정 대상 씬 기준)
    "diagnose_batch_size": 12,  # ★ v2.7 — DIAGNOSE 자동 분할 배치 사이즈 (시나리오 씬 기준)
    # v2.8 — Beat-Aware Diagnose
    "target_added_scenes": 0,    # 추가할 총 씬 수 (0이면 비트 인식 OFF, 양수면 ON)
    "beat_map": None,            # Pre-Diagnose 비트 매핑 결과
    "beat_distribution": None,   # 비트별 추가 씬 분배 결과
    "revise_result": None,      # Stage 2 통합 결과
    "verify_result": None,      # Stage 3 JSON 결과
    # v2.0 — 톤 레퍼런스 + Diff 학습 + Rewrite 메타
    "tone_ref_text": "",        # 톤 레퍼런스 DOCX의 전문
    "tone_ref_filename": "",
    "diff_orig_text": "",       # Diff 모드 — 원본 (이전 버전)
    "diff_orig_filename": "",
    "diff_refined_text": "",    # Diff 모드 — 손본본 (최신 버전)
    "diff_refined_filename": "",
    "rewrite_json_text": "",    # Rewrite Engine JSON 원문 (변환 + 흡수)
    "tone_dna": None,           # 추출된 톤 DNA (DIAGNOSE 자동 실행)
    "diff_analysis": None,      # Diff 학습 결과
    "distribution_diagnostic": None,  # 분포 진단
    "rewrite_metadata": None,   # Rewrite 메타 흡수
    # ★ v3.3 — Round N 검증 보고서 흡수
    "verify_report_text": "",     # 검증 보고서 통합 지시문
    "verify_report_metadata": None,  # 파싱된 검증 보고서 dict
    "round_n": 1,                  # 현재 라운드 번호 (1=초회, 2=Round 2 ...)
    "auto_fix_a29_enabled": True,  # 집필 후 A29 자동 후처리 ON/OFF
    # v2.1 — 장르 DNA (참고작 1~3편에서 추출)
    "genre_ref_texts": [],      # [참고작1 텍스트, 참고작2, ...] 최대 3편
    "genre_ref_filenames": [],  # 파일명 리스트
    "genre_dna": None,          # 추출된 장르 DNA
    # v2.1 — Diff 모드: Before로 원본을 자동 사용할지 옵션
    "diff_use_main_as_before": True,
    # v2.2 — 구간 지정 모드 (이어쓰기 + 부분 수정 통합)
    "work_mode": None,               # "full" / "continuation" / "partial" (카드로 선택)
    "section_mode": False,           # 내부 호환 (continuation/partial일 때 자동 True)
    "section_input_method": "auto",  # auto (자동 감지) / manual (수동 지정) / hybrid
    "section_detection": None,       # 자동 감지 결과
    "protected_ranges": [],          # [{"from":"S#1","to":"S#25","reason":"..."}, ...]
    "revision_ranges": [],           # [{"from":"S#26","to":"S#71","reason":"..."}, ...]
    "cascade_analysis": None,        # 연쇄 영향 분석 결과
    "boundary_info": "",             # 경계 매끄러움 정보 (자동 계산)
    "show_advanced": False,          # 고급 옵션 펼치기
}

for k, v in INIT_STATE.items():
    if k not in st.session_state:
        st.session_state[k] = v


def reset_workflow():
    """전체 워크플로우 리셋."""
    for k, v in INIT_STATE.items():
        st.session_state[k] = v


# =================================================================
# [2-B] 통합 파일 추출 헬퍼 — DOCX + PDF
# =================================================================
def extract_text_from_uploaded_file(uploaded_file) -> str:
    """업로드된 파일(DOCX 또는 PDF)에서 텍스트 추출.

    Args:
        uploaded_file: st.file_uploader가 반환한 파일 객체

    Returns:
        추출된 전체 텍스트. 실패 시 빈 문자열.
    """
    if uploaded_file is None:
        return ""

    name = uploaded_file.name.lower()

    # DOCX
    if name.endswith(".docx"):
        try:
            from docx import Document as _Doc
            _doc = _Doc(uploaded_file)
            return "\n".join(p.text for p in _doc.paragraphs if p.text.strip())
        except Exception as e:
            st.error(f"DOCX 읽기 실패 ({uploaded_file.name}): {e}")
            return ""

    # PDF
    elif name.endswith(".pdf"):
        try:
            from pypdf import PdfReader
            reader = PdfReader(uploaded_file)
            text_parts = []
            for page in reader.pages:
                t = page.extract_text() or ""
                if t.strip():
                    text_parts.append(t)
            full_text = "\n".join(text_parts)
            if not full_text.strip():
                st.warning(f"⚠️ PDF에서 텍스트를 추출하지 못했습니다. 스캔본/이미지 PDF는 지원되지 않습니다 ({uploaded_file.name})")
            return full_text
        except Exception as e:
            st.error(f"PDF 읽기 실패 ({uploaded_file.name}): {e}")
            return ""

    else:
        st.error(f"지원하지 않는 파일 형식: {uploaded_file.name} (DOCX 또는 PDF만 가능)")
        return ""


# =================================================================
# [3] API 클라이언트 & 호출
# =================================================================
def get_client():
    api_key = st.secrets.get("ANTHROPIC_API_KEY", os.getenv("ANTHROPIC_API_KEY"))
    if not api_key:
        st.error("❌ ANTHROPIC_API_KEY가 secrets에 없습니다. "
                 ".streamlit/secrets.toml 또는 환경변수에 추가해주세요.")
        return None
    return anthropic.Anthropic(api_key=api_key)


def call_claude(client, prompt_text: str, model: str, max_tokens: int = 32000, retries: int = 1):
    """Claude API 스트리밍 호출 + max_tokens 잘림 시 1회만 자동 증량 재시도.

    재시도 제한: 기본 1회 (토큰 낭비 방지).
    잘림이 반복되면 사용자에게 알리고 중단하여 결과를 그대로 반환.

    모델별 한도:
    - Sonnet 4.6: 최대 64,000 토큰
    - Opus 4.6:   최대 32,000 토큰
    """
    # 모델별 상한
    if "opus" in model.lower():
        absolute_cap = 32000
    else:  # sonnet, haiku 등
        absolute_cap = 64000

    current_tokens = min(max_tokens, absolute_cap)

    for attempt in range(1 + retries):
        try:
            collected = ""
            stop_reason = None
            with client.messages.stream(
                model=model,
                max_tokens=current_tokens,
                system=SYSTEM_PROMPT,
                messages=[{"role": "user", "content": prompt_text}],
            ) as stream:
                for text in stream.text_stream:
                    collected += text
                final_msg = stream.get_final_message()
                stop_reason = final_msg.stop_reason

            if stop_reason == "max_tokens":
                if attempt < retries and current_tokens < absolute_cap:
                    next_tokens = min(int(current_tokens * 1.5), absolute_cap)
                    if next_tokens == current_tokens:
                        # 이미 한도에 도달 — 즉시 중단
                        st.warning(f"⚠️ 모델 한도({absolute_cap:,} 토큰)에서 응답이 잘렸습니다. "
                                   "지금 결과로 진행합니다. 이상하면 시나리오를 더 짧게 분할하거나 "
                                   "지시문을 간소화하세요.")
                        return collected
                    st.info(f"🔄 응답이 {current_tokens:,} 토큰에서 잘렸습니다. "
                            f"{next_tokens:,} 토큰으로 1회만 재시도합니다.")
                    current_tokens = next_tokens
                    continue
                else:
                    # 1회 재시도 후에도 잘렸으면 그대로 진행 (추가 재시도 안 함)
                    st.warning(f"⚠️ 재시도 후에도 응답이 잘렸습니다 ({current_tokens:,} 토큰). "
                               "지금 결과로 진행합니다 — 추가 재시도는 토큰 낭비입니다. "
                               "결과가 불완전하면 입력을 간소화하세요.")
            return collected
        except Exception as e:
            st.error(f"❌ API 오류: {type(e).__name__} — {e}")
            return None
    return None


def parse_json(raw: str):
    """JSON 파싱 (마크다운 코드블록 제거 + 양끝 트리밍 + 잘린 JSON 복구 시도)."""
    if not raw:
        return None
    txt = raw.strip()
    # 코드블록 제거
    txt = re.sub(r'^```(?:json)?\s*\n', '', txt)
    txt = re.sub(r'\n```\s*$', '', txt)
    # 첫 { 부터 시작
    s = txt.find('{')
    if s == -1:
        return None

    # 1. 정상 파싱 시도 — 마지막 } 까지
    e = txt.rfind('}')
    if e > s:
        try:
            return json.loads(txt[s:e+1])
        except json.JSONDecodeError:
            pass  # 복구 시도로 진행

    # 2. 잘린 JSON 복구 시도 — 괄호 균형 맞추기
    candidate = txt[s:]
    open_braces = 0
    open_brackets = 0
    in_string = False
    escape_next = False
    last_valid_pos = -1

    for i, ch in enumerate(candidate):
        if escape_next:
            escape_next = False
            continue
        if ch == '\\' and in_string:
            escape_next = True
            continue
        if ch == '"' and not escape_next:
            in_string = not in_string
            continue
        if in_string:
            continue
        if ch == '{':
            open_braces += 1
        elif ch == '}':
            open_braces -= 1
            if open_braces == 0 and open_brackets == 0:
                last_valid_pos = i
        elif ch == '[':
            open_brackets += 1
        elif ch == ']':
            open_brackets -= 1

    # 마지막 유효 위치까지로 시도
    if last_valid_pos > 0:
        try:
            return json.loads(candidate[:last_valid_pos+1])
        except json.JSONDecodeError:
            pass

    # 3. 잘린 끝부분 강제 닫기 시도
    truncated = candidate.rstrip().rstrip(',')
    # 마지막에 미완성 문자열이 있으면 닫기
    if truncated.count('"') % 2 == 1:
        truncated += '"'
    # 열린 배열/객체 닫기
    truncated += ']' * max(0, open_brackets) + '}' * max(0, open_braces)

    try:
        result = json.loads(truncated)
        st.warning("⚠️ 응답이 잘렸지만 부분 복구를 시도했습니다. 결과 일부가 누락될 수 있습니다.")
        return result
    except json.JSONDecodeError as err:
        st.error(f"❌ JSON 파싱 실패: {err}")
        st.caption("응답 끝부분 (디버깅용):")
        st.code(candidate[-500:], language="json")
        return None


# =================================================================
# [4] DOCX 입력: 원본 시나리오 파싱
# =================================================================
def extract_docx_text(uploaded_file) -> str:
    """업로드된 DOCX에서 본문 텍스트를 추출."""
    try:
        doc = Document(uploaded_file)
        paragraphs = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                paragraphs.append(t)
        # 표 안의 텍스트도 수집 (씬 헤더·대사 표로 정리된 경우 대비)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        t = p.text.strip()
                        if t:
                            paragraphs.append(t)
        return "\n".join(paragraphs)
    except Exception as e:
        st.error(f"DOCX 추출 실패: {e}")
        return ""


# =================================================================
# [5] DOCX 출력: 수정본 & 검증 보고서
# =================================================================
def _set_font(run, font_name="맑은 고딕", size_pt=10.5, bold=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def create_revised_docx(revise_result: dict, title: str = "", genre: str = "",
                        original_text: str = "",
                        fact_based: bool = False,
                        historical: bool = False,
                        historical_type: str = "") -> bytes:
    """Stage 2 결과를 한국 시나리오 표준 서식의 DOCX로 변환.

    Writer Engine과 동일한 서식 사용:
    - 함초롬바탕 10pt 기본
    - 씬번호 / 대사 / 대사연속 / 지문 4가지 Word 스타일
    - A4, 20mm 여백
    - 캐릭터명\t\t대사 형식 자동 감지
    - 메타데이터 자동 차단
    """
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Mm, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from io import BytesIO

    doc = DocxDocument()

    # ── 페이지 설정 (A4, 20mm 마진) ──
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)

    # ── 기본 스타일: 함초롬바탕 10pt ──
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "함초롬바탕"
    style_normal.font.size = Pt(10)
    style_normal.paragraph_format.space_after = Pt(2)
    style_normal.paragraph_format.space_before = Pt(0)
    rpr = style_normal.element.rPr
    if rpr is None:
        rpr = style_normal.element.makeelement(qn('w:rPr'), {})
        style_normal.element.append(rpr)
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), '함초롬바탕')

    def _set_eastasia_font(rpr_elem, font_name='함초롬바탕'):
        rf = rpr_elem.find(qn('w:rFonts'))
        if rf is None:
            rf = rpr_elem.makeelement(qn('w:rFonts'), {})
            rpr_elem.append(rf)
        rf.set(qn('w:eastAsia'), font_name)

    # ── 커스텀 스타일: 씬번호 / 대사 / 대사연속 / 지문 ──
    style_scene = doc.styles.add_style('씬번호', WD_STYLE_TYPE.PARAGRAPH)
    style_scene.base_style = doc.styles['Normal']
    style_scene.font.name = '함초롬바탕'
    style_scene.font.size = Pt(11)
    style_scene.font.bold = True
    style_scene.paragraph_format.space_before = Pt(24)
    style_scene.paragraph_format.space_after = Pt(6)
    style_scene.paragraph_format.line_spacing = 1.5
    _set_eastasia_font(style_scene.element.get_or_add_rPr())

    style_dialogue = doc.styles.add_style('대사', WD_STYLE_TYPE.PARAGRAPH)
    style_dialogue.base_style = doc.styles['Normal']
    style_dialogue.font.name = '함초롬바탕'
    style_dialogue.font.size = Pt(10)
    style_dialogue.font.bold = True
    style_dialogue.paragraph_format.left_indent = Cm(1.25)
    style_dialogue.paragraph_format.space_before = Pt(8)
    style_dialogue.paragraph_format.space_after = Pt(2)
    style_dialogue.paragraph_format.line_spacing = 1.5
    _set_eastasia_font(style_dialogue.element.get_or_add_rPr())

    style_dialogue_cont = doc.styles.add_style('대사연속', WD_STYLE_TYPE.PARAGRAPH)
    style_dialogue_cont.base_style = style_dialogue
    style_dialogue_cont.paragraph_format.space_before = Pt(0)
    style_dialogue_cont.paragraph_format.space_after = Pt(0)

    style_action = doc.styles.add_style('지문', WD_STYLE_TYPE.PARAGRAPH)
    style_action.base_style = doc.styles['Normal']
    style_action.font.name = '함초롬바탕'
    style_action.font.size = Pt(10)
    style_action.font.bold = False
    style_action.paragraph_format.space_before = Pt(2)
    style_action.paragraph_format.space_after = Pt(2)
    _set_eastasia_font(style_action.element.get_or_add_rPr())

    # ★ Writer Engine v3.1.4 자산 — INSERT 전용 스타일 3종 이식
    # [스타일 5] INSERT 헤더 — 작은 대문자 느낌, 중간 들여쓰기, 굵게
    style_insert_header = doc.styles.add_style('인서트헤더', WD_STYLE_TYPE.PARAGRAPH)
    style_insert_header.base_style = doc.styles['Normal']
    style_insert_header.font.name = '함초롬바탕'
    style_insert_header.font.size = Pt(9)
    style_insert_header.font.bold = True
    style_insert_header.paragraph_format.left_indent = Cm(2.55)
    style_insert_header.paragraph_format.space_before = Pt(8)
    style_insert_header.paragraph_format.space_after = Pt(2)
    _set_eastasia_font(style_insert_header.element.get_or_add_rPr())

    # [스타일 6] INSERT 본문 — 깊은 들여쓰기, 이탤릭
    style_insert_body = doc.styles.add_style('인서트본문', WD_STYLE_TYPE.PARAGRAPH)
    style_insert_body.base_style = doc.styles['Normal']
    style_insert_body.font.name = '함초롬바탕'
    style_insert_body.font.size = Pt(10)
    style_insert_body.font.italic = True
    style_insert_body.paragraph_format.left_indent = Cm(2.55)
    style_insert_body.paragraph_format.space_before = Pt(2)
    style_insert_body.paragraph_format.space_after = Pt(2)
    style_insert_body.paragraph_format.line_spacing = 1.4
    _set_eastasia_font(style_insert_body.element.get_or_add_rPr())

    # [스타일 7] INSERT 라벨식 — 한 줄 짜리 [라벨] '본문' 형식
    style_insert_label = doc.styles.add_style('인서트라벨', WD_STYLE_TYPE.PARAGRAPH)
    style_insert_label.base_style = doc.styles['Normal']
    style_insert_label.font.name = '함초롬바탕'
    style_insert_label.font.size = Pt(10)
    style_insert_label.paragraph_format.left_indent = Cm(1.42)
    style_insert_label.paragraph_format.space_before = Pt(4)
    style_insert_label.paragraph_format.space_after = Pt(4)
    _set_eastasia_font(style_insert_label.element.get_or_add_rPr())

    # ── 헬퍼 ──
    def add_text(text, bold=False, size=None, color=None, align=None):
        p = doc.add_paragraph()
        if align:
            p.alignment = align
        r = p.add_run(text)
        r.font.name = "함초롬바탕"
        _set_eastasia_font(r._element.get_or_add_rPr())
        if bold:
            r.bold = True
        if size:
            r.font.size = size
        if color:
            r.font.color.rgb = color
        return p

    def add_scene_heading(text):
        p = doc.add_paragraph(style='씬번호')
        r = p.add_run(text)
        r.font.name = "함초롬바탕"
        _set_eastasia_font(r._element.get_or_add_rPr())
        return p

    def add_blank_line():
        """★ v3.2 — 지문↔대사 사이 빈 줄 (Writer Engine v3.5.1 자산 이식).
        함초롬바탕 10pt로 통일된 빈 단락. 워드/한글 모두 일관된 높이."""
        from docx.shared import Pt as _Pt
        p = doc.add_paragraph(style='지문')
        r = p.add_run("")
        r.font.name = "함초롬바탕"
        r.font.size = _Pt(10)
        _set_eastasia_font(r._element.get_or_add_rPr())
        return p

    def add_dialogue(char_name, parenthetical, line, continuation=False):
        """Writer Engine v3.1.3 자산 — 대사 본문 내 괄호 지시문은
        run을 분할하여 bold를 해제. 화자 표기의 괄호(예: V.O.)는 유지."""
        if continuation:
            p = doc.add_paragraph(style='대사연속')
            speaker_part = "\t\t"
        else:
            p = doc.add_paragraph(style='대사')
            speaker_part = f"{char_name}\t\t"

        # 대사 본문 영역 조립: parenthetical 인자 + 실제 대사
        body_parts = []  # [(text, is_paren), ...]
        if parenthetical:
            body_parts.append((f"({parenthetical}) ", True))
        if line:
            import re as _re_dlg
            chunks = _re_dlg.split(r'(\([^()]*\))', line)
            for chunk in chunks:
                if not chunk:
                    continue
                if chunk.startswith('(') and chunk.endswith(')'):
                    body_parts.append((chunk, True))
                else:
                    body_parts.append((chunk, False))

        # run 1: 화자 영역
        r_speaker = p.add_run(speaker_part)
        r_speaker.font.name = "함초롬바탕"
        _set_eastasia_font(r_speaker._element.get_or_add_rPr())

        # run 2~N: 대사 본문 — 괄호 부분은 bold=False
        for text, is_paren in body_parts:
            r = p.add_run(text)
            r.font.name = "함초롬바탕"
            _set_eastasia_font(r._element.get_or_add_rPr())
            if is_paren:
                r.bold = False

        return p

    def add_insert_block(header: str, body_lines: list):
        """Writer Engine v3.1.4 자산 — INSERT 형식 A 렌더링.
        헤더(작게·굵게·들여쓰기) + 본문(이탤릭·깊은 들여쓰기) + 자동 빈 줄.
        """
        doc.add_paragraph("")
        first_p = doc.add_paragraph(style='인서트헤더')
        r = first_p.add_run(header.strip())
        r.font.name = "함초롬바탕"
        _set_eastasia_font(r._element.get_or_add_rPr())

        for line in body_lines:
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph(style='인서트본문')
            r = p.add_run(line)
            r.font.name = "함초롬바탕"
            r.italic = True
            _set_eastasia_font(r._element.get_or_add_rPr())

        close_p = doc.add_paragraph(style='인서트헤더')
        cr = close_p.add_run('[/INSERT]')
        cr.font.name = "함초롬바탕"
        _set_eastasia_font(cr._element.get_or_add_rPr())
        doc.add_paragraph("")
        return first_p

    def add_insert_label_paragraph(text: str):
        """Writer Engine v3.1.4 자산 — 형식 B 라벨 한 줄 렌더링."""
        label, body = _parse_insert_label(text)
        p = doc.add_paragraph(style='인서트라벨')

        r_label = p.add_run(label + ' ')
        r_label.font.name = "함초롬바탕"
        r_label.font.size = Pt(9)
        r_label.bold = True
        _set_eastasia_font(r_label._element.get_or_add_rPr())

        if body:
            r_body = p.add_run(body)
            r_body.font.name = "함초롬바탕"
            r_body.italic = True
            _set_eastasia_font(r_body._element.get_or_add_rPr())
        return p

    def add_action(text):
        """Writer Engine v3.1.3+v3.1.4 자산 — 지문 출력.
        - 긴 단락은 의미 비트 단위로 자동 분단 (_split_action_paragraph)
        - INSERT 블록(형식 A·B) 자동 감지 → 전용 스타일로 분기
        - AI 시적 의도 보존: 짧은 단락(150자 미만, 7문장 미만)은 그대로 둠
        """
        # PROP 메모·CHECK 태그 정제 (안전망)
        text = _strip_prop_state_memos(text)

        # INSERT 블록 우선 분리
        items = _parse_insert_blocks(text)

        first_p = None
        for item in items:
            if item['type'] == 'insert_block':
                p = add_insert_block(item['data']['header'], item['data']['body'])
            elif item['type'] == 'insert_label':
                p = add_insert_label_paragraph(item['data'])
            else:
                # 일반 지문 — Writer Engine 분단 알고리즘 적용
                sub_paragraphs = _split_action_paragraph(item['data'])
                p = None
                for sub in sub_paragraphs:
                    sp = doc.add_paragraph(style='지문')
                    r = sp.add_run(sub)
                    r.font.name = "함초롬바탕"
                    _set_eastasia_font(r._element.get_or_add_rPr())
                    if p is None:
                        p = sp
            if first_p is None:
                first_p = p
        return first_p

    # ── 커버 페이지 ──
    for _ in range(6):
        doc.add_paragraph("")
    add_text("시나리오 (수정본)", size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph("")
    proj_title = title or f"<{genre}>"
    add_text(proj_title, bold=True, size=Pt(24), align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph("")
    doc.add_paragraph("")
    add_text("기획/제작 | 블루진픽처스", size=Pt(10),
             align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x8E, 0x8E, 0x99))
    rr = revise_result.get("revision_result", {})
    scenes = rr.get("revised_scenes", [])
    add_text(f"Revise Engine v2.2  ·  {len(scenes)}개 씬 수정",
             size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER,
             color=RGBColor(0x8E, 0x8E, 0x99))
    doc.add_page_break()

    # ── 면책 자막 (실화/팩션·퓨전) ──
    _need_disclaimer = fact_based or (
        historical and ("팩션" in (historical_type or "") or "퓨전" in (historical_type or ""))
    )
    if _need_disclaimer:
        for _ in range(10):
            doc.add_paragraph("")
        add_text("본 작품에 등장하는 인물, 단체, 지명, 상호, 사건은",
                 size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER)
        add_text("모두 허구이며, 실존하는 것과 관련이 있더라도",
                 size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER)
        add_text("극적 구성을 위해 각색되었습니다.",
                 size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_page_break()

    # ─────────────────────────────────────────────────────────
    # 본문 파싱: 원본 + 수정본을 통합한 "최종 시나리오" 생성
    # ─────────────────────────────────────────────────────────
    # 1) 원본 시나리오를 씬 단위로 분할
    # 2) 수정 대상 씬은 revised_content로 교체
    # 3) ADD 타입은 insert_position에 따라 삽입
    # 4) DELETE 타입은 제거
    #
    # 결과: 시나리오 처음부터 끝까지 자연스럽게 흐르는 통합본
    # ─────────────────────────────────────────────────────────

    import re as _re

    # 씬 헤딩 패턴 (Writer Engine과 동일)
    heading_re = _re.compile(r'^S?#?\d*\.?\s*(INT\.|EXT\.|INT\./EXT\.)\s*(.+)', _re.IGNORECASE)


    def _merge_header_content(header: str, content: str) -> str:
        """헤더와 본문이 중복되지 않도록 합친다.
        본문 첫 줄에 이미 헤더 핵심부(S#번호 또는 INT./EXT.)가 있으면 헤더 생략."""
        if not content.strip():
            return header
        first_line = content.strip().split('\n')[0].strip()
        # S#번호 추출
        import re as _re_inner
        h_num = _re_inner.search(r'S#\d+', header or "")
        c_num = _re_inner.search(r'S#\d+', first_line)
        if h_num and c_num and h_num.group() == c_num.group():
            return content  # 본문이 이미 헤더 포함
        # INT./EXT. 매칭
        if _re_inner.match(r'^(INT\.|EXT\.)', first_line, _re_inner.IGNORECASE):
            # 첫 줄이 씬 헤더 형태 → 본문 그대로
            return content
        return f"{header}\n{content}"

    def split_into_scenes(text: str):
        """원본 텍스트를 [(scene_id, scene_body), ...] 리스트로 분할."""
        if not text:
            return []
        lines = text.split('\n')
        scenes_list = []
        current_id = ""
        current_body = []

        for line in lines:
            stripped = line.strip()
            # S#숫자 또는 INT./EXT. 패턴이 씬 헤딩
            if heading_re.match(stripped) or _re.match(r'^S#\d+', stripped):
                # 이전 씬 저장
                if current_id or current_body:
                    scenes_list.append((current_id, '\n'.join(current_body)))
                # 새 씬 시작
                current_id = stripped
                current_body = [line]
            else:
                current_body.append(line)

        # 마지막 씬
        if current_id or current_body:
            scenes_list.append((current_id, '\n'.join(current_body)))

        return scenes_list

    def extract_scene_number(scene_id_str: str) -> str:
        """씬 식별자에서 'S#숫자' 패턴만 추출 (매칭용)."""
        m = _re.search(r'S#(\d+)', scene_id_str)
        return f"S#{m.group(1)}" if m else scene_id_str.strip()[:20]

    # 수정본 씬을 scene_number로 인덱싱
    revised_by_id = {}     # {"S#1": revised_content}
    deleted_ids = set()    # {"S#5", ...}
    add_after = {}         # {"S#42": [{header, content}, ...]}  ADD 타입

    for sc in scenes:
        sid = sc.get("scene_id", "")
        s_type = sc.get("type", "REWRITE")
        s_header = sc.get("scene_header", sid)
        s_content = sc.get("revised_content", "")
        s_num = extract_scene_number(sid)

        if s_type == "DELETE":
            deleted_ids.add(s_num)
        elif s_type == "ADD":
            insert_after = sc.get("insert_position", "") or sid
            insert_num = extract_scene_number(insert_after)
            if insert_num not in add_after:
                add_after[insert_num] = []
            # ADD 씬은 헤더가 본문에 없을 수 있으니 합쳐서
            full_text = _merge_header_content(s_header, s_content)
            add_after[insert_num].append(full_text)
        else:
            # REWRITE / MERGE / SPLIT: 단순 교체
            full_text = _merge_header_content(s_header, s_content)
            revised_by_id[s_num] = full_text

    # 원본을 씬 단위로 분할 후 통합 시나리오 구성
    final_scenes = []
    if original_text:
        original_scenes = split_into_scenes(original_text)
        for orig_id, orig_body in original_scenes:
            orig_num = extract_scene_number(orig_id)
            if orig_num in deleted_ids:
                continue  # 삭제
            if orig_num in revised_by_id:
                final_scenes.append(revised_by_id[orig_num])  # 교체
            else:
                final_scenes.append(orig_body)  # 원본 유지
            # ADD 처리
            if orig_num in add_after:
                for added in add_after[orig_num]:
                    final_scenes.append(added)
    else:
        # 원본이 없으면 수정본만 출력 (배치 부분만)
        for sc in scenes:
            s_type = sc.get("type", "REWRITE")
            if s_type == "DELETE":
                continue
            s_header = sc.get("scene_header", "")
            s_content = sc.get("revised_content", "")
            full_text = _merge_header_content(s_header, s_content)
            final_scenes.append(full_text)

    # ★ v3.3.4: 빈 씬 자동 제거
    # 헤더(S#16. INT. ...)만 있고 본문이 없는 씬을 자동 필터링.
    # 원인: LLM이 "S#16과 S#22 통합" 처방을 받으면 S#16을 비우고 S#22로 통합하는데,
    # 빈 헤더만 남으면 시나리오 흐름이 끊김.
    _filtered_scenes = []
    _empty_scene_count = 0
    import re as _re_empty
    for _scene in final_scenes:
        if not _scene or not _scene.strip():
            _empty_scene_count += 1
            continue
        # 헤더 한 줄 + 빈 줄만 있는 경우 검사
        _scene_lines = [l for l in _scene.strip().split('\n') if l.strip()]
        if len(_scene_lines) == 0:
            _empty_scene_count += 1
            continue
        # 한 줄만 있고 그게 씬 헤더면 빈 씬으로 간주
        if len(_scene_lines) == 1 and _re_empty.match(r'^\s*S#\d+', _scene_lines[0]):
            _empty_scene_count += 1
            continue
        _filtered_scenes.append(_scene)
    if _empty_scene_count > 0:
        try:
            st.warning(
                f"⚠️ 빈 씬 자동 제거: {_empty_scene_count}개 씬이 헤더만 있고 본문 없음 → 자동 필터링됨. "
                f"(LLM이 '통합' 처방 시 한쪽 씬을 비우는 패턴 — 정상 처리)"
            )
        except Exception:
            # st 컨텍스트 외부에서 호출되면 통과
            pass

    full_text = '\n\n'.join(_filtered_scenes)

    # ★ v3.3.5: 통째로 중복된 씬 블록 자동 제거 (A33 후처리)
    full_text, _dup_removed = auto_fix_duplicate_scene_blocks(full_text)
    if _dup_removed > 0:
        try:
            st.warning(
                f"⚠️ A33 위반 자동 처리: {_dup_removed}개 씬이 통째로 중복 출력됨 → 두 번째 블록 자동 제거. "
                f"(LLM이 ADD 처방 시 같은 씬을 두 번 쓰는 결함 — 정상 처리)"
            )
        except Exception:
            pass

    # ─────────────────────────────────────────────────────────
    # 본문 라인 단위 파싱 (Writer Engine 동일 로직)
    # ─────────────────────────────────────────────────────────
    # ★ v3.3.4: 캐릭터명 인식 강화
    # - 숫자 허용 (수강생1, 참석자2, 학생3)
    # - 괄호 형식 (F)/(M) 마커 인식 (미래(F), 봉식(M) 등 — 통화 상대 표시)
    # - 하이픈/언더스코어 허용 (S#36-2 같은 경우)
    char_re = _re.compile(
        r'^\s{2,}([가-힣a-zA-Z0-9\s]{1,15}?)\s*'
        r'(?:\(([FMfm]|V\.O\.|O\.S\.|CONT\'D|cont\'d|v\.o\.|o\.s\.|TEL|전화|VO|OS)\))?\s*$',
        _re.IGNORECASE
    )
    inline_dialogue_re = _re.compile(
        r'^([가-힣a-zA-Z0-9\s]{1,15}?)\s*'
        r'(?:\(([FMfm]|V\.O\.|O\.S\.|CONT\'D|cont\'d|v\.o\.|o\.s\.|TEL|전화|VO|OS)\))?\s*'
        r'\t{1,}\s*(?:\(([^)]*)\)\s*)?(.+)',
        _re.IGNORECASE
    )
    paren_re = _re.compile(r'^\s{2,}\((.+?)\)\s*$')

    # 문자열 그대로의 \n이 들어온 경우(JSON 이스케이프 잔존) 안전 처리
    full_text = full_text.replace('\\n', '\n').replace('\\t', '\t')

    # ═══════════════════════════════════════════════════════════
    # 대사 형식 붕괴 자동 복구 (Writer Engine v3.4 이식)
    # 버그: 긴 컨텍스트에서 AI가 대사 포맷 규칙을 잊고
    #       "캐릭터\n\n대사" 형식으로 출력
    # 복구: "캐릭터" 단독 라인 + 빈 라인 + 대사 라인 → "캐릭터\t\t대사"
    # ═══════════════════════════════════════════════════════════
    _CHAR_NAMES = {
        '유진', '진호', '세웅', '다은', '강회장', '민준', '박지영', '오현수',
        '이진호', '반세웅', '김사장', '비서', '편집자', '기사', '배달 기사',
        '사장', '민준 엄마', '박씨', '엄마', '아빠', '형', '누나', '아들', '딸',
        '김 여사', '지우', '여름', '최여름', '안경 아이', '동생', '아내',
        '집배원', '중개인', '최상진', '조민준', '강유진',
    }
    _broken_lines = full_text.split("\n")
    _fixed_lines = []
    _j = 0
    while _j < len(_broken_lines):
        _cur = _broken_lines[_j].strip()
        # 패턴 A: "캐릭터명" 단독 + 빈줄 + 대사 → "캐릭터\t\t대사"
        if (_cur in _CHAR_NAMES and
            _j + 2 < len(_broken_lines) and
            _broken_lines[_j + 1].strip() == "" and
            _broken_lines[_j + 2].strip() and
            not _broken_lines[_j + 2].strip().startswith("S#") and
            _broken_lines[_j + 2].strip() not in _CHAR_NAMES):
            _next_content = _broken_lines[_j + 2].strip()
            # 괄호 지시(예: "(잠깐 생각하고)")가 있으면 다음 줄이 진짜 대사
            if _next_content.startswith("(") and _next_content.endswith(")") and \
               _j + 4 < len(_broken_lines) and _broken_lines[_j + 3].strip() == "" and \
               _broken_lines[_j + 4].strip():
                _fixed_lines.append(f"{_cur}\t\t{_next_content} {_broken_lines[_j + 4].strip()}")
                _j += 5
                continue
            _fixed_lines.append(f"{_cur}\t\t{_next_content}")
            _j += 3
            continue
        _fixed_lines.append(_broken_lines[_j])
        _j += 1

    # ★ v3.2 — 지문↔대사 빈 줄 후처리 (3차 안전망, Writer Engine v3.5.1 자산)
    _normalized_text = _normalize_screenplay_blank_lines("\n".join(_fixed_lines))
    lines = _normalized_text.split("\n")

    i = 0
    # ★ v3.2 — 지문↔대사 사이 빈 줄 자동 삽입을 위한 직전 블록 타입 추적
    # (Writer Engine v3.5.1 자산 이식)
    # 가능한 값: None, "scene", "action", "dialogue", "insert"
    prev_block_type = None
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # 마크다운 강조 기호 제거 (** ** 등)
        if stripped.startswith('**') and stripped.endswith('**'):
            stripped = stripped[2:-2].strip()
            line = stripped

        # 메타 라인 차단 (간단 버전 — 메타 마커가 들어오면 스킵)
        if stripped.startswith('▸') or stripped.startswith('━━━'):
            i += 1
            continue
        if '내부 메모' in stripped or 'WRITER_NOTES' in stripped:
            i += 1
            continue
        # ★ v3.2: SPACE_DIVERSITY_CHECK 마커 차단 (Writer v3.5 자산)
        if 'SPACE_DIVERSITY_CHECK' in stripped:
            i += 1
            continue

        # 씬 헤딩
        m = heading_re.match(stripped)
        if m or _re.match(r'^S#\d+', stripped):
            add_scene_heading(stripped)
            prev_block_type = "scene"  # ★ v3.2
            i += 1
            continue

        # 인라인 대사
        im = inline_dialogue_re.match(stripped)
        if im:
            char_name = im.group(1).strip()
            vo_marker = im.group(2) or ""
            inline_paren = im.group(3) or ""
            inline_text = im.group(4).strip()
            if vo_marker:
                char_name = f"{char_name} ({vo_marker})"
            # ★ v3.2: 지문/insert 직후 대사면 빈 줄 1개 삽입
            if prev_block_type in ("action", "insert"):
                add_blank_line()
            add_dialogue(char_name, inline_paren, inline_text)
            prev_block_type = "dialogue"  # ★ v3.2
            i += 1
            continue

        # 들여쓰기 캐릭터명 + 대사
        cm = char_re.match(line)
        if cm:
            char_name = cm.group(1).strip()
            vo_marker = cm.group(2) or ""
            if vo_marker:
                char_name = f"{char_name} ({vo_marker})"
            parenthetical = ""
            dialogue_lines = []
            # ★ v3.2: 지문/insert 직후 대사면 빈 줄 1개 삽입
            if prev_block_type in ("action", "insert"):
                add_blank_line()
            i += 1
            if i < len(lines):
                pm = paren_re.match(lines[i])
                if pm:
                    parenthetical = pm.group(1)
                    i += 1
            while i < len(lines):
                dl = lines[i]
                ds = dl.strip()
                if not ds:
                    break
                if heading_re.match(ds) or _re.match(r'^S#\d+', ds):
                    break
                if char_re.match(dl):
                    break
                if inline_dialogue_re.match(ds):
                    break
                dialogue_lines.append(ds)
                i += 1
            if dialogue_lines:
                merged = " ".join(dialogue_lines)
                add_dialogue(char_name, parenthetical, merged)
            prev_block_type = "dialogue"  # ★ v3.2
            continue

        # 그 외: 지문
        # ★ v3.2: 대사 직후 지문이면 빈 줄 1개 삽입
        if prev_block_type == "dialogue":
            add_blank_line()
        add_action(stripped)
        prev_block_type = "action"  # ★ v3.2
        i += 1

    # ── 바이트 반환 ──
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# =================================================================
# ★★★ v3.3 — 검증 보고서 DOCX 파서 + JSON 출력 + 자동 후처리 ★★★
# =================================================================

def _get_round_aware_filename(title: str, kind: str = "revised", ext: str = "docx") -> str:
    """v3.3.2 — 라운드·시·분 인지 파일명 생성.

    형식: 제목_종류_R(라운드)_score(점수)_YYYYMMDD_HHMM_Blue.ext
    예시: 오랜만에_가제_수정본_R2_score78_20260504_2305_Blue.docx
          테이스티_러브_검증보고서_R1_score68_20260504_2310_Blue.json

    score는 verify_result에서 자동 추출되며, 없으면 생략.
    """
    import re as _re_filename
    from datetime import datetime as _dt_filename

    safe_title = _re_filename.sub(r'[/*?:"<>|]', '_', title.strip()) if title else "제목없음"
    timestamp = _dt_filename.now().strftime("%Y%m%d_%H%M")

    kind_map = {
        "revised": "수정본",
        "verify":  "검증보고서",
        "diagnose": "수정플랜",
        "backup": "백업",
    }
    kind_kor = kind_map.get(kind, kind)

    # 라운드 번호 (있으면)
    try:
        round_n = st.session_state.get("round_n", 1)
        round_str = f"_R{round_n}"
    except Exception:
        round_str = ""

    # 점수 (있으면, verify_result에서 자동 추출)
    score_str = ""
    try:
        verify_result = st.session_state.get("verify_result", {})
        if verify_result and isinstance(verify_result, dict):
            score = verify_result.get("overall_score")
            if isinstance(score, (int, float)):
                # 7.8 → "score78", 6.5 → "score65"
                score_str = f"_score{int(score * 10)}"
    except Exception:
        pass

    return f"{safe_title}_{kind_kor}{round_str}{score_str}_{timestamp}_Blue.{ext}"


def parse_verification_docx(file_obj_or_path) -> dict:
    """검증 보고서 DOCX를 파싱하여 다음 라운드 입력으로 변환.

    파싱 대상:
    - "✗ [N]" 항목 → 미반영 처방 (재수정 대상)
    - "△ [Partial]" 항목 → 부분 반영 (보강 대상)
    - "A29~A32 위반" 항목 → 자동 후처리 대상
    - "재수정 권고" 섹션 → 수정 지시문으로 변환
    - "원본 유지" 표기된 씬 ID → 다음 라운드 REWRITE 후보

    Args:
        file_obj_or_path: DOCX 파일 객체 또는 경로

    Returns:
        {
            "round_n": 1,
            "previous_score": 6.8,
            "previous_verdict": "NEEDS_REVISION",
            "not_reflected": [{"original": "...", "scene_ids": [...]}],
            "partial_reflected": [...],
            "ai_escape_violations": [{"pattern_id": "A29", "scene_id": "S#1", "quote": "..."}],
            "untouched_scenes": ["S#10", "S#14", ...],
            "next_recommendations": ["권고 1...", "권고 2..."],
            "instruction_text": "다음 라운드 진단·집필에 주입할 통합 지시문"
        }
    """
    from docx import Document as _DocxDoc
    import re as _re_v33

    try:
        doc = _DocxDoc(file_obj_or_path)
    except Exception as e:
        return {"error": f"DOCX 파싱 실패: {e}"}

    paragraphs = [p.text for p in doc.paragraphs]
    full_text = "\n".join(paragraphs)

    result = {
        "round_n": 1,
        "previous_score": None,
        "previous_verdict": "",
        "not_reflected": [],
        "partial_reflected": [],
        "ai_escape_violations": [],
        "untouched_scenes": [],
        "next_recommendations": [],
        "instruction_text": "",
    }

    # 1. 종합 점수 + 판정
    m = _re_v33.search(r'\[\s*(APPROVED|NEEDS_REVISION|REJECTED)\s*\]', full_text)
    if m:
        result["previous_verdict"] = m.group(1)
    m = _re_v33.search(r'종합\s*점수[:：]\s*(\d+\.?\d*)\s*/\s*10', full_text)
    if m:
        result["previous_score"] = float(m.group(1))

    # 2. 미반영 항목 (✗ [N])
    not_reflected_pattern = _re_v33.compile(r'✗\s*\[N\]\s*(.+?)(?=\n\s*✓|\n\s*△|\n\s*✗|\n\s*■|\Z)',
                                              _re_v33.DOTALL)
    for m in not_reflected_pattern.finditer(full_text):
        item_text = m.group(1).strip()
        # 첫 줄을 처방 요지로
        lines = [l.strip() for l in item_text.split('\n') if l.strip()]
        if lines:
            result["not_reflected"].append({
                "instruction": lines[0],
                "details": " ".join(lines[1:])[:300] if len(lines) > 1 else "",
                "scene_ids": _re_v33.findall(r'S#\d+', item_text)
            })

    # 3. 부분 반영 항목 (△ [Partial])
    partial_pattern = _re_v33.compile(r'△\s*\[Partial\]\s*(.+?)(?=\n\s*✓|\n\s*△|\n\s*✗|\n\s*■|\Z)',
                                        _re_v33.DOTALL)
    for m in partial_pattern.finditer(full_text):
        item_text = m.group(1).strip()
        lines = [l.strip() for l in item_text.split('\n') if l.strip()]
        if lines:
            result["partial_reflected"].append({
                "instruction": lines[0],
                "details": " ".join(lines[1:])[:300] if len(lines) > 1 else "",
                "scene_ids": _re_v33.findall(r'S#\d+', item_text)
            })

    # 4. AI ESCAPE 위반 항목 (A29~A32 등)
    # 패턴: "✗ [A29] 시간 정밀 표기 금지 — S#1 REWRITE (Medium)"
    escape_pattern = _re_v33.compile(
        r'✗\s*\[(A\d+)\][^—\n]*?—\s*([^(\n]+?)\s*\((\w+)\)\s*\n[^→]*→\s*"([^"]+)"',
        _re_v33.DOTALL
    )
    for m in escape_pattern.finditer(full_text):
        result["ai_escape_violations"].append({
            "pattern_id": m.group(1),
            "scene_ids": _re_v33.findall(r'S#\d+', m.group(2)),
            "scene_label": m.group(2).strip(),
            "severity": m.group(3),
            "quote": m.group(4).strip(),
        })

    # 5. 원본 유지 / 미처리 씬 추출
    # 패턴: "S#XX (원본 유지)" 또는 "원본 유지: S#10, S#14, ..."
    untouched_set = set()
    for m in _re_v33.finditer(r'(?:원본\s*유지|미처리|미수정).{0,120}?(S#\d+(?:\s*[,/·]\s*S#\d+)*)', full_text):
        scene_list = _re_v33.findall(r'S#\d+', m.group(1))
        untouched_set.update(scene_list)
    # 직접 언급된 패턴: "S#X / S#Y / S#Z (원본 유지)"
    for m in _re_v33.finditer(r'(S#\d+(?:\s*[/·,]\s*S#\d+)+)\s*\(원본\s*유지\)', full_text):
        scene_list = _re_v33.findall(r'S#\d+', m.group(1))
        untouched_set.update(scene_list)
    result["untouched_scenes"] = sorted(untouched_set,
                                         key=lambda s: int(_re_v33.search(r'\d+', s).group()))

    # 6. 재수정 권고 섹션
    recom_section = _re_v33.search(
        r'■\s*재수정\s*권고\s*\n?(.+?)(?=\n\s*■|\Z)', full_text, _re_v33.DOTALL
    )
    if recom_section:
        recom_text = recom_section.group(1)
        # 단락 단위로 시도
        # 패턴 1: "• N. ..." 명시적 항목 (각 줄 시작)
        items_v1 = _re_v33.findall(r'•\s*\d+\.\s*(.+?)(?=\n\s*•\s*\d+\.|\Z)',
                                    recom_text, _re_v33.DOTALL)
        # 패턴 2: 연속 텍스트에서 "• N." 또는 " N." 식별 (줄바꿈 없을 때)
        items_v2 = _re_v33.findall(r'(?:^|\s)•?\s*(\d+)\.\s*([^•]+?)(?=\s•?\s*\d+\.|\Z)',
                                    recom_text, _re_v33.DOTALL)
        # v1 우선, v2 폴백
        if items_v1 and len(items_v1) >= 2:
            for item in items_v1:
                cleaned = item.strip().replace('\n', ' ').strip()
                if cleaned and len(cleaned) > 10:
                    result["next_recommendations"].append(cleaned[:400])
        elif items_v2:
            for num_str, item in items_v2:
                cleaned = item.strip().replace('\n', ' ').strip()
                if cleaned and len(cleaned) > 10:
                    result["next_recommendations"].append(cleaned[:400])
        else:
            # 폴백 3: 단락 단위로 자르기
            paragraphs_split = [p.strip() for p in recom_text.split('\n\n') if p.strip()]
            for p in paragraphs_split:
                # "•" 또는 숫자.로 시작하는 것만
                if _re_v33.match(r'(?:•|\d+\.)', p):
                    cleaned = _re_v33.sub(r'^(?:•\s*)?\d+\.\s*', '', p).replace('\n', ' ').strip()
                    if cleaned and len(cleaned) > 10:
                        result["next_recommendations"].append(cleaned[:400])

    # 7. 통합 지시문 텍스트 생성
    parts = []
    if result["previous_score"]:
        parts.append(
            f"[Round 직전 검증 결과]\n"
            f"이전 라운드 종합 점수: {result['previous_score']}/10 ({result['previous_verdict']})\n"
            f"이번 라운드 목표: 8.0/10 도달."
        )

    if result["not_reflected"]:
        parts.append("\n[미반영 처방 — 이번 라운드에서 반드시 처리]")
        for i, item in enumerate(result["not_reflected"], 1):
            parts.append(f"{i}. {item['instruction']}")
            if item["details"]:
                parts.append(f"   세부: {item['details'][:200]}")

    if result["partial_reflected"]:
        parts.append("\n[부분 반영 처방 — 보강 필요]")
        for i, item in enumerate(result["partial_reflected"], 1):
            parts.append(f"{i}. {item['instruction']}")

    if result["ai_escape_violations"]:
        parts.append("\n[AI 작법 위반 — 자동 후처리 + 재집필 시 회피]")
        # pattern_id별 그룹화
        by_pattern = {}
        for v in result["ai_escape_violations"]:
            by_pattern.setdefault(v["pattern_id"], []).append(v)
        for pid, vs in sorted(by_pattern.items()):
            scene_set = set()
            for v in vs:
                scene_set.update(v["scene_ids"])
            scenes_str = ", ".join(sorted(scene_set,
                                            key=lambda s: int(_re_v33.search(r'\d+', s).group()))) if scene_set else "(전체)"
            quotes = " / ".join([f'"{v["quote"]}"' for v in vs[:3]])
            parts.append(f"- {pid}: {len(vs)}회 위반 — 위치 {scenes_str}\n  예시: {quotes}")

    if result["untouched_scenes"]:
        parts.append(f"\n[원본 유지 진단 씬 — 이번 라운드 우선 검토 대상]\n{', '.join(result['untouched_scenes'])}")

    if result["next_recommendations"]:
        parts.append("\n[다음 라운드 재수정 권고]")
        for i, r in enumerate(result["next_recommendations"], 1):
            parts.append(f"{i}. {r}")

    result["instruction_text"] = "\n".join(parts)

    return result


def auto_fix_a29_violations(text: str) -> tuple:
    """v3.3 — A29 시간 정밀 표기 자동 후처리.

    "한 박자", "찰나", "0.3초" 등을 모호한 시간어로 자동 치환.

    Args:
        text: 시나리오 본문 텍스트

    Returns:
        (fixed_text, replacement_count)
    """
    import re as _re_v33

    replacements = [
        # 정밀 시간 표기 → 모호한 시간어
        (r'한\s*박자', '잠깐'),
        (r'반\s*박자', '잠깐'),
        (r'두\s*박자', '잠시'),
        (r'세\s*박자', '잠시'),
        (r'찰나(다|이다)?', '잠깐'),
        (r'(\d+(?:\.\d+)?)\s*초간\b', '잠깐'),
        (r'(\d+(?:\.\d+)?)\s*초\s*(?:후|뒤)', '잠시 후'),
        (r'(\d+(?:\.\d+)?)\s*초\s*(?:만에|동안)', '잠깐'),
        # 단독 "0.3초", "1초" 같은 표현 (단어 경계 + 한국어 동사 직전)
        (r'(\d+(?:\.\d+)?)\s*초\s+(?=[가-힣])', '잠깐 '),
        (r'순간(이다|적으로)?', '갑자기'),  # "순간이다" 같은 연출 표현
    ]

    count = 0
    fixed = text
    for pat, repl in replacements:
        new_fixed, n = _re_v33.subn(pat, repl, fixed)
        if n > 0:
            fixed = new_fixed
            count += n

    return fixed, count


# =================================================================
# ★★★ v3.3.5 — A33 자동 검증·수정 (Verbatim Repetition Prohibition) ★★★
# =================================================================

def auto_fix_a33_violations(scenes_dict: dict) -> tuple:
    """v3.3.5 — A33 위반 자동 후처리.

    3가지 결함 패턴 자동 처리:
    1. 같은 씬 ID가 두 번 등장 → 두 번째 제거
    2. 같은 씬 내 같은 화자 동일 대사 → 두 번째 제거
    3. 작품 전체에서 동일 대사 verbatim 반복 → 경고만 (변주는 작가 결정 필요)

    Args:
        scenes_dict: {"scenes": [{"scene_id": "S#1", "scene_header": "...",
                                  "revised_content": "..."}], ...}

    Returns:
        (fixed_dict, stats) — 수정된 dict + 통계
    """
    import re as _re_v335
    from collections import defaultdict

    if not isinstance(scenes_dict, dict):
        return scenes_dict, {"errors": "invalid_input"}

    rr = scenes_dict.get("revise_result", scenes_dict)
    scenes = rr.get("scenes", []) if isinstance(rr, dict) else []
    if not scenes:
        return scenes_dict, {"no_scenes": True}

    stats = {
        "duplicate_scenes_removed": 0,
        "in_scene_duplicate_dialogues_removed": 0,
        "verbatim_repetitions_warned": [],
    }

    # === Phase 1: 같은 씬 ID 두 번 등장 → 두 번째 제거 ===
    seen_scene_ids = set()
    deduped_scenes = []
    for sc in scenes:
        sid = sc.get("scene_id", "")
        if sid in seen_scene_ids:
            stats["duplicate_scenes_removed"] += 1
            continue
        if sid:
            seen_scene_ids.add(sid)
        deduped_scenes.append(sc)
    scenes = deduped_scenes

    # === Phase 2: 같은 씬 내 동일 대사 중복 제거 ===
    dialogue_pattern = _re_v335.compile(
        r'^([가-힣a-zA-Z0-9\s]{1,15}?)\s*(?:\([^)]*\))?\s*\t+(?:\([^)]*\))?\s*(.+)$'
    )
    for sc in scenes:
        content = sc.get("revised_content", "")
        if not content:
            continue
        lines = content.split('\n')
        seen_dialogues = {}
        new_lines = []
        for ln in lines:
            m = dialogue_pattern.match(ln.strip())
            if m and "\t\t" in ln:
                speaker = m.group(1).strip()
                dialogue = _re_v335.sub(r'[\s.,!?…\-]', '', m.group(2).strip())
                key = f"{speaker}::{dialogue}"
                if dialogue and len(dialogue) >= 5 and key in seen_dialogues:
                    # 같은 씬 내 동일 대사 중복 — 제거
                    stats["in_scene_duplicate_dialogues_removed"] += 1
                    continue
                seen_dialogues[key] = True
            new_lines.append(ln)
        sc["revised_content"] = '\n'.join(new_lines)

    # === Phase 3: 작품 전체 verbatim 반복 검출 → 경고만 ===
    # (자동 변주는 LLM 영역이므로 코드는 식별만 함)
    global_dialogues = defaultdict(list)
    for sc in scenes:
        content = sc.get("revised_content", "")
        sid = sc.get("scene_id", "")
        for ln in content.split('\n'):
            m = dialogue_pattern.match(ln.strip())
            if m and "\t\t" in ln:
                speaker = m.group(1).strip()
                dialogue = _re_v335.sub(r'[\s.,!?…\-]', '', m.group(2).strip())
                if dialogue and len(dialogue) >= 8:  # 8자 이상만 (짧은 응답 제외)
                    key = f"{speaker}::{dialogue}"
                    global_dialogues[key].append(sid)

    for key, scene_ids in global_dialogues.items():
        if len(scene_ids) >= 2:
            speaker = key.split("::")[0]
            stats["verbatim_repetitions_warned"].append({
                "speaker": speaker,
                "occurrence_count": len(scene_ids),
                "scenes": scene_ids,
            })

    rr["scenes"] = scenes
    rr["_a33_auto_fix_stats"] = stats

    return scenes_dict, stats


def auto_fix_duplicate_scene_blocks(full_text: str) -> tuple:
    """v3.3.5 — 통합 본문(full_text)에서 씬 블록 통째로 중복 자동 제거.

    같은 씬 헤더(예: 'S#36-2')가 본문과 함께 두 번 출현하면
    두 번째 블록을 제거.

    Args:
        full_text: 통합된 시나리오 본문 텍스트

    Returns:
        (fixed_text, removed_count)
    """
    import re as _re_v335

    # 씬 단위로 분할
    scene_pat = _re_v335.compile(r'(?=^S#\d+(?:-\d+)?\.\s)', _re_v335.MULTILINE)
    blocks = scene_pat.split(full_text)

    seen_blocks = {}  # scene_id → first_block_normalized
    output_blocks = []
    removed = 0

    for blk in blocks:
        if not blk.strip():
            output_blocks.append(blk)
            continue
        # 첫 줄에서 씬 ID 추출
        first_line = blk.strip().split('\n')[0] if blk.strip() else ""
        m = _re_v335.match(r'^(S#\d+(?:-\d+)?)\.', first_line)
        if not m:
            output_blocks.append(blk)
            continue
        sid = m.group(1)

        # 본문 정규화 (공백·구두점 제거)
        body_normalized = _re_v335.sub(r'\s+', '', blk)

        if sid in seen_blocks:
            # 같은 씬 ID 두 번 등장 — 본문도 같은지 확인
            if seen_blocks[sid] == body_normalized:
                # 통째로 중복 — 제거
                removed += 1
                continue
            # 본문이 다르면 — 그래도 같은 씬 ID 중복은 결함이므로 두 번째 제거
            # (의도적 분할 씬은 S#XX-1, S#XX-2 같이 다른 ID를 써야 함)
            removed += 1
            continue

        seen_blocks[sid] = body_normalized
        output_blocks.append(blk)

    return ''.join(output_blocks), removed


def export_verify_json(verify_result: dict, title: str = "",
                        round_n: int = 1) -> bytes:
    """v3.3 — Stage 3 검증 결과를 JSON으로 내보내기.

    다음 라운드(Round N+1)에서 자동 흡수 가능한 형식.

    v3.3.3 핫픽스: 실제 verify_result 키 구조에 맞춰 폴백 체인 추가.
    - 새 구조: verify_result["verify_report"]["overall_verdict"|"overall_score"|...]
    - 구 구조: verify_result["verdict"|"overall_score"|...]
    - 둘 다 시도해서 데이터를 정확히 추출.

    Args:
        verify_result: Stage 3 verify 결과 dict
        title: 작품 제목
        round_n: 현재 라운드 번호

    Returns:
        JSON 바이트 (UTF-8)
    """
    import json as _json
    from datetime import datetime as _dt

    if not isinstance(verify_result, dict):
        verify_result = {}

    # ★ v3.3.3: verify_report 래핑 자동 인식
    # 실제 LLM 응답 구조: {"verify_report": {...실제 데이터...}}
    # 일부 코드 경로 구조: {...실제 데이터 평면 배치...}
    inner = verify_result.get("verify_report")
    if isinstance(inner, dict):
        # 새 구조 — verify_report 안에 실제 데이터
        src = inner
    else:
        # 구 구조 — 평면 배치
        src = verify_result

    # ★ v3.3.3: 키 폴백 체인 (다양한 키명 시도)
    verdict = (
        src.get("overall_verdict") or
        src.get("verdict") or
        src.get("final_verdict", {}).get("judgment") if isinstance(src.get("final_verdict"), dict) else None or
        src.get("judgment") or
        ""
    )

    overall_score = (
        src.get("overall_score") or
        src.get("score") or
        (src.get("final_verdict", {}).get("score") if isinstance(src.get("final_verdict"), dict) else None)
    )

    instruction_compliance = (
        src.get("instruction_compliance") or
        src.get("stage1_instruction") or
        {}
    )

    locked_preservation = (
        src.get("locked_preservation") or
        src.get("stage2_locked") or
        {}
    )

    ai_escape_check = (
        src.get("ai_escape_check") or
        src.get("stage3_ai_escape") or
        {}
    )

    genre_compliance = (
        src.get("genre_compliance") or
        src.get("stage4_genre") or
        {}
    )

    key_changes = (
        src.get("side_by_side_highlights") or  # 실제 키
        src.get("key_changes") or
        []
    )

    next_recommendations = (
        src.get("recommendations") or  # 실제 키
        src.get("next_round_recommendations") or
        []
    )

    # 진단 신뢰도, 판정 이유 등 부가 정보 추가
    verdict_reason = src.get("verdict_reason", "")

    export = {
        "schema_version": "revise_verify_v1.0",
        "engine_version": "BLUE JEANS REVISE ENGINE v3.3.3",
        "title": title or "(제목 미지정)",
        "round_n": round_n,
        "report_date": _dt.now().isoformat(),
        "verdict": verdict,
        "verdict_reason": verdict_reason,
        "overall_score": overall_score,
        "instruction_compliance": instruction_compliance,
        "locked_preservation": locked_preservation,
        "ai_escape_check": ai_escape_check,
        "genre_compliance": genre_compliance,
        "key_changes": key_changes,
        "next_round_recommendations": next_recommendations,
        # 다음 라운드 자동 흡수용 정규화 필드
        "normalized_for_next_round": {
            "untouched_scenes": _extract_untouched_from_verify(verify_result),
            "ai_escape_violations": ai_escape_check.get("violations", []) if isinstance(ai_escape_check, dict) else [],
            "target_score_next_round": (overall_score or 7.0) + 1.0 if isinstance(overall_score, (int, float)) else 8.0,
        }
    }

    return _json.dumps(export, ensure_ascii=False, indent=2).encode("utf-8")


def _extract_untouched_from_verify(verify_result: dict) -> list:
    """verify_result에서 '원본 유지' 진단된 씬 ID 추출.

    v3.3.3 핫픽스: verify_report 래핑 자동 인식.
    """
    import re as _re_v33

    if not isinstance(verify_result, dict):
        return []

    # ★ v3.3.3: verify_report 래핑 자동 인식
    inner = verify_result.get("verify_report")
    src = inner if isinstance(inner, dict) else verify_result

    untouched = set()
    # ai_escape_check.violations 내 (원본 유지) 표기
    ai_check = src.get("ai_escape_check") or src.get("stage3_ai_escape") or {}
    violations = ai_check.get("violations", []) if isinstance(ai_check, dict) else []
    for v in violations:
        if not isinstance(v, dict):
            continue
        scene_id = v.get("scene_id", "")
        if "원본 유지" in str(scene_id) or "원본유지" in str(scene_id):
            ids = _re_v33.findall(r'S#\d+', str(scene_id))
            untouched.update(ids)

    # instruction_compliance 내 [N] 항목
    ic = src.get("instruction_compliance") or src.get("stage1_instruction") or {}
    items = ic.get("items", []) if isinstance(ic, dict) else []
    items = items or []
    for item in items:
        if not isinstance(item, dict):
            continue
        status = item.get("status", "")
        if status == "N" or status == "Partial":
            details = str(item.get("details", "")) + " " + str(item.get("instruction", ""))
            if "원본 유지" in details or "원본유지" in details:
                ids = _re_v33.findall(r'S#\d+', details)
                untouched.update(ids)

    return sorted(untouched, key=lambda s: int(_re_v33.search(r'\d+', s).group()))


def create_verify_docx(verify_result: dict, title: str = "") -> bytes:
    """Stage 3 검증 보고서 DOCX."""
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── 표지 ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BLUE JEANS PICTURES · REVISE ENGINE")
    _set_font(r, size_pt=9, bold=True, color=(255, 203, 5))

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title if title else "검증 보고서")
    _set_font(r, size_pt=20, bold=True, color=(25, 25, 112))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Verification Report")
    _set_font(r, size_pt=11, color=(142, 142, 153))

    doc.add_paragraph()

    vr = verify_result.get("verify_report", {})

    # ── 종합 판정 ──
    verdict = vr.get("overall_verdict", "")
    score = vr.get("overall_score", "")
    reason = vr.get("verdict_reason", "")

    verdict_color = (46, 196, 132) if verdict == "APPROVED" else \
                    (255, 203, 5) if verdict == "NEEDS_REVISION" else \
                    (225, 68, 68)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"[ {verdict} ]")
    _set_font(r, size_pt=16, bold=True, color=verdict_color)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"종합 점수: {score} / 10.0")
    _set_font(r, size_pt=13, bold=True, color=(25, 25, 112))

    if reason:
        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run(reason)
        _set_font(r, size_pt=10.5)

    doc.add_paragraph()

    # ── 4축 검증 ──
    def section_score(title_text, key):
        block = vr.get(key, {})
        sc = block.get("score", "")
        p = doc.add_paragraph()
        r = p.add_run(f"■ {title_text}  —  {sc}/10")
        _set_font(r, size_pt=13, bold=True, color=(25, 25, 112))
        return block

    # 1. 지시사항 반영도
    block = section_score("1. 지시사항 반영도 (Instruction Compliance)", "instruction_compliance")
    for item in block.get("items", []):
        status = item.get("status", "")
        inst = item.get("instruction_item", "")
        ev = item.get("evidence", "")
        mark = "✓" if status == "Y" else ("△" if status == "Partial" else "✗")
        p = doc.add_paragraph()
        r = p.add_run(f"  {mark} [{status}] {inst}")
        color = (46, 196, 132) if status == "Y" else (255, 140, 0) if status == "Partial" else (225, 68, 68)
        _set_font(r, size_pt=10, bold=True, color=color)
        if ev:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            r = p.add_run(f"    → {ev}")
            _set_font(r, size_pt=9.5, color=(80, 80, 90))

    doc.add_paragraph()

    # 2. LOCKED 보존도
    block = section_score("2. LOCKED 보존도 (Locked Preservation)", "locked_preservation")
    for item in block.get("items", []):
        status = item.get("status", "")
        locked_item = item.get("locked_item", "")
        ev = item.get("evidence", "")
        mark = "✓" if status == "Preserved" else ("—" if status == "N/A" else "✗")
        p = doc.add_paragraph()
        r = p.add_run(f"  {mark} [{status}] {locked_item}")
        color = (46, 196, 132) if status == "Preserved" else (142, 142, 153) if status == "N/A" else (225, 68, 68)
        _set_font(r, size_pt=10, bold=True, color=color)
        if ev:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            r = p.add_run(f"    → {ev}")
            _set_font(r, size_pt=9.5, color=(80, 80, 90))

    doc.add_paragraph()

    # 3. AI ESCAPE
    block = section_score("3. AI SCREENPLAY ESCAPE 준수도", "ai_escape_check")
    clean = block.get("clean_patterns", "")
    if clean:
        p = doc.add_paragraph()
        r = p.add_run(f"  ✓ 위반 없는 패턴: {clean} / 20")
        _set_font(r, size_pt=10, color=(46, 196, 132))
    violations = block.get("violations", [])
    if violations:
        for v in violations:
            pid = v.get("pattern_id", "")
            pname = v.get("pattern_name", "")
            scene = v.get("scene_id", "")
            quote = v.get("quote", "")
            sev = v.get("severity", "")
            sev_color = (225, 68, 68) if sev == "High" else (255, 140, 0) if sev == "Medium" else (142, 142, 153)
            p = doc.add_paragraph()
            r = p.add_run(f"  ✗ [{pid}] {pname} — {scene} ({sev})")
            _set_font(r, size_pt=10, bold=True, color=sev_color)
            if quote:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.8)
                r = p.add_run(f'    → "{quote}"')
                _set_font(r, size_pt=9.5, color=(80, 80, 90))
    else:
        p = doc.add_paragraph()
        r = p.add_run("  위반 사항 없음.")
        _set_font(r, size_pt=10, color=(46, 196, 132))

    doc.add_paragraph()

    # 4. 장르 준수도
    block = section_score("4. 장르 RULE PACK 준수도", "genre_compliance")
    p = doc.add_paragraph()
    r = p.add_run("  [Must Have]")
    _set_font(r, size_pt=10.5, bold=True, color=(25, 25, 112))
    for item in block.get("must_have_check", []):
        st_val = item.get("status", "")
        it = item.get("item", "")
        nt = item.get("note", "")
        mark = "✓" if st_val == "Met" else ("△" if st_val == "Partial" else "✗")
        color = (46, 196, 132) if st_val == "Met" else (255, 140, 0) if st_val == "Partial" else (225, 68, 68)
        p = doc.add_paragraph()
        r = p.add_run(f"    {mark} [{st_val}] {it}  —  {nt}")
        _set_font(r, size_pt=9.5, color=color)

    p = doc.add_paragraph()
    r = p.add_run("  [Fails to Avoid]")
    _set_font(r, size_pt=10.5, bold=True, color=(25, 25, 112))
    for item in block.get("fails_check", []):
        st_val = item.get("status", "")
        it = item.get("item", "")
        nt = item.get("note", "")
        mark = "✓" if st_val == "Avoided" else ("△" if st_val == "Improved" else "✗")
        color = (46, 196, 132) if st_val == "Avoided" else (255, 140, 0) if st_val == "Improved" else (225, 68, 68)
        p = doc.add_paragraph()
        r = p.add_run(f"    {mark} [{st_val}] {it}  —  {nt}")
        _set_font(r, size_pt=9.5, color=color)

    doc.add_paragraph()

    # ── 주요 변화 ──
    highlights = vr.get("side_by_side_highlights", [])
    if highlights:
        p = doc.add_paragraph()
        r = p.add_run("■ 핵심 변화 요약")
        _set_font(r, size_pt=13, bold=True, color=(25, 25, 112))
        for h in highlights:
            sid = h.get("scene_id", "")
            kc = h.get("key_change", "")
            imp = h.get("improvement_note", "")
            p = doc.add_paragraph()
            r = p.add_run(f"  • {sid}: {kc}")
            _set_font(r, size_pt=10, bold=True)
            if imp:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.8)
                r = p.add_run(f"    {imp}")
                _set_font(r, size_pt=9.5, color=(80, 80, 90))

    # ── 재수정 권고 ──
    recs = vr.get("recommendations", [])
    if recs:
        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run("■ 재수정 권고")
        _set_font(r, size_pt=13, bold=True, color=(225, 68, 68))
        for rec in recs:
            p = doc.add_paragraph()
            r = p.add_run(f"  • {rec}")
            _set_font(r, size_pt=10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# =================================================================
# [6] Stage 실행 함수
# =================================================================
def run_v2_pre_analyses(client) -> dict:
    """v2.0/v2.1 자동 분석을 실행하고 결과를 세션에 저장.

    실행 항목 (각각 독립적으로 캐시 체크):
    1. 톤 DNA 추출 (tone_ref_text 있고 캐시 없을 때)
    2. Diff 학습 (After 있고 캐시 없을 때)
    3. 분포 진단 (항상 실행, 캐시 없을 때)
    4. Rewrite 메타 흡수 (rewrite_json_text 있고 캐시 없을 때)
    5. 장르 DNA 추출 (genre_ref_texts 있고 캐시 없을 때) [v2.1]

    각 항목은 독립적으로 캐시되므로, 일부만 추가 업로드해도 정확히 그 부분만 재추출됨.

    Returns:
        {tone_dna, diff_analysis, distribution_diagnostic, rewrite_metadata, genre_dna}
    """
    results = {
        "tone_dna": st.session_state.tone_dna,
        "diff_analysis": st.session_state.diff_analysis,
        "distribution_diagnostic": st.session_state.distribution_diagnostic,
        "rewrite_metadata": st.session_state.rewrite_metadata,
        "genre_dna": st.session_state.genre_dna,
    }

    # 1. 톤 DNA 추출 (캐시 없을 때만)
    if (st.session_state.tone_ref_text and st.session_state.tone_ref_text.strip()
            and not st.session_state.tone_dna):
        with st.spinner("📐 톤 DNA 자동 추출 중... (Sonnet 4.6)"):
            prompt_text = build_tone_dna_extraction_prompt(st.session_state.tone_ref_text)
            raw = call_claude(client, prompt_text, model=MODEL_ANALYZE, max_tokens=5000)
            if raw:
                tone_dna = parse_json(raw)
                if tone_dna:
                    results["tone_dna"] = tone_dna
                    st.session_state.tone_dna = tone_dna
                    summary = tone_dna.get("tone_dna", {}).get("summary", "톤 추출 완료")
                    st.success(f"✅ 톤 DNA 추출 완료: {summary[:120]}")

    # 2. Diff 학습 (Before는 옵션에 따라 main 시나리오 또는 별도 업로드)
    diff_before = ""
    diff_after = st.session_state.diff_refined_text.strip() if st.session_state.diff_refined_text else ""
    if diff_after:
        if st.session_state.diff_use_main_as_before:
            diff_before = st.session_state.raw_text.strip() if st.session_state.raw_text else ""
        else:
            diff_before = st.session_state.diff_orig_text.strip() if st.session_state.diff_orig_text else ""

    if diff_before and diff_after:
        with st.spinner("🔬 작가 편집 패턴 학습 중... (Sonnet 4.6)"):
            prompt_text = build_diff_analysis_prompt(diff_before, diff_after)
            raw = call_claude(client, prompt_text, model=MODEL_ANALYZE, max_tokens=8000)
            if raw:
                diff = parse_json(raw)
                if diff:
                    results["diff_analysis"] = diff
                    st.session_state.diff_analysis = diff
                    summary = diff.get("diff_analysis", {}).get("summary", "편집 패턴 학습 완료")
                    st.success(f"✅ Diff 학습 완료: {summary[:120]}")

    # 3. 분포 진단 (전체 각색 모드에서만 실행 — 구간 모드는 보호 영역이 이미 OK라 불필요)
    if (not st.session_state.section_mode
            and not st.session_state.distribution_diagnostic):
        with st.spinner("📊 장르 메트릭 + 캐릭터 분포 진단 중... (Sonnet 4.6)"):
            prompt_text = build_distribution_diagnostic_prompt(
                st.session_state.raw_text,
                st.session_state.genre
            )
            raw = call_claude(client, prompt_text, model=MODEL_ANALYZE, max_tokens=5000)
            if raw:
                dist = parse_json(raw)
                if dist:
                    results["distribution_diagnostic"] = dist
                    st.session_state.distribution_diagnostic = dist
                    summary = dist.get("distribution_diagnostic", {}).get("summary", "분포 진단 완료")
                    upgrades = dist.get("distribution_diagnostic", {}).get("auto_priority_upgrades", [])
                    st.success(f"✅ 분포 진단 완료: {summary[:100]}  (자동 격상 {len(upgrades)}개)")

    # 4. Rewrite 메타 흡수
    if st.session_state.rewrite_json_text and st.session_state.rewrite_json_text.strip():
        meta = absorb_rewrite_engine_metadata(st.session_state.rewrite_json_text)
        if meta and any([
            meta.get("preserve_notes_by_seq"),
            meta.get("weak_zone_scenes"),
            meta.get("auto_priority_high"),
        ]):
            results["rewrite_metadata"] = meta
            st.session_state.rewrite_metadata = meta
            preserve_count = len(meta.get("preserve_notes_by_seq", {}))
            weak_count = len(meta.get("weak_zone_scenes", []))
            st.success(f"✅ Rewrite 메타 흡수: preserve_notes {preserve_count}개, weak_zones {weak_count}개")

    # 5. 장르 DNA 추출 (v2.1)
    if st.session_state.genre_ref_texts and any(t.strip() for t in st.session_state.genre_ref_texts):
        ref_count = len(st.session_state.genre_ref_texts)
        with st.spinner(f"🎬 장르 DNA 추출 중... 참고작 {ref_count}편 분석 (Sonnet 4.6)"):
            prompt_text = build_genre_dna_extraction_prompt(
                st.session_state.genre_ref_texts,
                st.session_state.genre
            )
            raw = call_claude(client, prompt_text, model=MODEL_ANALYZE, max_tokens=5000)
            if raw:
                gd = parse_json(raw)
                if gd:
                    results["genre_dna"] = gd
                    st.session_state.genre_dna = gd
                    summary = gd.get("genre_dna", {}).get("summary", "장르 DNA 추출 완료")
                    st.success(f"✅ 장르 DNA 추출 완료: {summary[:120]}")

    # 6. 구간 자동 감지 (v2.2 — step1 가벼운 호출 + 코드 계산)
    if (st.session_state.section_mode
            and st.session_state.section_input_method in ("auto", "hybrid")
            and st.session_state.diff_refined_text
            and not st.session_state.section_detection):
        with st.spinner("🔍 구간 자동 감지 (1단계: 새로 쓴 부분 식별)..."):
            prompt_text = build_section_detect_step1_prompt(
                refined_text=st.session_state.diff_refined_text,
                original_text=st.session_state.raw_text,
            )
            # step1은 출력이 매우 짧음 — 2000 토큰이면 충분
            raw = call_claude(client, prompt_text, model=MODEL_ANALYZE, max_tokens=2000)
            if raw:
                step1 = parse_json(raw)
                if step1:
                    # 코드로 protected/revision 범위 자동 계산 (API 호출 불필요)
                    derived = derive_section_ranges_from_step1(
                        step1_result=step1,
                        refined_text=st.session_state.diff_refined_text,
                        original_text=st.session_state.raw_text,
                    )
                    rec_prot = derived.get("protected_ranges", [])
                    rec_rev = derived.get("revision_ranges", [])

                    # 캐시에 저장 (구버전 형식과 호환)
                    st.session_state.section_detection = {
                        "section_detection": {
                            "summary": derived.get("continuation_point", {}).get("explanation", ""),
                            "continuation_point": derived.get("continuation_point", {}),
                            "recommended_protected_range": rec_prot,
                            "recommended_revision_range": rec_rev,
                        }
                    }

                    # auto + hybrid 모두 자동 적용
                    if st.session_state.section_input_method in ("auto", "hybrid"):
                        if rec_prot and not st.session_state.protected_ranges:
                            st.session_state.protected_ranges = rec_prot
                        if rec_rev and not st.session_state.revision_ranges:
                            st.session_state.revision_ranges = rec_rev

                    cp = derived.get("continuation_point", {})
                    if cp.get("detected") == "true":
                        prot_str_log = ", ".join(
                            f"{r.get('from','')}~{r.get('to','')}" for r in rec_prot
                        )
                        rev_str_log = ", ".join(
                            f"{r.get('from','')}~{r.get('to','')}" for r in rec_rev
                        )
                        st.success(
                            f"✅ **이어쓰기 시작점 감지 완료**\n\n"
                            f"🔒 보호 구간: `{prot_str_log}`\n\n"
                            f"✏️ 재집필 구간: `{rev_str_log}`\n\n"
                            f"→ {cp.get('explanation', '')[:200]}"
                        )
                    else:
                        st.warning(
                            "⚠️ 자동 감지 결과가 명확하지 않습니다. "
                            "**부분 수정 모드 + 직접 지정**으로 재집필 구간을 입력하세요."
                        )

    # 7. 연쇄 영향 분석 (v2.2 — 보호/재집필 영역 모두 있을 때)
    if (st.session_state.section_mode
            and st.session_state.protected_ranges
            and st.session_state.revision_ranges
            and not st.session_state.cascade_analysis):
        with st.spinner("🔬 연쇄 영향 분석 중... 보호 구간과의 모순 점검 (Sonnet 4.6)"):
            prompt_text = build_cascade_analysis_prompt(
                revision_ranges=st.session_state.revision_ranges,
                protected_ranges=st.session_state.protected_ranges,
                raw_text=st.session_state.raw_text,
            )
            raw = call_claude(client, prompt_text, model=MODEL_ANALYZE, max_tokens=5000)
            if raw:
                ca = parse_json(raw)
                if ca:
                    st.session_state.cascade_analysis = ca
                    summary = ca.get("cascade_analysis", {}).get("summary", "연쇄 분석 완료")
                    must_preserve_count = len(ca.get("cascade_analysis", {}).get("must_preserve", []))
                    conflict_count = len(ca.get("cascade_analysis", {}).get("potential_conflicts", []))
                    st.success(
                        f"✅ 연쇄 영향 분석 완료: 보존 요소 {must_preserve_count}개 / 잠재 모순 {conflict_count}개"
                    )

    # 8. 경계 매끄러움 정보 자동 계산 (v2.2)
    if (st.session_state.section_mode
            and st.session_state.protected_ranges
            and st.session_state.revision_ranges
            and not st.session_state.boundary_info):
        boundary = build_boundary_smoothness_block(
            protected_ranges=st.session_state.protected_ranges,
            revision_ranges=st.session_state.revision_ranges,
            raw_text=st.session_state.raw_text,
        )
        if boundary:
            st.session_state.boundary_info = boundary

    return results


def _build_auto_diagnose_for_section_mode() -> dict:
    """구간 모드일 때 Stage 1 진단을 코드로 자동 생성.

    이어쓰기/부분수정에서는 어디를 다시 쓸지 이미 정해졌으므로,
    AI에게 진단 시키는 대신 코드로 즉시 revision_plan을 만든다.
    토큰 절약 + 잘림 방지 + 즉시 응답.
    """
    import re as _re

    revision_ranges = st.session_state.revision_ranges or []
    protected_ranges = st.session_state.protected_ranges or []

    # 재집필 구간 안의 모든 씬을 target_scenes로 자동 등록
    target_scenes = []
    raw = st.session_state.raw_text or ""

    # 원본에서 모든 씬 헤더 + 헤더 텍스트 추출
    scene_pattern = _re.compile(r'^\s*\*?\*?S#?(\d+)[\.\s]([^\n]*)', _re.MULTILINE)
    all_scenes = {int(m.group(1)): m.group(2).strip() for m in scene_pattern.finditer(raw)}

    for rev_range in revision_ranges:
        from_str = rev_range.get("from", "")
        to_str = rev_range.get("to", "")
        from_num = int(_re.search(r'\d+', from_str).group()) if from_str else 0
        to_num = int(_re.search(r'\d+', to_str).group()) if to_str else 0

        for scene_num in range(from_num, to_num + 1):
            if scene_num in all_scenes:
                target_scenes.append({
                    "scene_id": f"S#{scene_num}",
                    "header": f"S#{scene_num}. {all_scenes[scene_num][:80]}",
                    "priority": "HIGH",
                    "type": "REWRITE",
                    "what_to_change": [
                        "보호 구간의 톤·스타일을 그대로 따라 다시 쓴다",
                        "원본 씬의 핵심 사건·인물·장소는 유지",
                        "대사·지문은 작가 톤(절제된 현재형 단문, 미시 시간 표기)으로 재작성"
                    ],
                    "preservation_notes": [
                        "원본 씬의 이벤트·정보 누락 금지",
                        "캐릭터 정체성·관계 변경 금지"
                    ],
                })

    # LOCKED 인식 (입력값 그대로)
    locked_text = (st.session_state.locked or "").strip()
    locked_recognition = []
    if locked_text:
        # LOCKED 텍스트를 줄 단위로 분리해 인식 항목으로
        for line in locked_text.split("\n"):
            line = line.strip()
            if line:
                locked_recognition.append({
                    "category": "사용자 명시",
                    "item": line,
                    "scope": "전체"
                })

    # 자동 생성된 진단 결과
    prot_str = ", ".join(f"{r.get('from','')}~{r.get('to','')}" for r in protected_ranges)
    rev_str = ", ".join(f"{r.get('from','')}~{r.get('to','')}" for r in revision_ranges)

    work_mode = st.session_state.work_mode or "continuation"
    if work_mode == "continuation":
        summary = (
            f"이어쓰기 모드 — 작가가 직접 손본 보호 구간({prot_str})의 톤·스타일을 학습하여, "
            f"재집필 구간({rev_str})을 같은 결로 다시 쓴다. "
            f"보호 구간은 한 글자도 건드리지 않으며, 재집필 구간의 모든 씬은 priority HIGH로 처리된다. "
            f"원본의 핵심 사건은 유지하되, 표현은 작가의 절제된 현재형 단문 + 미시 시간 표기 스타일로 재작성한다."
        )
    else:
        summary = (
            f"부분 수정 모드 — 사용자가 지정한 구간({rev_str})만 재집필한다. "
            f"나머지 구간({prot_str})은 보호되어 변경되지 않는다. "
            f"재집필 시 인접한 보호 구간과 자연스럽게 이어지도록 경계 조건을 따른다."
        )

    return {
        "revision_plan": {
            "summary": summary,
            "estimated_scene_count": len(target_scenes),
            "confidence": 9,
            "auto_generated": True,
            "auto_generation_reason": "구간 모드에서는 진단 대신 자동 플랜 생성 (토큰 절약·즉시 응답)",
            "locked_recognition": locked_recognition,
            "locked_conflicts": [],
            "target_scenes": target_scenes,
            "out_of_scope": [],
            "section_mode_info": {
                "mode": work_mode,
                "protected_ranges": protected_ranges,
                "revision_ranges": revision_ranges,
            }
        }
    }


def _build_auto_diagnose_from_rewrite_metadata(rewrite_metadata: dict) -> dict:
    """Rewrite Engine JSON 메타데이터로부터 Stage 1 진단을 코드로 자동 생성.

    REWRITE 제안 6개 + ADD 제안 4개 + MOON 의견을 받아
    AI 호출 없이 즉시 revision_plan을 완성한다.

    토큰 절감 + 잘림 방지 + 즉시 응답.
    """
    import re as _re

    target_scenes = []
    raw = st.session_state.raw_text or ""

    # 원본 시나리오에서 모든 씬 헤더 추출 (insert_after 검증용)
    scene_pattern = _re.compile(r'^\s*\*?\*?S#?(\d+)[\.\s]([^\n]*)', _re.MULTILINE)
    all_scenes = {int(m.group(1)): m.group(2).strip() for m in scene_pattern.finditer(raw)}

    # ── 1. REWRITE 제안 → priority HIGH로 자동 등록 ──
    rewrite_suggestions = rewrite_metadata.get("rewrite_suggestions", [])
    for item in rewrite_suggestions:
        scene_id = item.get("scene_id", "")
        what = item.get("what_to_change", "")
        why = item.get("why", "")

        # scene_id에서 숫자 추출
        m = _re.search(r'\d+', scene_id)
        if not m:
            continue
        scene_num = int(m.group())
        header = f"S#{scene_num}. {all_scenes.get(scene_num, '')[:80]}"

        target_scenes.append({
            "scene_id": f"S#{scene_num}",
            "header": header,
            "priority": "HIGH",
            "type": "REWRITE",
            "what_to_change": [what] if what else ["Rewrite Engine 처방 적용"],
            "preservation_notes": [
                "원본 씬의 핵심 사건·인물·장소 유지",
                "캐릭터 정체성·관계 변경 금지"
            ],
            "rewrite_engine_reason": why,
            "rewrite_engine_source": "rewrite_suggestions",
        })

    # ── 2. ADD 제안 → type=ADD로 자동 등록 ──
    add_suggestions = rewrite_metadata.get("add_suggestions", [])
    add_counter = 0
    for item in add_suggestions:
        insert_after = item.get("insert_after", "")
        scene_type = item.get("type", "추가 시퀀스")
        content_plan = item.get("content_plan", "")
        why = item.get("why", "")

        m = _re.search(r'\d+', insert_after)
        if not m:
            continue
        after_num = int(m.group())
        add_counter += 1

        target_scenes.append({
            "scene_id": f"(NEW) ADD-{add_counter}",
            "header": f"(NEW) S#{after_num}-{add_counter} [{scene_type}]",
            "priority": "HIGH",
            "type": "ADD",
            "insert_after": f"S#{after_num}",
            "what_to_change": [content_plan] if content_plan else [f"{scene_type} 추가"],
            "preservation_notes": [
                "기존 씬과 자연스럽게 연결",
                "보호 구간의 톤·캐릭터 정체성 유지",
                "헐리우드 작법 표준 시퀀스 구성"
            ],
            "context_before": f"S#{after_num} 직후",
            "context_after": f"S#{after_num + 1}" if (after_num + 1) in all_scenes else "(다음 씬)",
            "rewrite_engine_reason": why,
            "rewrite_engine_source": "add_suggestions",
        })

    # ── 3. weak_zone 시퀀스 → priority HIGH로 추가 등록 ──
    weak_zones = rewrite_metadata.get("weak_zone_scenes", [])
    for zone in weak_zones:
        seq_ref = zone.get("seq_ref", "")
        m = _re.search(r'\d+', seq_ref)
        if not m:
            continue
        scene_num = int(m.group())
        # 이미 REWRITE에 있으면 스킵
        if any(s.get("scene_id") == f"S#{scene_num}" for s in target_scenes):
            continue

        header = f"S#{scene_num}. {all_scenes.get(scene_num, '')[:80]}"
        hook = zone.get("hook_suggestion", "")
        punch = zone.get("punch_suggestion", "")

        target_scenes.append({
            "scene_id": f"S#{scene_num}",
            "header": header,
            "priority": "HIGH",
            "type": "REWRITE",
            "what_to_change": [
                f"훅 강화: {hook[:80]}" if hook else "약점 영역 강화",
                f"펀치 보강: {punch[:80]}" if punch else "장르 재미 회복"
            ],
            "preservation_notes": ["원본 핵심 보존", "캐릭터 일관성 유지"],
            "rewrite_engine_source": "weak_zones",
        })

    # ── 4. MOON 의견 → summary에 통합 ──
    moon_text = rewrite_metadata.get("moon_opinion_text", "")
    moon_market = rewrite_metadata.get("moon_market_direction", "")
    moon_genre = rewrite_metadata.get("moon_genre_strengthening", "")
    moon_unique = rewrite_metadata.get("moon_unique_value", "")

    moon_parts = []
    if moon_text:
        moon_parts.append(f"전체 방향: {moon_text}")
    if moon_market:
        moon_parts.append(f"시장: {moon_market}")
    if moon_genre:
        moon_parts.append(f"장르 강화: {moon_genre}")
    if moon_unique:
        moon_parts.append(f"차별성: {moon_unique}")
    moon_summary = " / ".join(moon_parts) if moon_parts else ""

    # ── 5. LOCKED 인식 ──
    locked_text = (st.session_state.locked or "").strip()
    locked_recognition = []
    if locked_text:
        for line in locked_text.split("\n"):
            line = line.strip()
            if line:
                locked_recognition.append({
                    "category": "사용자 명시",
                    "item": line,
                    "scope": "전체"
                })

    # ── 6. summary 생성 ──
    rewrite_count = len([s for s in target_scenes if s.get("type") == "REWRITE"])
    add_count = len([s for s in target_scenes if s.get("type") == "ADD"])

    summary = (
        f"Rewrite Engine 처방 자동 흡수 — REWRITE {rewrite_count}개 씬 수정 + ADD {add_count}개 씬 추가. "
        f"AI 호출 없이 코드로 진단 즉시 생성. "
        f"MOON 의견 → 전체 방향에 강제 반영. "
        f"보호 구간(작가가 손본 부분)은 그대로 유지하며, 위 처방대로 정밀 수정·추가 진행."
    )
    if moon_summary:
        summary += f"\n\n[MOON 처방] {moon_summary}"

    return {
        "revision_plan": {
            "summary": summary,
            "estimated_scene_count": len(target_scenes),
            "confidence": 9,
            "auto_generated": True,
            "auto_generation_reason": "Rewrite Engine JSON 메타데이터로부터 코드로 즉시 생성 (토큰 절약·잘림 방지)",
            "locked_recognition": locked_recognition,
            "locked_conflicts": [],
            "target_scenes": target_scenes,
            "out_of_scope": [],
            "moon_directives": {
                "overall": moon_text,
                "market": moon_market,
                "genre": moon_genre,
                "unique": moon_unique,
            },
            "rewrite_engine_absorbed": True,
        }
    }


# =================================================================
# ★★★ v2.7 자동 배치 분할 시스템 ★★★
# 71씬 같은 대형 시나리오를 자동으로 N씬 단위로 쪼개
# 각 배치를 별도 LLM 호출로 진단·집필 후 결과 병합
# =================================================================

import re as _re_v27


def _detect_scene_count(scenario_text: str) -> int:
    """시나리오에서 씬 헤더 패턴을 카운트하여 총 씬 수 반환.

    인식 패턴 (우선순위):
    1. 'S#숫자.' 한국 시나리오 표준 형식 (S#1., S#1, S# 1 등)
    2. 'EXT./INT. 장소 — DAY/NIGHT' 헐리우드 형식 (테이스티 러브 v3.2 같은 케이스)
    3. '씬 숫자' 또는 'Scene 숫자' 폴백

    Returns:
        총 씬 개수 (정수). 인식 실패 시 0.
    """
    if not scenario_text:
        return 0

    # 패턴 1: 'S#숫자' (가장 흔한 한국 시나리오 형식)
    pattern_a = r'(?:^|\n)\s*S\s*#\s*(\d+)\b'
    matches = _re_v27.findall(pattern_a, scenario_text, flags=_re_v27.MULTILINE)
    if matches:
        nums = sorted(set(int(m) for m in matches))
        return len(nums)

    # 패턴 2: 'EXT./INT. ... — DAY/NIGHT' 헐리우드 형식
    # 예: "EXT. 한남시장 (한남동 재래시장) — DAY"
    #     "INT. 카페 - 낮"
    pattern_b = r'(?:^|\n)\s*(?:EXT|INT|EXT\./INT|I/E)\.\s+\S'
    matches_b = _re_v27.findall(pattern_b, scenario_text, flags=_re_v27.MULTILINE | _re_v27.IGNORECASE)
    if matches_b:
        return len(matches_b)

    # 패턴 3: 폴백 — '씬 숫자' 또는 'Scene 숫자'
    pattern_c = r'(?:^|\n)\s*(?:씬|Scene)\s*(\d+)'
    matches_c = _re_v27.findall(pattern_c, scenario_text, flags=_re_v27.MULTILINE | _re_v27.IGNORECASE)
    if matches_c:
        nums = sorted(set(int(m) for m in matches_c))
        return len(nums)

    return 0


def _split_scenario_by_scenes(scenario_text: str, batch_size: int = 12) -> list:
    """시나리오 본문을 N씬 단위로 분할.

    인식 형식:
    - 'S#숫자' 한국 시나리오 표준 (씬 번호 사용)
    - 'EXT./INT. ... — DAY/NIGHT' 헐리우드 (순서 인덱스 사용)
    - '씬/Scene 숫자' 폴백

    Args:
        scenario_text: 전체 시나리오 텍스트
        batch_size: 한 배치당 씬 개수 (기본 12)

    Returns:
        [
            {
                "batch_index": 1,
                "first_scene": 1,
                "last_scene": 12,
                "scene_range": "S#1~S#12",
                "scenario_chunk": "...",
                "scene_count": 12,
                "scene_format": "S#" | "EXT/INT" | "FALLBACK"
            },
            ...
        ]
    """
    if not scenario_text:
        return []

    scene_positions = []  # [(scene_label, scene_num, start_char_index), ...]
    scene_format = None

    # ─────────────────────────────────────────
    # 1차 시도: 'S#숫자' 한국 시나리오 형식
    # ─────────────────────────────────────────
    pattern_a = r'(?:^|\n)\s*S\s*#\s*(\d+)\b'
    for m in _re_v27.finditer(pattern_a, scenario_text, flags=_re_v27.MULTILINE):
        scene_num = int(m.group(1))
        actual_start = m.start()
        if actual_start < len(scenario_text) and scenario_text[actual_start] == '\n':
            actual_start += 1
        scene_positions.append((f"S#{scene_num}", scene_num, actual_start))

    if scene_positions:
        scene_format = "S#"
    else:
        # ─────────────────────────────────────────
        # 2차 시도: 'EXT./INT.' 헐리우드 형식 (순서 인덱스 사용)
        # ─────────────────────────────────────────
        pattern_b = r'(?:^|\n)\s*(?:EXT|INT|EXT\./INT|I/E)\.\s+\S'
        seq_idx = 0
        for m in _re_v27.finditer(pattern_b, scenario_text, flags=_re_v27.MULTILINE | _re_v27.IGNORECASE):
            seq_idx += 1
            actual_start = m.start()
            if actual_start < len(scenario_text) and scenario_text[actual_start] == '\n':
                actual_start += 1
            scene_positions.append((f"S#{seq_idx}", seq_idx, actual_start))

        if scene_positions:
            scene_format = "EXT/INT"
        else:
            # ─────────────────────────────────────────
            # 3차 시도: '씬/Scene 숫자' 폴백
            # ─────────────────────────────────────────
            pattern_c = r'(?:^|\n)\s*(?:씬|Scene)\s*(\d+)'
            for m in _re_v27.finditer(pattern_c, scenario_text, flags=_re_v27.MULTILINE | _re_v27.IGNORECASE):
                scene_num = int(m.group(1))
                actual_start = m.start()
                if actual_start < len(scenario_text) and scenario_text[actual_start] == '\n':
                    actual_start += 1
                scene_positions.append((f"S#{scene_num}", scene_num, actual_start))
            if scene_positions:
                scene_format = "FALLBACK"

    if not scene_positions:
        # 그래도 못 찾음 → 단일 배치로 반환
        return [{
            "batch_index": 1,
            "first_scene": 0,
            "last_scene": 0,
            "scene_range": "(씬 인식 실패)",
            "scenario_chunk": scenario_text,
            "scene_count": 0,
            "scene_format": "NONE",
        }]

    # 위치 기준 정렬 후 중복 제거 (S#숫자 모드에서만 의미 있음)
    scene_positions.sort(key=lambda x: x[2])

    # S# 모드일 때만 동일 번호 중복 제거 (EXT/INT는 모든 출현이 별개 씬)
    if scene_format in ("S#", "FALLBACK"):
        seen_nums = set()
        unique_positions = []
        for label, num, pos in scene_positions:
            if num not in seen_nums:
                seen_nums.add(num)
                unique_positions.append((label, num, pos))
        scene_positions = unique_positions

    # 배치 분할
    batches = []
    n_scenes = len(scene_positions)
    n_batches = (n_scenes + batch_size - 1) // batch_size

    for batch_idx in range(n_batches):
        start_i = batch_idx * batch_size
        end_i = min(start_i + batch_size, n_scenes)

        first_label, first_scene_num, first_pos = scene_positions[start_i]
        last_label, last_scene_num, _ = scene_positions[end_i - 1]

        if end_i < n_scenes:
            _, _, chunk_end_pos = scene_positions[end_i]
        else:
            chunk_end_pos = len(scenario_text)

        chunk = scenario_text[first_pos:chunk_end_pos].strip()

        batches.append({
            "batch_index": batch_idx + 1,
            "first_scene": first_scene_num,
            "last_scene": last_scene_num,
            "scene_range": f"{first_label}~{last_label}",
            "scenario_chunk": chunk,
            "scene_count": end_i - start_i,
            "scene_format": scene_format,
        })

    return batches


def run_diagnose_with_auto_batch(client, batch_size: int = 12):
    """v2.7 — 시나리오를 자동 분할하여 배치별 진단 후 결과 병합.

    호출 흐름:
    1. _detect_scene_count로 총 씬 수 파악
    2. _split_scenario_by_scenes로 청크 분할
    3. 각 청크마다 build_diagnose_prompt(batch_info=...) 호출
    4. revision_items[]를 모두 병합하여 통합 diagnose_result 반환

    Args:
        client: Anthropic 클라이언트
        batch_size: 한 배치당 씬 개수 (기본 12)

    Returns:
        통합된 diagnose_result dict, 실패 시 None
    """
    # v2.0/v2.1/v2.2 — 사전 분석 항상 호출 (각 항목 독립 캐싱)
    pre_results = run_v2_pre_analyses(client)

    raw_text = st.session_state.raw_text
    if not raw_text:
        st.error("원본 시나리오가 비어 있습니다.")
        return None

    # 씬 수 감지
    scene_count = _detect_scene_count(raw_text)

    # 분할
    batches = _split_scenario_by_scenes(raw_text, batch_size=batch_size)
    n_batches = len(batches)

    if n_batches == 0:
        st.error("씬 인식에 실패했습니다. 'S#숫자' 형식이 시나리오에 있어야 합니다.")
        return None

    # 단일 배치면 기존 방식과 동일 (배치 정보 없이 호출)
    if n_batches == 1:
        st.info(f"🔍 시나리오 분석: {scene_count}씬 → 단일 배치 처리 (분할 불필요)")
        return _run_diagnose_single(client, raw_text, pre_results, batch_info=None)

    # 다중 배치 처리
    st.info(
        f"🔍 시나리오 분석: **{scene_count}씬** → "
        f"**{n_batches}배치**로 자동 분할 (배치당 약 {batch_size}씬)"
    )

    progress_bar = st.progress(0.0, text="진단 준비 중...")
    merged_target_scenes = []
    summaries = []
    out_of_scope_all = []
    confidences = []
    locked_summaries = []
    conflicts_all = []

    for i, batch in enumerate(batches, start=1):
        progress_bar.progress(
            (i - 1) / n_batches,
            text=f"🔬 배치 {i}/{n_batches} 진단 중... ({batch['scene_range']})"
        )

        batch_info = {
            "batch_index": i,
            "total_batches": n_batches,
            "scene_range": batch["scene_range"],
            "first_scene": batch["first_scene"],
            "last_scene": batch["last_scene"],
            "scene_format": batch.get("scene_format", "S#"),
        }

        result = _run_diagnose_single(
            client,
            raw_text=batch["scenario_chunk"],
            pre_results=pre_results,
            batch_info=batch_info,
            retry_count=3,
        )

        if not result:
            st.warning(
                f"⚠️ 배치 {i}/{n_batches} ({batch['scene_range']}) 진단 실패. "
                f"이 배치는 빈 결과로 처리됩니다."
            )
            continue

        rp = result.get("revision_plan", {})
        merged_target_scenes.extend(rp.get("target_scenes", []))
        s = rp.get("summary", "").strip()
        if s:
            summaries.append(f"[{batch['scene_range']}] {s}")
        oos = rp.get("out_of_scope", [])
        if oos:
            out_of_scope_all.extend(oos)
        c = rp.get("confidence")
        if isinstance(c, (int, float)):
            confidences.append(c)
        ls = rp.get("locked_summary", "").strip()
        if ls:
            locked_summaries.append(ls)
        confs = rp.get("conflicts", [])
        if confs:
            conflicts_all.extend(confs)

    progress_bar.progress(1.0, text="✅ 모든 배치 진단 완료. 결과 통합 중...")

    # 통합 결과 구성
    avg_confidence = round(sum(confidences) / len(confidences), 1) if confidences else 7.0
    merged_result = {
        "revision_plan": {
            "summary": "\n\n".join(summaries) if summaries else "자동 배치 진단 완료.",
            "locked_summary": "\n".join(locked_summaries) if locked_summaries else "",
            "conflicts": conflicts_all,
            "target_scenes": merged_target_scenes,
            "out_of_scope": out_of_scope_all,
            "confidence": avg_confidence,
            "estimated_scene_count": str(len(merged_target_scenes)),
            "total_scenes": len(merged_target_scenes),
            "recommended_batches": max(1, (len(merged_target_scenes) + 9) // 10),
            "batch_strategy": (
                f"v2.7 자동 분할 진단: 총 {n_batches}개 진단 배치를 통합 → "
                f"수정 대상 {len(merged_target_scenes)}개 씬 식별. "
                f"REVISE 단계는 별도 배치 정책 적용."
            ),
            "auto_batch_diagnose": {
                "diagnose_batch_count": n_batches,
                "diagnose_batch_size": batch_size,
                "total_scenes_detected": scene_count,
            }
        }
    }

    progress_bar.empty()
    st.success(
        f"✅ 자동 배치 진단 완료 — {n_batches}배치 → "
        f"수정 대상 **{len(merged_target_scenes)}개 씬** 식별"
    )
    return merged_result


def _run_diagnose_single(client, raw_text: str, pre_results: dict,
                         batch_info: dict = None, retry_count: int = 1):
    """단일 진단 호출 (배치 단위 또는 비배치). 재시도 로직 포함."""
    prompt_text = build_diagnose_prompt(
        raw_text=raw_text,
        instruction=st.session_state.instruction,
        locked=st.session_state.locked,
        genre=st.session_state.genre,
        intensity=st.session_state.intensity,
        profession_input=st.session_state.profession_input,
        period_key=st.session_state.period_key,
        historical_type=st.session_state.historical_type,
        fact_based=st.session_state.fact_based,
        tone_dna=pre_results.get("tone_dna"),
        diff_analysis=pre_results.get("diff_analysis"),
        distribution_diagnostic=pre_results.get("distribution_diagnostic"),
        rewrite_metadata=pre_results.get("rewrite_metadata"),
        genre_dna=pre_results.get("genre_dna"),
        section_mode=st.session_state.section_mode,
        protected_ranges=st.session_state.protected_ranges,
        revision_ranges=st.session_state.revision_ranges,
        cascade_analysis=st.session_state.cascade_analysis,
        boundary_info=st.session_state.boundary_info,
        batch_info=batch_info,
    )

    for attempt in range(retry_count):
        raw = call_claude(client, prompt_text, model=MODEL_ANALYZE, max_tokens=32000)
        if not raw:
            continue
        parsed = parse_json(raw)
        if parsed:
            return parsed
        if attempt < retry_count - 1:
            # 재시도 전 잠깐 대기 (rate limit 대비)
            import time as _t
            _t.sleep(1.5)

    return None


# =================================================================
# ★★★ v2.8 BEAT-AWARE DIAGNOSE 시스템 ★★★
# 71→100씬 같은 대규모 확장 시 비트별 약점 진단 + 자동 분배
# =================================================================

def _run_pre_diagnose_beat_map(client, scenario_text: str, genre: str = "로맨틱 코미디",
                                retry_count: int = 2):
    """v2.8 Pre-Diagnose: 시나리오 전체를 15-Beat 구조에 매핑.

    출력은 작다 (씬 ID 리스트 + 짧은 진단 = 약 3~5K 토큰).
    토큰 잘림 위험 없음 — 단일 호출로 처리.

    Args:
        client: Anthropic 클라이언트
        scenario_text: 전체 시나리오 텍스트
        genre: 작품 장르
        retry_count: 재시도 횟수

    Returns:
        beat_map dict 또는 None
    """
    if not scenario_text or not client:
        return None

    prompt_text = build_beat_mapping_prompt(
        scenario_text=scenario_text,
        genre=genre,
    )

    for attempt in range(retry_count):
        # 비트 매핑은 출력이 작으므로 max_tokens=8000으로 충분
        raw = call_claude(client, prompt_text, model=MODEL_ANALYZE, max_tokens=8000)
        if not raw:
            continue
        parsed = parse_json(raw)
        if parsed and parsed.get("beat_mapping"):
            return parsed
        if attempt < retry_count - 1:
            import time as _t
            _t.sleep(1.5)

    return None


def run_diagnose_with_beat_aware_batch(client, batch_size: int = 12,
                                        target_added_scenes: int = 0):
    """v2.8 — 비트 인식 자동 배치 분할 진단.

    호출 흐름:
    1. Pre-Diagnose: 시나리오 전체를 15-Beat 매핑 (단일 호출)
    2. distribute_added_scenes_across_beats: target_added_scenes를 비트별 분배
    3. 시나리오를 N씬 단위로 분할 (v2.7 로직)
    4. 각 배치마다 build_diagnose_prompt(batch_info, beat_map, beat_distribution, target_added_scenes) 호출
    5. 결과 병합 → 통합 diagnose_result + beat_map + distribution 반환

    Args:
        client: Anthropic 클라이언트
        batch_size: DIAGNOSE 배치당 씬 개수 (기본 12)
        target_added_scenes: 추가할 총 씬 수 (예: 29)

    Returns:
        통합된 diagnose_result dict (beat_map, beat_distribution 포함)
    """
    pre_results = run_v2_pre_analyses(client)

    raw_text = st.session_state.raw_text
    if not raw_text:
        st.error("원본 시나리오가 비어 있습니다.")
        return None

    scene_count = _detect_scene_count(raw_text)
    if scene_count == 0:
        st.error("씬 인식 실패. 씬 헤더 형식을 확인해주세요.")
        return None

    # ─────────────────────────────────────────────────
    # Phase 1: Pre-Diagnose 비트 매핑
    # ─────────────────────────────────────────────────
    st.info(
        f"🎯 **v2.8 Beat-Aware Diagnose 시작** — "
        f"시나리오 {scene_count}씬 → {target_added_scenes}씬 추가 목표"
    )

    with st.spinner("🎬 Phase 1: 시나리오 전체를 15-Beat 구조에 매핑 중... (Sonnet 4.6)"):
        beat_map = _run_pre_diagnose_beat_map(
            client=client,
            scenario_text=raw_text,
            genre=st.session_state.genre,
        )

    if not beat_map:
        st.error("❌ 비트 매핑 실패. v2.7 모드로 폴백합니다.")
        return run_diagnose_with_auto_batch(client, batch_size=batch_size)

    # 비트 매핑 결과 캐싱
    st.session_state.beat_map = beat_map

    # 분배 산출
    beat_distribution = distribute_added_scenes_across_beats(
        beat_map=beat_map,
        target_added=target_added_scenes,
    )
    st.session_state.beat_distribution = beat_distribution

    # 비트 매핑 결과 표시
    bm = beat_map.get("beat_mapping", {})
    weak = beat_map.get("weak_beats", [])
    missing = beat_map.get("missing_essentials", [])

    st.success(
        f"✅ Phase 1 완료 — 비트 매핑 (총 {beat_map.get('total_scenes', scene_count)}씬, "
        f"장르 준수도 {beat_map.get('genre_compliance_score', 'N/A')}/10)"
    )

    with st.expander("📊 비트 매핑 결과 보기 (15-Beat)", expanded=False):
        for k, name, pct, desc in SAVE_THE_CAT_15_BEATS:
            beat_info = bm.get(k, {})
            sids = beat_info.get("scene_ids", [])
            sc = beat_info.get("scene_count", 0)
            strength = beat_info.get("strength", "")
            strength_emoji = {
                "STRONG": "🟢", "ADEQUATE": "🟡",
                "WEAK": "🟠", "MISSING": "🔴"
            }.get(strength, "⚪")
            sids_text = ", ".join(sids[:6]) + (f"... (+{len(sids)-6})" if len(sids) > 6 else "")
            st.markdown(
                f"**{strength_emoji} {name}** ({pct}) — {sc}씬 / {strength}  \n"
                f"<span style='color:#666; font-size:0.85em;'>{sids_text}</span>  \n"
                f"<span style='color:#999; font-size:0.85em;'>{beat_info.get('function_check', '')}</span>",
                unsafe_allow_html=True
            )

        if weak:
            st.markdown("---\n**🔴 약점 비트:**")
            for wb in weak:
                st.markdown(
                    f"- **{wb.get('beat_name', '')}**: {wb.get('current_scenes', 0)}씬 → "
                    f"권장 {wb.get('recommended_min_scenes', 0)}씬 (deficit {wb.get('deficit', 0)}) — "
                    f"{wb.get('weakness_reason', '')}"
                )

        if missing:
            st.markdown("---\n**❌ 누락 필수 요소:**")
            for me in missing:
                st.markdown(
                    f"- **{me.get('essential', '')}** ({me.get('severity', '')}, "
                    f"위치: {me.get('located_in_beat', '')}): {me.get('fix_direction', '')}"
                )

    # 분배 결과 표시
    st.markdown("---")
    st.markdown(f"**🎯 추가 씬 분배 정책 — 총 +{target_added_scenes}씬**")
    dist = beat_distribution.get("distribution", {})
    if dist:
        cols = st.columns(min(len(dist), 4))
        for i, (beat_key, count) in enumerate(sorted(dist.items(), key=lambda x: -x[1])):
            beat_kr = beat_key
            for k, name, pct, desc in SAVE_THE_CAT_15_BEATS:
                if k == beat_key:
                    beat_kr = name
                    break
            with cols[i % len(cols)]:
                st.metric(beat_kr, f"+{count}씬")

    # ─────────────────────────────────────────────────
    # Phase 2: 비트 인식 배치 진단
    # ─────────────────────────────────────────────────
    batches = _split_scenario_by_scenes(raw_text, batch_size=batch_size)
    n_batches = len(batches)

    st.info(
        f"🔬 **Phase 2: 비트 인식 배치 진단 시작** — "
        f"{n_batches}배치로 분할, 각 배치마다 비트 분배 정책에 따라 ADD 위치 제안"
    )

    progress_bar = st.progress(0.0, text="비트 인식 진단 준비 중...")
    merged_target_scenes = []
    summaries = []
    out_of_scope_all = []
    confidences = []
    locked_summaries = []
    conflicts_all = []

    for i, batch in enumerate(batches, start=1):
        progress_bar.progress(
            (i - 1) / n_batches,
            text=f"🔬 배치 {i}/{n_batches} 비트 인식 진단 중... ({batch['scene_range']})"
        )

        batch_info = {
            "batch_index": i,
            "total_batches": n_batches,
            "scene_range": batch["scene_range"],
            "first_scene": batch["first_scene"],
            "last_scene": batch["last_scene"],
            "scene_format": batch.get("scene_format", "S#"),
        }

        prompt_text = build_diagnose_prompt(
            raw_text=batch["scenario_chunk"],
            instruction=st.session_state.instruction,
            locked=st.session_state.locked,
            genre=st.session_state.genre,
            intensity=st.session_state.intensity,
            profession_input=st.session_state.profession_input,
            period_key=st.session_state.period_key,
            historical_type=st.session_state.historical_type,
            fact_based=st.session_state.fact_based,
            tone_dna=pre_results.get("tone_dna"),
            diff_analysis=pre_results.get("diff_analysis"),
            distribution_diagnostic=pre_results.get("distribution_diagnostic"),
            rewrite_metadata=pre_results.get("rewrite_metadata"),
            genre_dna=pre_results.get("genre_dna"),
            section_mode=st.session_state.section_mode,
            protected_ranges=st.session_state.protected_ranges,
            revision_ranges=st.session_state.revision_ranges,
            cascade_analysis=st.session_state.cascade_analysis,
            boundary_info=st.session_state.boundary_info,
            batch_info=batch_info,
            beat_map=beat_map,
            beat_distribution=beat_distribution,
            target_added_scenes=target_added_scenes,
        )

        result = None
        for attempt in range(3):
            raw = call_claude(client, prompt_text, model=MODEL_ANALYZE, max_tokens=32000)
            if not raw:
                continue
            parsed = parse_json(raw)
            if parsed:
                result = parsed
                break
            if attempt < 2:
                import time as _t
                _t.sleep(1.5)

        if not result:
            st.warning(f"⚠️ 배치 {i}/{n_batches} 진단 실패. 빈 결과로 처리.")
            continue

        rp = result.get("revision_plan", {})
        merged_target_scenes.extend(rp.get("target_scenes", []))
        s = rp.get("summary", "").strip()
        if s:
            summaries.append(f"[{batch['scene_range']}] {s}")
        oos = rp.get("out_of_scope", [])
        if oos:
            out_of_scope_all.extend(oos)
        c = rp.get("confidence")
        if isinstance(c, (int, float)):
            confidences.append(c)
        ls = rp.get("locked_summary", "").strip()
        if ls:
            locked_summaries.append(ls)
        confs = rp.get("conflicts", [])
        if confs:
            conflicts_all.extend(confs)

    progress_bar.progress(1.0, text="✅ 모든 배치 진단 완료. 결과 통합 중...")

    # 통합 결과
    avg_confidence = round(sum(confidences) / len(confidences), 1) if confidences else 7.0

    # 추가 씬(ADD) vs 수정 씬(REWRITE) 카운트
    add_count = sum(1 for s in merged_target_scenes if s.get("type") == "ADD")
    rewrite_count = sum(1 for s in merged_target_scenes if s.get("type") == "REWRITE")
    other_count = len(merged_target_scenes) - add_count - rewrite_count

    merged_result = {
        "revision_plan": {
            "summary": "\n\n".join(summaries) if summaries else "v2.8 비트 인식 진단 완료.",
            "locked_summary": "\n".join(locked_summaries) if locked_summaries else "",
            "conflicts": conflicts_all,
            "target_scenes": merged_target_scenes,
            "out_of_scope": out_of_scope_all,
            "confidence": avg_confidence,
            "estimated_scene_count": str(len(merged_target_scenes)),
            "total_scenes": len(merged_target_scenes),
            "recommended_batches": max(1, (len(merged_target_scenes) + 4) // 5),
            "batch_strategy": (
                f"v2.8 Beat-Aware 진단: 총 {n_batches}개 배치 → "
                f"ADD {add_count}개 + REWRITE {rewrite_count}개 + 기타 {other_count}개"
            ),
            "auto_batch_diagnose": {
                "diagnose_batch_count": n_batches,
                "diagnose_batch_size": batch_size,
                "total_scenes_detected": scene_count,
            },
            # v2.8 추가 메타
            "beat_aware": {
                "beat_map": beat_map,
                "beat_distribution": beat_distribution,
                "target_added_scenes": target_added_scenes,
                "actual_add_count": add_count,
                "actual_rewrite_count": rewrite_count,
                "deficit_or_surplus": add_count - target_added_scenes,
            }
        }
    }

    progress_bar.empty()

    # 최종 결과 표시
    if add_count == target_added_scenes:
        st.success(
            f"✅ v2.8 Beat-Aware 진단 완료 — ADD **{add_count}씬** (목표 {target_added_scenes}씬 정확 일치) "
            f"+ REWRITE {rewrite_count}씬"
        )
    elif abs(add_count - target_added_scenes) <= 3:
        st.success(
            f"✅ v2.8 Beat-Aware 진단 완료 — ADD **{add_count}씬** "
            f"(목표 {target_added_scenes}씬, 차이 {add_count - target_added_scenes:+d}) "
            f"+ REWRITE {rewrite_count}씬"
        )
    else:
        st.warning(
            f"⚠️ v2.8 Beat-Aware 진단 완료 — ADD **{add_count}씬** "
            f"(목표 {target_added_scenes}씬, 차이 {add_count - target_added_scenes:+d}) "
            f"+ REWRITE {rewrite_count}씬. 수정 지시문 또는 배치 사이즈 조정 권장."
        )

    return merged_result


# =================================================================
# ★★★ v2.9 BEAT EXPANSION MODE ★★★
# 비트 보강 확장 모드 = Beat-Aware Diagnose + LOCKED 영역 강제 차단
# 71→100씬 같은 작업: 보호 구간(S#1~S#25)은 진단·집필에서 완전 제외
# =================================================================

def _parse_scene_range_to_int(range_str: str) -> int:
    """'S#25' → 25 같은 변환. 실패 시 0 반환."""
    if not range_str:
        return 0
    import re as _re_temp
    m = _re_temp.search(r'\d+', str(range_str))
    return int(m.group()) if m else 0


def _filter_target_scenes_against_protected(target_scenes: list,
                                              protected_ranges: list) -> tuple:
    """target_scenes에서 LOCKED 영역에 속하는 씬을 필터링.

    Args:
        target_scenes: 진단 결과의 target_scenes 리스트
        protected_ranges: [{"from":"S#1","to":"S#25",...}, ...]

    Returns:
        (filtered_scenes, removed_scenes) 튜플
    """
    if not protected_ranges:
        return target_scenes, []

    # 보호 범위를 (from_int, to_int) 튜플로 변환
    prot_intervals = []
    for pr in protected_ranges:
        f = _parse_scene_range_to_int(pr.get("from", ""))
        t = _parse_scene_range_to_int(pr.get("to", ""))
        if f > 0 and t > 0:
            prot_intervals.append((min(f, t), max(f, t)))

    if not prot_intervals:
        return target_scenes, []

    filtered = []
    removed = []
    for ts in target_scenes:
        scene_id = ts.get("scene_id", "")
        ts_type = ts.get("type", "REWRITE")

        # ADD의 경우 insert_after 기준
        if ts_type == "ADD":
            # ADD는 "이 씬 뒤에 삽입"이므로 insert_after 위치가 LOCKED 끝 이전이면 작업 영역 첫 씬으로 우회
            ia = ts.get("insert_after", "")
            ia_num = _parse_scene_range_to_int(ia)
            if ia_num <= 0:
                # insert_after 인식 못 함 → 일단 통과
                filtered.append(ts)
                continue
            # 보호 구간 내부에 ADD 위치가 있으면 → 보호 끝 직후로 우회
            in_protected = False
            max_prot_end = 0
            for f_int, t_int in prot_intervals:
                if f_int <= ia_num <= t_int:
                    in_protected = True
                    max_prot_end = max(max_prot_end, t_int)
            if in_protected:
                # 우회: insert_after를 보호 영역 마지막 씬으로 변경
                ts = dict(ts)
                ts["insert_after"] = f"S#{max_prot_end}"
                ts["_redirected_from_locked"] = ia
                filtered.append(ts)
            else:
                filtered.append(ts)
        else:
            # REWRITE/DELETE/MERGE/SPLIT: scene_id 위치가 LOCKED 영역 내면 차단
            sid_num = _parse_scene_range_to_int(scene_id)
            if sid_num <= 0:
                filtered.append(ts)
                continue
            in_protected = any(f_int <= sid_num <= t_int for f_int, t_int in prot_intervals)
            if in_protected:
                removed.append(ts)
            else:
                filtered.append(ts)

    return filtered, removed


def run_diagnose_beat_expansion(client, batch_size: int = 12,
                                 target_added_scenes: int = 0):
    """v2.9 비트 보강 확장 모드 — Beat-Aware Diagnose + LOCKED 강제 차단.

    동작:
    1. v2.8 run_diagnose_with_beat_aware_batch 호출
    2. 결과의 target_scenes에서 protected_ranges 영역 침범 항목 필터링
    3. ADD 위치가 LOCKED 영역이면 작업 영역 시작점으로 자동 우회
    4. REWRITE/DELETE 등이 LOCKED 영역이면 out_of_scope로 이동

    Args:
        client: Anthropic 클라이언트
        batch_size: DIAGNOSE 배치 사이즈
        target_added_scenes: 추가할 총 씬 수

    Returns:
        통합된 diagnose_result dict (LOCKED 영역 침범 차단됨)
    """
    protected_ranges = st.session_state.get("protected_ranges", []) or []
    revision_ranges = st.session_state.get("revision_ranges", []) or []

    # 보호/작업 영역 안내
    if protected_ranges:
        prot_str = ", ".join(
            f"{r.get('from','')}~{r.get('to','')}" for r in protected_ranges
        )
    else:
        prot_str = "(없음)"

    if revision_ranges:
        rev_str = ", ".join(
            f"{r.get('from','')}~{r.get('to','')}" for r in revision_ranges
        )
    else:
        rev_str = "(자동 산출)"

    st.info(
        f"🎯 **v2.9 비트 보강 확장 모드 시작**\n\n"
        f"🔒 보호 구간: {prot_str} (진단·집필 완전 제외)\n\n"
        f"✏️ 작업 영역: {rev_str} (ADD/REWRITE 가능)\n\n"
        f"➕ 추가 목표: +{target_added_scenes}씬"
    )

    # ─────────────────────────────────────────────────
    # Phase 1~4: v2.8 Beat-Aware Diagnose 호출
    # ─────────────────────────────────────────────────
    base_result = run_diagnose_with_beat_aware_batch(
        client=client,
        batch_size=batch_size,
        target_added_scenes=target_added_scenes,
    )

    if not base_result:
        st.error("❌ Beat-Aware 진단 실패. v2.8 모드로 폴백합니다.")
        return None

    # ─────────────────────────────────────────────────
    # Phase 5: LOCKED 영역 강제 차단
    # ─────────────────────────────────────────────────
    rp = base_result.get("revision_plan", {})
    target_scenes = rp.get("target_scenes", []) or []

    if not protected_ranges:
        # 보호 구간 없으면 그대로 반환
        return base_result

    filtered, removed = _filter_target_scenes_against_protected(
        target_scenes, protected_ranges
    )

    # 우회된 ADD 카운트
    redirected_count = sum(
        1 for s in filtered if s.get("_redirected_from_locked")
    )

    # 통계
    original_count = len(target_scenes)
    final_count = len(filtered)
    removed_count = len(removed)

    # 결과 반영
    rp["target_scenes"] = filtered

    # out_of_scope에 차단된 항목 추가
    out_of_scope = rp.get("out_of_scope", []) or []
    for r in removed:
        out_of_scope.append(
            f"⚠️ LOCKED 영역 침범 차단: {r.get('scene_id', '')} "
            f"({r.get('type', '')}) — {r.get('preservation_notes', '')[:80]}"
        )
    rp["out_of_scope"] = out_of_scope

    # 메타데이터 업데이트
    rp["total_scenes"] = final_count
    rp["estimated_scene_count"] = str(final_count)

    # v2.9 메타 추가
    rp["beat_expansion"] = {
        "mode": "expansion",
        "protected_ranges": protected_ranges,
        "revision_ranges": revision_ranges,
        "target_added_scenes": target_added_scenes,
        "original_target_scenes_count": original_count,
        "filtered_target_scenes_count": final_count,
        "removed_count_due_to_locked": removed_count,
        "redirected_count_due_to_locked": redirected_count,
    }

    # 최종 ADD/REWRITE 카운트
    add_count = sum(1 for s in filtered if s.get("type") == "ADD")
    rewrite_count = sum(1 for s in filtered if s.get("type") == "REWRITE")

    # 사용자 안내
    if removed_count > 0 or redirected_count > 0:
        msg_parts = []
        if removed_count > 0:
            msg_parts.append(f"🚫 LOCKED 침범 차단: {removed_count}건 (out_of_scope 이동)")
        if redirected_count > 0:
            msg_parts.append(f"↪️ ADD 위치 자동 우회: {redirected_count}건 (작업 영역 시작점으로 이동)")
        st.info("**v2.9 LOCKED 차단 결과**\n\n" + "\n\n".join(msg_parts))

    if add_count == target_added_scenes:
        st.success(
            f"✅ v2.9 비트 보강 확장 진단 완료 — ADD **{add_count}씬** "
            f"(목표 정확 일치) + REWRITE {rewrite_count}씬"
        )
    else:
        diff = add_count - target_added_scenes
        st.success(
            f"✅ v2.9 비트 보강 확장 진단 완료 — ADD **{add_count}씬** "
            f"(목표 {target_added_scenes}씬, 차이 {diff:+d}) "
            f"+ REWRITE {rewrite_count}씬"
        )

    return base_result


def run_diagnose(client):
    """Stage 1: 진단 + 수정 플랜 생성.

    Fast Path 우선순위 (v2.9 업데이트):
    0. ★ v2.9 expansion 모드 (비트 보강 확장) → Beat-Aware + LOCKED 차단
    1. 구간 모드(이어쓰기/부분수정) → 코드 자동 생성
    2. Rewrite JSON 있음 (REWRITE/ADD 제안) → 코드 자동 생성
    3. 그 외 (순수 전체 각색) → AI 진단
       - target_added_scenes > 0 → v2.8 Beat-Aware
       - target_added_scenes == 0 → v2.7 자동 배치
    """

    # v2.0/v2.1/v2.2 — 사전 분석 항상 호출 (각 항목 독립 캐싱)
    pre_results = run_v2_pre_analyses(client)

    work_mode = st.session_state.get("work_mode", "")
    target_added_raw = st.session_state.get("target_added_scenes", 0)
    try:
        target_added = int(target_added_raw) if target_added_raw else 0
    except (ValueError, TypeError):
        target_added = 0
    batch_size = st.session_state.get("diagnose_batch_size", 12)

    # ★★★ v3.0 라우팅 디버그 안내 — 어느 Fast Path를 탔는지 명시
    st.markdown(
        f'<div style="background:#F3F4F6; padding:8px 12px; border-radius:4px; '
        f'border-left:3px solid #6B7280; margin:6px 0; font-size:0.82rem; color:#374151;">'
        f'🔍 <b>라우팅 상태:</b> work_mode=<code>{work_mode}</code>, '
        f'target_added=<code>{target_added}</code>, '
        f'section_mode=<code>{st.session_state.get("section_mode", False)}</code>, '
        f'protected={len(st.session_state.get("protected_ranges", []))}건, '
        f'revision={len(st.session_state.get("revision_ranges", []))}건'
        f'</div>',
        unsafe_allow_html=True
    )

    # ★ Fast Path 0 — v3.0 비트 보강 확장 모드 (최우선, 강화 조건)
    if work_mode == "expansion" and target_added > 0:
        st.markdown(
            '<div style="background:#FCE7F3; padding:8px 12px; border-radius:4px; '
            'border-left:3px solid #EC4899; margin:6px 0; font-size:0.85rem;">'
            '🎯 <b>Fast Path 0 진입:</b> v3.0 비트 보강 확장 모드 → run_diagnose_beat_expansion()'
            '</div>',
            unsafe_allow_html=True
        )
        return run_diagnose_beat_expansion(
            client,
            batch_size=batch_size,
            target_added_scenes=target_added,
        )

    # ★ Fast Path 1 — 구간 모드 (이어쓰기/부분수정만)
    # ★ v3.0: expansion 모드는 절대 여기로 떨어지지 않도록 명시 차단
    if (st.session_state.section_mode
            and st.session_state.revision_ranges
            and work_mode in ("continuation", "partial")):
        st.success(
            f"⚡ Fast Path 1 진입: 구간 모드(<b>{work_mode}</b>) — "
            f"진단 자동 생성 (AI 호출 없이 즉시 완료)"
        )
        return _build_auto_diagnose_for_section_mode()

    # ★ Fast Path 2 — Rewrite Engine JSON 흡수 시
    rewrite_meta = pre_results.get("rewrite_metadata") or st.session_state.rewrite_metadata
    if rewrite_meta and isinstance(rewrite_meta, dict):
        has_rewrite = bool(rewrite_meta.get("rewrite_suggestions"))
        has_add = bool(rewrite_meta.get("add_suggestions"))
        has_weak = bool(rewrite_meta.get("weak_zone_scenes"))
        has_moon = bool(rewrite_meta.get("moon_opinion_text"))

        # REWRITE 제안 또는 ADD 제안 또는 weak_zone 중 하나라도 있으면 Fast Path 사용
        if has_rewrite or has_add or has_weak:
            r_count = len(rewrite_meta.get("rewrite_suggestions", []))
            a_count = len(rewrite_meta.get("add_suggestions", []))
            w_count = len(rewrite_meta.get("weak_zone_scenes", []))
            st.success(
                f"⚡ Rewrite Engine 처방 자동 흡수 — 진단 즉시 완료 (AI 호출 없이)\n\n"
                f"  • REWRITE 제안 {r_count}개 → priority HIGH 자동 등록\n"
                f"  • ADD 제안 {a_count}개 → type=ADD 자동 등록\n"
                f"  • weak_zone {w_count}개 → priority HIGH 자동 격상\n"
                f"  • MOON 의견 {'반영됨' if has_moon else '없음'}"
            )
            return _build_auto_diagnose_from_rewrite_metadata(rewrite_meta)

    # ★ Fast Path 3 — 일반 전체 각색
    if target_added > 0:
        # v2.8 Beat-Aware Diagnose 모드
        return run_diagnose_with_beat_aware_batch(
            client,
            batch_size=batch_size,
            target_added_scenes=target_added,
        )
    else:
        # v2.7 자동 배치 분할 모드
        return run_diagnose_with_auto_batch(client, batch_size=batch_size)


def run_revise_batch(client, batch_index: int, batch_scenes: list, total_batches: int):
    """Stage 2: Opus로 특정 배치만 집필.

    Args:
        client: Anthropic 클라이언트
        batch_index: 배치 번호 (1부터)
        batch_scenes: 이번 배치에서 처리할 씬 리스트
        total_batches: 전체 배치 수
    """
    prompt_text = build_revise_prompt(
        raw_text=st.session_state.raw_text,
        diagnose_result=st.session_state.diagnose_result,
        genre=st.session_state.genre,
        intensity=st.session_state.intensity,
        locked=st.session_state.locked,
        profession_input=st.session_state.profession_input,
        period_key=st.session_state.period_key,
        historical_type=st.session_state.historical_type,
        fact_based=st.session_state.fact_based,
        batch_scenes=batch_scenes,
        batch_index=batch_index,
        total_batches=total_batches,
        tone_dna=st.session_state.tone_dna,
        diff_analysis=st.session_state.diff_analysis,
        genre_dna=st.session_state.genre_dna,
        section_mode=st.session_state.section_mode,
        protected_ranges=st.session_state.protected_ranges,
        revision_ranges=st.session_state.revision_ranges,
        cascade_analysis=st.session_state.cascade_analysis,
        boundary_info=st.session_state.boundary_info,
    )
    raw = call_claude(client, prompt_text, model=MODEL_WRITE, max_tokens=32000)
    if not raw:
        return None
    parsed = parse_json(raw)
    if not parsed:
        return None

    # ★ 포맷 검증 — revised_content가 통짜 출력인지 체크 + 자동 보강
    parsed = _validate_and_fix_revised_format(parsed)

    return parsed



# ═══════════════════════════════════════════════════════════
# ★ Writer Engine v3.1.3 자산 이식 — 지문 자동 분단 시스템
# (Writer Engine main.py 63~204행 그대로 가져옴)
# ═══════════════════════════════════════════════════════════

# 분단 임계값 — Writer Engine 데이터 분석 기반 (실제 결과물 481개 단락 분포 검증)
_ACTION_SPLIT_CHAR_THRESHOLD = 150     # 자수 임계
_ACTION_SPLIT_SENTENCE_THRESHOLD = 7   # 문장수 임계
_ACTION_SPLIT_HARD_CHARS = 240         # 하드 임계 (이 이상은 무조건 분단)
_ACTION_SPLIT_HARD_SENTENCES = 9       # 하드 문장수 임계


def _split_sentences(text: str):
    """한국어 지문을 문장 단위로 쪼갠다. 마침표·물음표·느낌표 + 공백 기준.
    문장 부호 뒤 공백이 없는 경우(약어 등)는 분리하지 않는다.
    """
    import re as _re
    parts = _re.split(r'(?<=[.!?])\s+', text.strip())
    return [s.strip() for s in parts if s.strip()]


def _detect_paragraph_break_index(sentences: list) -> int:
    """
    문장 리스트에서 가장 자연스러운 분단 위치(인덱스)를 찾는다.
    분단 트리거 4종에 따라 우선순위로 탐색.
    찾지 못하면 -1 반환.

    트리거 우선순위:
      1) 시간 압축 종료: "수업이 진행된다" / "수업이 끝난다" 같은 큰 사건 변화
      2) 동작 주체 변경: 단체 → 개인, 또는 인물 ↔ 다른 인물
      3) 카메라 시점 변경: 인물 동작 ↔ 정물·공간 묘사 (지시 대명사 없는 명사구 시작)
    """
    import re as _re
    n = len(sentences)
    if n < 4:
        return -1

    candidates = []

    for i in range(2, n - 1):
        cur = sentences[i]
        prev = sentences[i - 1]
        score = 0

        # 트리거 1: 시간 압축 / 큰 상황 전환 키워드
        time_break_patterns = [
            r'수업이?\s*(끝나|진행)',
            r'시간이\s*(흐른|지난|경과)',
            r'(다음\s*날|이튿날|새벽|아침|저녁|밤)',
            r'몇\s*(분|시간|일)\s*(후|뒤)',
            r'(직후|잠시\s*후|곧)',
        ]
        for pat in time_break_patterns:
            if _re.search(pat, cur):
                score += 10
                break

        # 트리거 2: 인물명 + 단독/혼자/홀로 (주체 전환 신호)
        if _re.search(r'(혼자|홀로|단독)', cur):
            score += 6

        # 트리거 3: 정물·공간 묘사 시작 (인물 주어 없는 명사구)
        space_patterns = [
            r'^(긴\s|넓은\s|좁은\s|텅\s|빈\s|새|작은\s|커다란\s)',
            r'^(테이블|벽|창|문|바닥|천장|복도|골목)\s',
            r'^(아일랜드|카운터|책상|의자|침대|소파)\s',
            r'^[가-힣]+\s위에',  # "~ 위에" 시작
        ]
        prev_has_actor = bool(_re.search(r'[가-힣]{2,4}이\s', prev) or _re.search(r'[가-힣]{2,4}가\s', prev))
        cur_is_space = any(_re.search(pat, cur) for pat in space_patterns)
        if prev_has_actor and cur_is_space:
            score += 5

        # 트리거 4: 인물 주체 변경 ("A가 ~한다." → "B가 ~한다.")
        prev_actor = _re.match(r'^([가-힣]{2,4})(이|가)\s', prev)
        cur_actor = _re.match(r'^([가-힣]{2,4})(이|가)\s', cur)
        if prev_actor and cur_actor and prev_actor.group(1) != cur_actor.group(1):
            score += 4

        # 위치 보정: 단락의 정중앙 근처가 가장 자연스러운 분단 위치
        center_distance = abs(i - n // 2)
        position_bonus = max(0, 3 - center_distance)
        score += position_bonus

        if score >= 5:
            candidates.append((i, score))

    if not candidates:
        return -1

    candidates.sort(key=lambda x: -x[1])
    return candidates[0][0]


# ═══════════════════════════════════════════════════════════
# ★ v3.2 — 시나리오 줄간격 후처리 (Writer Engine v3.5.1 자산 이식)
# 지문↔대사 사이에 빈 줄을 자동 삽입한다 (한국 시나리오 표준 포맷).
# AI가 prompt.py 지시를 따라 빈 줄을 넣어 출력하는 것이 1차 안전망이고,
# DOCX 빌더의 prev_block_type 추적이 2차 안전망,
# 이 함수는 누락 시 보완하는 3차 안전망이다.
# ═══════════════════════════════════════════════════════════

def _normalize_screenplay_blank_lines(text: str) -> str:
    """시나리오 본문에서 지문↔대사 사이 빈 줄을 보정한다.

    규칙:
    - 지문 다음 줄이 대사면 사이에 빈 줄 1개
    - 대사 다음 줄이 지문이면 사이에 빈 줄 1개
    - 같은 화자/다른 화자 대사 연속은 빈 줄 없이 유지
    - 씬 헤딩 직전/직후는 기존 빈 줄 처리 유지
    - 이미 빈 줄이 있으면 추가 삽입 안 함 (중복 방지)
    """
    import re as _re_norm

    # 라인 분류 함수
    heading_pat = _re_norm.compile(r'^S#\d+', _re_norm.UNICODE)
    # 대사 패턴: "캐릭터명\t\t대사" (탭 1~3개 허용)
    dialogue_pat = _re_norm.compile(r'^[^\t]+\t{1,}\S', _re_norm.UNICODE)

    def line_type(line: str) -> str:
        s = line.rstrip()
        if not s.strip():
            return "blank"
        if heading_pat.match(s.strip()):
            return "scene"
        if dialogue_pat.match(s):
            return "dialogue"
        return "action"

    lines = text.split('\n')
    result = []
    for idx, line in enumerate(lines):
        cur_type = line_type(line)
        # 이전 의미 있는 라인 타입 찾기 (빈 줄 건너뛰기)
        prev_meaningful = None
        for r_line in reversed(result):
            r_type = line_type(r_line)
            if r_type != "blank":
                prev_meaningful = r_type
                break

        # 직전 줄이 빈 줄인지 (이미 분리된 상태인지)
        already_separated = bool(result) and not result[-1].strip()

        # 빈 줄 삽입 결정
        need_blank = False
        if prev_meaningful and not already_separated:
            # 지문 → 대사 또는 대사 → 지문 전환
            if (prev_meaningful == "action" and cur_type == "dialogue") or \
               (prev_meaningful == "dialogue" and cur_type == "action"):
                need_blank = True

        if need_blank:
            result.append("")
        result.append(line)

    return '\n'.join(result)


def _split_action_paragraph(text: str) -> list:
    """
    지문 단락이 임계값을 넘으면 의미 비트 단위로 분할.
    임계값 미만이거나 적절한 분단 지점을 못 찾으면 [text] 그대로 반환.

    분단 조건 (둘 중 하나 충족):
      - 자수 >= 150자
      - 자수 >= 100자 AND 문장수 >= 7 (단문 리듬 보존을 위한 하한)

    Returns:
        분할된 단락 리스트 (1개 또는 그 이상)
    """
    text = text.strip()
    if not text:
        return [text]

    char_len = len(text)
    sentences = _split_sentences(text)
    sent_count = len(sentences)

    triggered_by_length = char_len >= _ACTION_SPLIT_CHAR_THRESHOLD
    triggered_by_sentence = (char_len >= 100 and sent_count >= _ACTION_SPLIT_SENTENCE_THRESHOLD)

    if not (triggered_by_length or triggered_by_sentence):
        return [text]

    split_idx = _detect_paragraph_break_index(sentences)

    if split_idx < 0:
        if char_len < _ACTION_SPLIT_HARD_CHARS and sent_count < _ACTION_SPLIT_HARD_SENTENCES:
            return [text]
        split_idx = sent_count // 2

    part1 = ' '.join(sentences[:split_idx])
    part2 = ' '.join(sentences[split_idx:])
    result = [part1] + _split_action_paragraph(part2)
    return result


# ═══════════════════════════════════════════════════════════
# ★ Writer Engine v3.1.5 자산 이식 — PROP 메모·CHECK 태그 정제
# ═══════════════════════════════════════════════════════════

def _strip_prop_state_memos(text: str) -> str:
    """
    [소품 상태 / S#N 종료 시점] / GENRE_BOOSTER_CHECK / HELPER_CHARACTER_CHECK
    같은 INTERNAL 메모 블록을 본문에서 제거.

    AI가 비트 끝에 작성한 자가 검증 메모는 본문 노출 금지.
    """
    import re as _re_prop
    if not text:
        return text

    # 패턴 1: 코드블록 안 [소품 상태 ...] (```으로 감싼 형태)
    pattern_codeblock = _re_prop.compile(
        r'```[^\n]*\n\[소품\s*상태[^\]]*\][\s\S]*?```',
        _re_prop.MULTILINE
    )
    text = pattern_codeblock.sub('', text)

    # 패턴 2: 일반 텍스트 안의 [소품 상태] 블록
    pattern_inline = _re_prop.compile(
        r'\n*\[소품\s*상태[^\]]*\]\s*\n'
        r'(?:[\s]*[-•·][^\n]*\n?)+',
        _re_prop.MULTILINE
    )
    text = pattern_inline.sub('\n', text)

    # 패턴 3: INTERNAL / 작가 노트 / 소품 추적
    pattern_internal = _re_prop.compile(
        r'\n*\[?(?:INTERNAL|작가\s*노트|작가노트|소품\s*추적)[^\]]*\]?\s*\n'
        r'(?:\[소품\s*상태[^\]]*\]\s*\n)?'
        r'(?:[\s]*[-•·][^\n]*\n?)+',
        _re_prop.IGNORECASE | _re_prop.MULTILINE
    )
    text = pattern_internal.sub('\n', text)

    # 패턴 4: GENRE_BOOSTER_CHECK 태그
    pattern_booster = _re_prop.compile(
        r'\n*<GENRE_BOOSTER_CHECK>[\s\S]*?</GENRE_BOOSTER_CHECK>\n*',
        _re_prop.IGNORECASE
    )
    text = pattern_booster.sub('\n', text)

    # 패턴 5: HELPER_CHARACTER_CHECK 태그
    pattern_helper = _re_prop.compile(
        r'\n*<HELPER_CHARACTER_CHECK>[\s\S]*?</HELPER_CHARACTER_CHECK>\n*',
        _re_prop.IGNORECASE
    )
    text = pattern_helper.sub('\n', text)

    # 패턴 6: 닫기 태그 없이 떠도는 자가 검증 헤더
    pattern_check_header = _re_prop.compile(
        r'\n*\[★?\s*비트\s*종료[^\]]*GENRE_BOOSTER_CHECK[^\]]*\][\s\S]*?(?=\n\[|\nS#|\n$|\Z)',
        _re_prop.IGNORECASE
    )
    text = pattern_check_header.sub('\n', text)

    # 연속된 빈 줄 정리 (3개 이상 → 2개)
    text = _re_prop.sub(r'\n{3,}', '\n\n', text)

    return text.strip()


# ═══════════════════════════════════════════════════════════
# ★ Writer Engine v3.1.4 자산 이식 — INSERT 시스템
# 카톡·문자·이메일·유튜브·뉴스 등 화면 인서트 자동 감지
# ═══════════════════════════════════════════════════════════

_INSERT_LABEL_KEYWORDS = [
    '카톡', '메신저', '라인', '디스코드', '카카오톡',
    '문자', 'SMS', 'MMS',
    '이메일', '메일',
    '유튜브', 'YouTube', 'youtube', 'TV', '뉴스', '방송',
    'SNS', '인스타', '인스타그램', '페이스북', '트위터', 'X', '틱톡', 'DM',
    '검색', '구글', '네이버', '다음',
    '노트', '일기', '메모', '편지', '손글씨', '쪽지',
    '신문', '잡지', '기사', '헤드라인',
    '자막',
    '알림',
    '핸드폰', '핸드폰 화면', '폰', '폰 화면', '화면',
]


def _is_insert_label(text: str) -> bool:
    """[...] 형식 라벨인지 판단 — 형식 B 감지."""
    import re as _re_insert
    text = text.strip()
    if not (text.startswith('[') and ']' in text):
        return False
    label_match = _re_insert.match(r'^\[([^\]]+)\]', text)
    if not label_match:
        return False
    label_inner = label_match.group(1)
    return any(kw in label_inner for kw in _INSERT_LABEL_KEYWORDS)


def _parse_insert_blocks(text: str) -> list:
    """
    여러 줄 텍스트를 받아 INSERT 블록과 일반 텍스트로 분리.

    Returns:
        [{'type': 'action'|'insert_block'|'insert_label', 'data': ...}, ...]
    """
    import re as _re_insert
    if not text or not text.strip():
        return []

    lines = text.split('\n')
    items = []
    i = 0
    n = len(lines)
    accumulated_action = []

    def flush_action():
        if accumulated_action:
            joined = '\n'.join(accumulated_action).strip()
            if joined:
                items.append({'type': 'action', 'data': joined})
            accumulated_action.clear()

    while i < n:
        line = lines[i]
        line_stripped = line.strip()

        # 형식 A: INSERT — / INSERT - / INSERT:
        if _re_insert.match(r'^INSERT\s*[—\-:]', line_stripped, _re_insert.IGNORECASE):
            flush_action()
            header = line_stripped
            body_lines = []
            i += 1
            while i < n:
                bl = lines[i].strip()
                if _re_insert.match(r'^\[/INSERT\]?$', bl, _re_insert.IGNORECASE):
                    i += 1
                    break
                if not bl:
                    j = i + 1
                    while j < n and not lines[j].strip():
                        j += 1
                    if j >= n:
                        i = j
                        break
                    next_line = lines[j].strip()
                    if _re_insert.match(r'^\[/INSERT\]?$', next_line, _re_insert.IGNORECASE):
                        i = j + 1
                        break
                    if not _re_insert.match(r"^['\"\u2018\u2019\u201C\u201D]", next_line):
                        i = j
                        break
                    i += 1
                    continue
                body_lines.append(bl)
                i += 1
            items.append({
                'type': 'insert_block',
                'data': {'header': header, 'body': body_lines}
            })
            continue

        # 형식 B: [...] 라벨
        if _is_insert_label(line_stripped):
            flush_action()
            items.append({'type': 'insert_label', 'data': line_stripped})
            i += 1
            continue

        # 떠도는 [/INSERT] 단독 라인 무시
        if _re_insert.match(r'^\[/INSERT\]?$', line_stripped, _re_insert.IGNORECASE):
            i += 1
            continue

        accumulated_action.append(line)
        i += 1

    flush_action()
    return items


def _parse_insert_label(text: str) -> tuple:
    """
    형식 B 라벨 한 줄을 (label, body)로 분리.
    예: "[핸드폰 / 카톡] '아빠: 임대료 30프로 올린다.'"
        → ("[핸드폰 / 카톡]", "'아빠: 임대료 30프로 올린다.'")
    """
    import re as _re_insert
    m = _re_insert.match(r'^(\[[^\]]+\])\s*(.*)$', text.strip())
    if m:
        return m.group(1), m.group(2).strip()
    return text, ""


# ═══════════════════════════════════════════════════════════
# 씬 헤더 시간 표기 정규화 (Revise Engine 전용)
# ═══════════════════════════════════════════════════════════

def _normalize_scene_time_marker(content: str) -> str:
    """씬 헤더 시간 표기를 DAY/NIGHT 표준으로 정규화.

    매핑:
      아침 / 이른 아침 / 오전 / 낮 / 오후 / 정오 → DAY
      저녁 / 밤 / 야경 / 새벽 / 자정              → NIGHT
      요일만 표기                                  → DAY (보수적 기본값)
    """
    import re as _re

    day_keywords = ['이른 아침', '아침', '오전', '낮', '오후', '정오', '대낮']
    night_keywords = ['이른 저녁', '저녁', '밤', '야경', '새벽', '자정', '한밤', '심야']
    weekdays = ['월요일', '화요일', '수요일', '목요일', '금요일', '토요일', '일요일']

    def fix_header(match):
        scene_num = match.group(1)
        in_ex = match.group(2)
        location = match.group(3).strip()
        time_part = match.group(4).strip()
        for kw in ['야경', '새벽']:
            if location.endswith(' ' + kw):
                location = location[:-len(kw)].rstrip()
                time_part = kw + ' ' + time_part
        is_night = any(kw in time_part for kw in night_keywords)
        is_day = any(kw in time_part for kw in day_keywords)
        has_weekday = any(wd in time_part for wd in weekdays)
        if is_night:
            time_normalized = "NIGHT"
        elif is_day:
            time_normalized = "DAY"
        elif has_weekday:
            time_normalized = "DAY"
        else:
            return match.group(0)
        return f"S#{scene_num}. {in_ex}. {location} — {time_normalized}"

    pattern = _re.compile(
        r'S#(\d+(?:-\d+)?)\.\s*(INT|EXT)\.\s*([^\n—]+?)\s*—\s*([^\n]{1,40}?)(?=\n|$|\.\s+[가-힣A-Z])',
        _re.MULTILINE
    )
    return pattern.sub(fix_header, content)


# ═══════════════════════════════════════════════════════════
# 대사·지문 융합 분리 (A22 위반 자동 수정)
# ═══════════════════════════════════════════════════════════

def _split_dialog_action_fusion(content: str, char_names_pattern: str) -> tuple:
    """A22 위반 — 대사 라인에 지문이 따라붙은 경우 분리.

    예:
      유진\t\t자, 면수 버리지 마세요. 지난번에도 말씀드렸죠? 수업이 진행된다. 토마토가...
                                                              ↑ 종결문 후 지문 시작 → 분리
    """
    import re as _re
    fusion_count = 0
    lines = content.split('\n')
    new_lines = []

    action_indicators = [
        r'수업이\s+(?:진행|시작|끝)',
        r'(?:유진|진호|세웅|다은|강회장|지우|여름|최여름|조민준|최상진|집배원)이?\s+(?:웃|보|돈|섰|들|놓|떨|넘|건|받|먹|쳐|일어|앉|걷|뛰|달|들어|나|덮|연|닫|꺾)',
        r'(?:토마토|면|접시|냄비|국|된장|소스)가?\s+(?:끓|졸|담|풀)',
        r'(?:바람|음악|가로등|시계|핸드폰|폰|문|커튼|불|조명)이?\s+(?:불|울|흐|들|나|꺼|켜|보)',
        r'CUT\s+TO:', r'DISSOLVE\s+TO:', r'FADE\s+(?:IN|OUT)', r'\[인서트', r'INSERT\s*[—\-:]',
    ]
    action_pattern = _re.compile('|'.join(f'({p})' for p in action_indicators))
    char_line_pattern = _re.compile(rf'^({char_names_pattern})\t\t(.+)$')

    for line in lines:
        m = char_line_pattern.match(line)
        if m:
            speaker = m.group(1)
            dialog_text = m.group(2)
            split_match = _re.search(
                r'([\.\?\!…])\s+(?=' + action_pattern.pattern + ')',
                dialog_text
            )
            if split_match:
                cut_pos = split_match.end(1)
                pure_dialog = dialog_text[:cut_pos].rstrip()
                trailing_action = dialog_text[cut_pos:].strip()
                if pure_dialog and trailing_action:
                    new_lines.append(f"{speaker}\t\t{pure_dialog}")
                    new_lines.append("")
                    new_lines.append(trailing_action)
                    fusion_count += 1
                    continue
        new_lines.append(line)
    return '\n'.join(new_lines), fusion_count


# ═══════════════════════════════════════════════════════════
# 통짜 지문 자동 분단 (Writer Engine _split_action_paragraph 활용)
# ═══════════════════════════════════════════════════════════

def _split_clumping_action_lines(content: str) -> tuple:
    """REVISE 결과 본문 전체에 Writer Engine의 _split_action_paragraph 적용.

    문단 단위로 분리한 다음, 각 문단을 _split_action_paragraph로 검사.
    임계값 (150자/7문장 또는 240자/9문장 하드)을 넘으면 자동 분단.
    """
    import re as _re
    paragraphs = _re.split(r'\n\n+', content)
    new_paragraphs = []
    clump_count = 0

    for para in paragraphs:
        # 대사 라인이거나 씬 헤더면 통과
        if '\t\t' in para or _re.match(r'^S#\d+\.', para.strip()):
            new_paragraphs.append(para)
            continue
        # INSERT 블록도 통과 (자체 처리됨)
        if _re.match(r'^INSERT\s*[—\-:]', para.strip(), _re.IGNORECASE):
            new_paragraphs.append(para)
            continue
        if _re.match(r'^\[', para.strip()) and _is_insert_label(para.strip().split('\n')[0]):
            new_paragraphs.append(para)
            continue

        # Writer Engine 정품 분단 알고리즘 호출
        sub_paragraphs = _split_action_paragraph(para)
        if len(sub_paragraphs) > 1:
            new_paragraphs.extend(sub_paragraphs)
            clump_count += 1
        else:
            new_paragraphs.append(para)

    return '\n\n'.join(new_paragraphs), clump_count



def _validate_and_fix_revised_format(revise_result: dict) -> dict:
    """REVISE 결과의 revised_content 포맷을 검증하고 자동 보강.

    AI가 한 줄 통짜로 출력한 경우를 감지해 자동 줄바꿈 삽입.
    Writer Engine 표준에 맞춰 캐릭터명 + 탭2 + 대사 형식 강제.
    추가:
      - 씬 헤더 시간 표기 → DAY/NIGHT 정규화
      - A22 위반(대사·지문 융합) 자동 분리
      - A21 위반(통짜 지문) 자동 비트 분리
    """
    if not isinstance(revise_result, dict):
        return revise_result

    rr = revise_result.get("revision_result", {})
    if not isinstance(rr, dict):
        return revise_result

    revised_scenes = rr.get("revised_scenes", [])
    if not isinstance(revised_scenes, list):
        return revise_result

    import re as _re

    # 캐릭터명 패턴 (대표님 작품 + 일반)
    char_names_pattern = (
        '유진|진호|세웅|다은|강회장|김 여사|김여사|지우|여름|최여름|'
        '안경 아이|동생|아내|집배원|중개인|최상진|조민준|강유진|이진호|반세웅|'
        '편집자|엄마|아빠|손님|상인|남자|여자|수강생\\s*\\d?|아이|보호자'
    )

    fixed_count = 0
    time_normalized_count = 0
    fusion_split_count = 0
    clump_split_count = 0
    prop_strip_count = 0
    
    for scene in revised_scenes:
        if not isinstance(scene, dict):
            continue
        content = scene.get("revised_content", "")
        if not content or len(content) < 100:
            continue

        # ★ 1. PROP 메모·CHECK 태그 정제 (Writer Engine v3.1.5 / v3.2.0 자산)
        before_prop = content
        content = _strip_prop_state_memos(content)
        if content != before_prop:
            prop_strip_count += 1

        # ★ 2. 씬 헤더 시간 표기 정규화 (DAY/NIGHT)
        before = content
        content = _normalize_scene_time_marker(content)
        if content != before:
            time_normalized_count += 1
        if "scene_header" in scene and isinstance(scene["scene_header"], str):
            scene["scene_header"] = _normalize_scene_time_marker(scene["scene_header"])

        # ★ 3. A22 — 대사·지문 융합 분리
        content, fcount = _split_dialog_action_fusion(content, char_names_pattern)
        fusion_split_count += fcount

        # ★ 4. A21 — 통짜 지문 비트 분리 (Writer Engine v3.1.3 분단 알고리즘)
        content, ccount = _split_clumping_action_lines(content)
        clump_split_count += ccount

        # 줄바꿈 비율 체크
        line_count = content.count('\n')
        char_count = len(content)

        # 200자당 줄바꿈 1개 미만이면 통짜 출력 → 자동 보강
        if char_count > 500 and line_count < char_count / 200:
            # 1. 씬 헤더(S#XX) 다음 줄바꿈
            content = _re.sub(
                r'(S#\d+\.\s+(?:INT|EXT)\.\s+[^\n—]+—\s+[^\n]{2,30}?(?:아침|점심|저녁|밤|낮|오전|오후|새벽))\s+([가-힣A-Z\[])',
                r'\1\n\n\2',
                content
            )
            # 2. CUT TO / DISSOLVE TO / FADE 앞뒤 줄바꿈
            content = _re.sub(r'(?<!\n)\s+(CUT TO:|DISSOLVE TO:|FADE\s)', r'\n\n\1', content)
            content = _re.sub(r'(CUT TO:|DISSOLVE TO:|FADE\s\w+)\s+(?=[^\n])', r'\1\n', content)
            # 3. [인서트] 앞뒤 줄바꿈
            content = _re.sub(r'(?<!\n)\s+(\[인서트)', r'\n\n\1', content)
            # 4. 문장 끝 + 캐릭터명 + (괄호) → 줄바꿈 + 탭2
            content = _re.sub(
                rf'([\.\?\!…])\s+({char_names_pattern})\s+(\([^)]+\))',
                r'\1\n\n\2\t\t\3',
                content
            )
            # 5. 문장 끝 + 캐릭터명 + 공백 2+ → 줄바꿈 + 탭2
            content = _re.sub(
                rf'(?<!^)([\.\?\!…])\s+({char_names_pattern})\s{{2,}}',
                r'\1\n\n\2\t\t',
                content
            )
            # 6. 줄 시작 캐릭터명 + 공백 2+ → 캐릭터명\t\t
            content = _re.sub(
                rf'(^|\n)({char_names_pattern})\s{{2,}}',
                r'\1\2\t\t',
                content
            )
            # 7. 연속 빈 줄 정리
            content = _re.sub(r'\n{3,}', '\n\n', content)

            fixed_count += 1

        # ★ v3.3 — A29 자동 후처리 (시간 정밀 표기 → 모호한 시간어)
        # 세션 옵션이 ON일 때만 작동 (기본 ON)
        try:
            if st.session_state.get("auto_fix_a29_enabled", True):
                content_a29, a29_count = auto_fix_a29_violations(content)
                if a29_count > 0:
                    content = content_a29
                    if "_a29_auto_fixed_total" not in rr:
                        rr["_a29_auto_fixed_total"] = 0
                    rr["_a29_auto_fixed_total"] += a29_count
        except Exception:
            # 세션 외부에서 호출되면 통과
            pass

        # ★ v3.3.5 — A33 자동 후처리: 같은 씬 내 동일 대사 중복 제거
        try:
            import re as _re_a33
            _dialogue_pat = _re_a33.compile(
                r'^([가-힣a-zA-Z0-9\s]{1,15}?)\s*(?:\([^)]*\))?\s*\t+(?:\([^)]*\))?\s*(.+)$'
            )
            _content_lines = content.split('\n')
            _seen_dialogues = {}
            _new_lines = []
            _a33_removed_in_scene = 0
            for _ln in _content_lines:
                _m = _dialogue_pat.match(_ln.strip())
                if _m and "\t\t" in _ln:
                    _speaker = _m.group(1).strip()
                    _norm = _re_a33.sub(r'[\s.,!?…\-]', '', _m.group(2).strip())
                    _key = f"{_speaker}::{_norm}"
                    if _norm and len(_norm) >= 5 and _key in _seen_dialogues:
                        _a33_removed_in_scene += 1
                        continue
                    _seen_dialogues[_key] = True
                _new_lines.append(_ln)
            if _a33_removed_in_scene > 0:
                content = '\n'.join(_new_lines)
                if "_a33_auto_fixed_total" not in rr:
                    rr["_a33_auto_fixed_total"] = 0
                rr["_a33_auto_fixed_total"] += _a33_removed_in_scene
        except Exception:
            pass

        # ★ A21/A22/시간 정규화로 변경된 내용도 무조건 저장
        scene["revised_content"] = content

    if fixed_count > 0:
        rr["_format_auto_fixed"] = fixed_count
    if time_normalized_count > 0:
        rr["_time_marker_normalized"] = time_normalized_count
    if fusion_split_count > 0:
        rr["_dialog_action_fusion_split"] = fusion_split_count
    if clump_split_count > 0:
        rr["_clumping_action_split"] = clump_split_count
    if prop_strip_count > 0:
        rr["_prop_memo_stripped"] = prop_strip_count

    return revise_result


def run_revise(client):
    """Stage 2: 전체 배치 처리 (구버전 호환 — 사용하지 않는 것을 권장)."""
    prompt_text = build_revise_prompt(
        raw_text=st.session_state.raw_text,
        diagnose_result=st.session_state.diagnose_result,
        genre=st.session_state.genre,
        intensity=st.session_state.intensity,
        locked=st.session_state.locked,
        profession_input=st.session_state.profession_input,
        period_key=st.session_state.period_key,
        historical_type=st.session_state.historical_type,
        fact_based=st.session_state.fact_based,
    )
    raw = call_claude(client, prompt_text, model=MODEL_WRITE, max_tokens=32000)
    if not raw:
        return None
    return parse_json(raw)


def run_verify(client):
    """Stage 3: Sonnet으로 검증 보고서 생성."""
    prompt_text = build_verify_prompt(
        raw_text=st.session_state.raw_text,
        revise_result=st.session_state.revise_result,
        instruction=st.session_state.instruction,
        locked=st.session_state.locked,
        genre=st.session_state.genre,
    )
    raw = call_claude(client, prompt_text, model=MODEL_ANALYZE, max_tokens=32000)
    if not raw:
        return None
    return parse_json(raw)


# =================================================================
# [7] UI 컴포넌트
# =================================================================
def render_hero():
    st.markdown("""
    <div class="rev-hero">
        <div class="brand">B L U E &nbsp; J E A N S &nbsp; P I C T U R E S</div>
        <div class="title">REVISE ENGINE <span style="font-size:0.45em; vertical-align:middle; background:#FFCB05; color:#191970; padding:3px 10px; border-radius:12px; margin-left:10px; font-weight:700; letter-spacing:1px;">v3.3.5</span></div>
        <div class="tag">D I A G N O S E &nbsp; · &nbsp; R E V I S E &nbsp; · &nbsp; V E R I F Y</div>
    </div>
    """, unsafe_allow_html=True)


def render_stepbar():
    step = st.session_state.step
    steps = [
        (0, "입력"),
        (1, "진단"),
        (2, "집필"),
        (3, "검증"),
        (4, "완료"),
    ]
    html = '<div class="rev-stepbar">'
    for num, name in steps:
        cls = "rev-step"
        if num < step:
            cls += " done"
        elif num == step:
            cls += " active"
        html += f'<div class="{cls}"><span class="num">{num+1}</span>{name}</div>'
    html += '</div>'
    st.markdown(html, unsafe_allow_html=True)


# =================================================================
# [8] Step 화면들
# =================================================================
def show_step_0_input():
    """Step 0: 작업 모드 선택 → 모드별 입력."""

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # ① 작업 모드 선택 (카드 3개)
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    st.markdown(
        '<div style="text-align:center; margin:8px 0 16px;">'
        '<div style="font-size:1.1rem; color:#191970; font-weight:700;">무엇을 하시겠습니까?</div>'
        '<div style="font-size:0.85rem; color:#8E8E99; margin-top:4px;">'
        '먼저 작업 종류를 선택하세요. 선택한 모드에 필요한 입력만 표시됩니다.'
        '</div></div>',
        unsafe_allow_html=True
    )

    # 카드 스타일 정의
    st.markdown("""
    <style>
    .work-card-row { display: flex; gap: 12px; margin-bottom: 12px; }
    .stButton > button.work-card {
        height: auto !important;
        padding: 18px 14px !important;
        white-space: normal !important;
        line-height: 1.45 !important;
    }
    </style>
    """, unsafe_allow_html=True)

    cw1, cw2, cw3, cw4 = st.columns(4)

    def _card(col, mode_id: str, emoji: str, title: str, sub: str, desc: str):
        with col:
            is_selected = (st.session_state.work_mode == mode_id)
            label = f"{emoji}\n\n**{title}**\n\n{sub}\n\n{desc}"
            btn_type = "primary" if is_selected else "secondary"
            if st.button(
                label,
                key=f"card_{mode_id}",
                use_container_width=True,
                type=btn_type,
                help=desc,
            ):
                # ★★★ v3.0: 모드 전환 시 이전 모드의 잔존 데이터 완전 초기화 ★★★
                # (이전 이어쓰기 세션의 자동 감지 결과가 expansion 모드에 잔존하던 버그 수정)
                st.session_state.work_mode = mode_id
                st.session_state.section_detection = None
                st.session_state.protected_ranges = []
                st.session_state.revision_ranges = []
                st.session_state.cascade_analysis = None
                st.session_state.boundary_info = ""

                if mode_id == "full":
                    st.session_state.section_mode = False
                    st.session_state.target_added_scenes = 0
                elif mode_id == "expansion":
                    # 비트 보강 확장: section_mode 켜되 revision_ranges는 사용자 입력 후 채워짐
                    st.session_state.section_mode = True
                    if st.session_state.target_added_scenes <= 0:
                        st.session_state.target_added_scenes = 29  # 기본값
                else:
                    # continuation / partial
                    st.session_state.section_mode = True
                    st.session_state.target_added_scenes = 0
                st.rerun()

    _card(cw1, "full",
          "📝", "전체 각색",
          "시나리오 전체를 다시 씀",
          "기본 모드. 시나리오 한 편을 통째로 진단하고 수정합니다.")

    _card(cw2, "continuation",
          "✍️", "이어쓰기",
          "손본 부분 다음을 이어씀",
          "1막은 이미 손봤고, 그 톤으로 나머지를 이어 쓰고 싶을 때.")

    _card(cw3, "partial",
          "✂️", "부분 수정",
          "특정 구간만 다시 씀",
          "2막 엔딩만 약하다 같은 경우. 일부 씬만 핀포인트로 수정.")

    _card(cw4, "expansion",
          "🎯", "비트 보강 확장",
          "씬 추가로 분량 확장",
          "71→100씬처럼 약점 비트에 ADD 씬 자동 분배. 보호 구간 LOCKED.")

    # 모드 선택 안 했으면 안내 후 종료
    if not st.session_state.work_mode:
        st.info("👆 위에서 작업 모드를 선택하세요.")
        st.stop()

    # 선택된 모드 표시 + 변경 버튼
    mode_labels = {
        "full": "📝 전체 각색",
        "continuation": "✍️ 이어쓰기",
        "partial": "✂️ 부분 수정",
        "expansion": "🎯 비트 보강 확장",
    }
    cm1, cm2 = st.columns([5, 1])
    with cm1:
        st.markdown(
            f'<div style="background:#EAF3DE; padding:10px 14px; border-radius:6px; '
            f'border-left:3px solid #2EC484; margin:8px 0; font-size:0.9rem;">'
            f'<b>선택된 모드: {mode_labels.get(st.session_state.work_mode, "")}</b>'
            f'</div>',
            unsafe_allow_html=True
        )
    with cm2:
        if st.button("🔄 모드 변경", key="change_mode", use_container_width=True):
            st.session_state.work_mode = None
            st.rerun()

    st.markdown("---")

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # ② 모드별 입력 영역
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

    # 1. 시나리오 업로드 (모든 모드 공통)
    st.markdown('<div class="rev-card-title">1. 원본 시나리오 업로드 (DOCX 또는 PDF)</div>', unsafe_allow_html=True)
    st.markdown('<div class="rev-caption">Writer Engine 출력 DOCX, 외부 시나리오 PDF 모두 가능합니다.</div>', unsafe_allow_html=True)

    uploaded = st.file_uploader(
        "DOCX 또는 PDF 파일 선택",
        type=["docx", "pdf"],
        key="docx_uploader",
        label_visibility="collapsed"
    )

    if uploaded:
        text = extract_text_from_uploaded_file(uploaded)
        if text:
            st.session_state.raw_text = text
            st.session_state.raw_filename = uploaded.name
            # 제목은 파일명에서 추출 (확장자 제거)
            st.session_state.title = re.sub(r'\.(docx|pdf)$', '', uploaded.name, flags=re.IGNORECASE)
            st.success(f"✅ 업로드 완료: {uploaded.name}  ({len(text):,}자)")
            with st.expander("📄 추출된 본문 미리보기 (앞 1,000자)"):
                st.text(text[:1000] + ("..." if len(text) > 1000 else ""))
        else:
            st.error("❌ 본문 추출 실패")

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 모드별 추가 필수 입력 (시나리오 바로 아래)
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

    # 이어쓰기 모드: 손본본 업로드
    if st.session_state.work_mode == "continuation":
        st.markdown("---")
        st.markdown('<div class="rev-card-title">2. 손본 시나리오 업로드 ✍️</div>', unsafe_allow_html=True)
        st.markdown(
            '<div class="rev-caption">앞부분(예: 1막)을 직접 손보신 시나리오를 올리세요. '
            'AI가 자동으로 손본 부분을 감지해 보호하고, 그 다음을 같은 톤으로 이어 씁니다.</div>',
            unsafe_allow_html=True
        )

        ref_file = st.file_uploader(
            "손본 시나리오 (DOCX 또는 PDF)",
            type=["docx", "pdf"],
            key="continuation_uploader",
            help="예: v2_3 — 1막을 새로 쓴 PDF 또는 DOCX",
        )
        if ref_file:
            _text = extract_text_from_uploaded_file(ref_file)
            if _text:
                # 이어쓰기 모드는 diff_refined_text를 손본본 베이스로 사용
                st.session_state.diff_refined_text = _text
                st.session_state.diff_refined_filename = ref_file.name
                # 톤 학습용으로도 같이 등록
                st.session_state.tone_ref_text = _text
                st.session_state.tone_ref_filename = ref_file.name
                # 자동 감지 모드로 설정
                st.session_state.section_input_method = "auto"
                st.session_state.diff_use_main_as_before = True

                st.success(f"✅ 손본본 등록: {ref_file.name} ({len(_text):,}자)")

                # 이미 자동 감지가 됐다면 결과 표시
                if st.session_state.section_detection:
                    detection = st.session_state.section_detection.get("section_detection", {})
                    cp = detection.get("continuation_point", {})
                    if cp.get("detected") in (True, "true"):
                        prot_str = ", ".join(
                            f"{r.get('from','')}~{r.get('to','')}"
                            for r in st.session_state.protected_ranges
                        )
                        rev_str = ", ".join(
                            f"{r.get('from','')}~{r.get('to','')}"
                            for r in st.session_state.revision_ranges
                        )
                        st.info(
                            f"✓ 자동 감지 완료\n\n"
                            f"🔒 보호 (안 건드림): {prot_str}\n\n"
                            f"✏️ 재집필 (이어 씀): {rev_str}"
                        )
                else:
                    st.caption("→ 분석 시작 시 자동 감지가 진행됩니다.")
        elif st.session_state.diff_refined_filename:
            st.info(f"📎 등록됨: {st.session_state.diff_refined_filename} "
                    f"({len(st.session_state.diff_refined_text):,}자)")
            if st.button("🗑️ 손본본 제거", key="btn_clear_continuation"):
                st.session_state.diff_refined_text = ""
                st.session_state.diff_refined_filename = ""
                st.session_state.tone_ref_text = ""
                st.session_state.tone_ref_filename = ""
                st.session_state.section_detection = None
                st.session_state.protected_ranges = []
                st.session_state.revision_ranges = []
                st.rerun()

    # 비트 보강 확장 모드: LOCKED 범위 + 추가 씬 수
    elif st.session_state.work_mode == "expansion":
        st.markdown("---")
        st.markdown('<div class="rev-card-title">2. 비트 보강 확장 설정 🎯</div>', unsafe_allow_html=True)
        st.markdown(
            '<div class="rev-caption">'
            '약점 비트에 ADD 씬을 자동 분배하여 분량을 확장합니다. '
            '보호 구간(LOCKED)은 한 글자도 건드리지 않습니다.'
            '</div>',
            unsafe_allow_html=True
        )

        # 안내 박스
        st.markdown(
            '<div style="background:#FEF3C7; padding:12px 16px; border-radius:6px; '
            'border-left:3px solid #F59E0B; margin:12px 0; font-size:0.88rem;">'
            '<b>🎯 비트 보강 확장이란?</b><br>'
            '시나리오 전체를 15-Beat 구조에 매핑 → 약점 비트(deficit) 진단 → '
            '추가할 씬 수를 비트별로 자동 분배 → ADD 위치 자동 산출.<br>'
            '<span style="color:#666;">예시: 71씬 → 100씬 확장 시 +29씬을 약점 비트에 자동 배치</span>'
            '</div>',
            unsafe_allow_html=True
        )

        # ① 추가 씬 수 입력
        col_t1, col_t2 = st.columns([1, 1])
        with col_t1:
            current_target = st.session_state.get("target_added_scenes", 29)
            new_target = st.number_input(
                "🎯 추가할 씬 수",
                min_value=1,
                max_value=200,
                value=max(1, current_target),
                step=1,
                help="예: 71씬 → 100씬 = 29 입력",
                key="expansion_target_scenes",
            )
            if new_target != current_target:
                st.session_state.target_added_scenes = new_target
                st.rerun()

        with col_t2:
            # 현재 시나리오 씬 수 자동 계산
            if st.session_state.raw_text:
                current_scenes = _detect_scene_count(st.session_state.raw_text)
                if current_scenes > 0:
                    expected_total = current_scenes + new_target
                    st.metric(
                        "확장 후 총 씬 수",
                        f"{expected_total}씬",
                        delta=f"+{new_target}씬"
                    )

        # ② LOCKED 보호 구간 입력
        st.markdown("**🔒 보호 구간 (LOCKED) 설정**")
        st.caption("이 범위는 진단·집필에서 절대 변경하지 않습니다.")

        # 기존 protected_ranges 확인
        if st.session_state.protected_ranges:
            current_protected = st.session_state.protected_ranges[0]
            default_from = current_protected.get("from", "S#1")
            default_to = current_protected.get("to", "S#25")
        else:
            default_from = "S#1"
            default_to = "S#25"

        col_p1, col_p2 = st.columns([1, 1])
        with col_p1:
            prot_from = st.text_input(
                "보호 시작",
                value=default_from,
                key="expansion_prot_from",
                placeholder="S#1",
                help="예: S#1"
            )
        with col_p2:
            prot_to = st.text_input(
                "보호 끝",
                value=default_to,
                key="expansion_prot_to",
                placeholder="S#25",
                help="예: S#25"
            )

        # 자동으로 protected_ranges 갱신
        if prot_from and prot_to:
            new_protected = [{"from": prot_from, "to": prot_to,
                              "reason": "비트 보강 확장 모드 — 사용자 지정 보호 구간"}]
            if st.session_state.protected_ranges != new_protected:
                st.session_state.protected_ranges = new_protected

            # 작업 영역 자동 산출 (보호 끝 다음 ~ 마지막 씬)
            try:
                # S#숫자 형식에서 숫자만 추출
                import re as _re_temp
                m = _re_temp.search(r'\d+', prot_to)
                if m and st.session_state.raw_text:
                    prot_end_num = int(m.group())
                    total_scenes = _detect_scene_count(st.session_state.raw_text)
                    if total_scenes > prot_end_num:
                        rev_from = f"S#{prot_end_num + 1}"
                        rev_to = f"S#{total_scenes}"
                        new_revision = [{"from": rev_from, "to": rev_to,
                                         "reason": "비트 보강 확장 모드 — 자동 산출 작업 영역"}]
                        if st.session_state.revision_ranges != new_revision:
                            st.session_state.revision_ranges = new_revision

                        st.success(
                            f"✅ 자동 설정 완료\n\n"
                            f"🔒 보호 구간: {prot_from} ~ {prot_to}\n\n"
                            f"✏️ 작업 영역: {rev_from} ~ {rev_to} (ADD/REWRITE 가능)\n\n"
                            f"➕ 추가 목표: +{new_target}씬"
                        )
            except Exception as e:
                st.warning(f"보호 구간 형식을 확인해주세요: {e}")

        # ③ 진단 미리보기
        if st.session_state.raw_text and new_target > 0:
            current_scenes = _detect_scene_count(st.session_state.raw_text)
            if current_scenes > 0:
                bs = st.session_state.get("diagnose_batch_size", 12)
                expected_batches = max(1, (current_scenes + bs - 1) // bs)
                st.caption(
                    f"🔬 진단 시 진행: Phase 1(15-Beat 매핑, 단일 호출) → "
                    f"Phase 2(비트 인식 배치 진단, {expected_batches}배치). "
                    f"보호 구간은 진단 결과에서 자동 제외됩니다."
                )

    # 부분 수정 모드: 어디를 다시 쓸지 선택
    elif st.session_state.work_mode == "partial":
        st.markdown("---")
        st.markdown('<div class="rev-card-title">2. 어디를 다시 쓸까요? ✂️</div>', unsafe_allow_html=True)
        st.markdown(
            '<div class="rev-caption">수정할 부분을 선택하세요. 나머지는 자동으로 보호됩니다.</div>',
            unsafe_allow_html=True
        )

        # 막 단위 빠른 선택
        st.markdown("**빠른 선택 (작품 분량에 따라 자동 계산):**")
        scene_count = 0
        if st.session_state.raw_text:
            import re as _re
            scene_count = len(_re.findall(r'^\s*\*?\*?S#?\d+', st.session_state.raw_text, _re.MULTILINE))

        if scene_count > 0:
            # 막 분할: 1막 = 1~25%, 2막 = 25~75%, 3막 = 75~100%
            act1_end = max(1, int(scene_count * 0.25))
            act2_start = act1_end + 1
            act2_end = max(act2_start, int(scene_count * 0.75))
            act3_start = act2_end + 1

            cb1, cb2, cb3 = st.columns(3)
            with cb1:
                if st.button(f"1막 다시 쓰기\n(S#1~S#{act1_end})",
                             key="act1_btn", use_container_width=True):
                    st.session_state.revision_ranges = [{
                        "from": "S#1", "to": f"S#{act1_end}", "reason": "1막 재집필"
                    }]
                    st.session_state.section_input_method = "manual"
                    st.rerun()
            with cb2:
                if st.button(f"2막 다시 쓰기\n(S#{act2_start}~S#{act2_end})",
                             key="act2_btn", use_container_width=True):
                    st.session_state.revision_ranges = [{
                        "from": f"S#{act2_start}", "to": f"S#{act2_end}", "reason": "2막 재집필"
                    }]
                    st.session_state.section_input_method = "manual"
                    st.rerun()
            with cb3:
                if st.button(f"3막 다시 쓰기\n(S#{act3_start}~S#{scene_count})",
                             key="act3_btn", use_container_width=True):
                    st.session_state.revision_ranges = [{
                        "from": f"S#{act3_start}", "to": f"S#{scene_count}", "reason": "3막 재집필"
                    }]
                    st.session_state.section_input_method = "manual"
                    st.rerun()
        else:
            st.caption("⚠️ 1번 시나리오를 먼저 업로드하면 막 단위 빠른 선택이 활성화됩니다.")

        # 직접 지정
        st.markdown("**또는 직접 지정 (씬 번호 입력):**")
        st.caption("형식: `S#41-S#55` 또는 여러 구간은 콤마로 `S#41-S#55, S#67-S#67`")

        current_rev_str = ", ".join(
            f"{r.get('from','')}-{r.get('to','')}"
            for r in st.session_state.revision_ranges
        ) if st.session_state.revision_ranges else ""

        rev_input = st.text_input(
            "재집필 구간 직접 지정",
            value=current_rev_str,
            placeholder="예: S#41-S#55",
            key="manual_rev_ranges",
            label_visibility="collapsed",
        )

        if rev_input.strip() and rev_input != current_rev_str:
            try:
                import re as _re
                new_ranges = []
                for piece in rev_input.split(","):
                    piece = piece.strip()
                    m = _re.match(r'(S#\d+)\s*[-~]\s*(S#\d+)', piece)
                    if m:
                        new_ranges.append({
                            "from": m.group(1),
                            "to": m.group(2),
                            "reason": "직접 지정",
                        })
                if new_ranges:
                    st.session_state.revision_ranges = new_ranges
                    st.session_state.section_input_method = "manual"
                    st.success(f"✅ 재집필 구간 {len(new_ranges)}개 등록됨")
            except Exception as e:
                st.error(f"구간 파싱 실패: {e}")

        # 등록된 구간 표시
        if st.session_state.revision_ranges:
            rev_str = ", ".join(
                f"{r.get('from','')}~{r.get('to','')}"
                for r in st.session_state.revision_ranges
            )
            st.info(f"✏️ 재집필 구간: **{rev_str}** (나머지는 자동 보호)")

            if st.button("🗑️ 구간 초기화", key="btn_clear_partial"):
                st.session_state.revision_ranges = []
                st.session_state.protected_ranges = []
                st.rerun()

    st.markdown("---")

    # ── 제목 ──
    st.session_state.title = st.text_input(
        "프로젝트 제목",
        value=st.session_state.title,
        placeholder="예: 쿠킹클래스 러브 스토리",
    )

    # ── 장르 + 강도 ──
    c1, c2 = st.columns(2)
    with c1:
        st.session_state.genre = st.selectbox(
            "장르",
            options=list(GENRE_RULES.keys()),
            index=list(GENRE_RULES.keys()).index(st.session_state.genre)
                if st.session_state.genre in GENRE_RULES else 0,
        )
    with c2:
        intensity_options = list(INTENSITY_RULES.keys())
        intensity_labels = {
            "CONSERVATIVE": "CONSERVATIVE (원본 70%+ 보존)",
            "BALANCED":     "BALANCED (원본 50% 보존)",
            "AGGRESSIVE":   "AGGRESSIVE (원본 20~30% 유지)",
        }
        selected_label = st.selectbox(
            "수정 강도 (Intensity)",
            options=[intensity_labels[k] for k in intensity_options],
            index=intensity_options.index(st.session_state.intensity),
        )
        for k, lbl in intensity_labels.items():
            if lbl == selected_label:
                st.session_state.intensity = k
                break
        st.caption(f"→ {INTENSITY_RULES[st.session_state.intensity]['description']}")

    st.markdown("---")

    # ── 피드백 자료 (v2.1: 선택사항으로 격하) ──
    st.markdown('<div class="rev-card-title">2. 피드백 자료 <span style="font-weight:400; color:#8E8E99; font-size:0.85rem;">(선택사항)</span></div>',
                unsafe_allow_html=True)
    st.markdown('<div class="rev-caption">모니터 보고서·투자사 피드백·본인 메모·Rewrite Engine 진단 등 수정 방향에 도움되는 자료를 자유롭게 입력하세요. '
                '비워두어도 자동 진단이 작동합니다.</div>',
                unsafe_allow_html=True)

    # ★ v3.3.1 — 워크플로우 안내 박스 (어느 입력을 언제 쓰는지 명확히)
    st.markdown(
        '<div style="background:#F8F9FB; padding:12px 16px; border-radius:8px; '
        'border-left:4px solid #191970; margin:10px 0; font-size:0.88rem;">'
        '<b style="color:#191970;">📋 입력 자료 사용 가이드</b><br>'
        '<table style="width:100%; margin-top:6px; border-collapse:collapse;">'
        '<tr style="border-bottom:1px solid #E5E7EB;">'
        '<td style="padding:6px 0; width:80px;"><b style="color:#0EA5E9;">1차 수정</b></td>'
        '<td style="padding:6px 0;">Writer Engine 결과 → <b>Rewrite JSON</b> 사용 (아래 첫 expander)</td></tr>'
        '<tr><td style="padding:6px 0;"><b style="color:#EC4899;">N차 수정</b></td>'
        '<td style="padding:6px 0;">이전 각색본 + 검증 보고서 → <b>Revise 검증 JSON</b> 사용 (아래 두 번째 expander)</td></tr>'
        '</table>'
        '<span style="color:#666; font-size:0.82rem;">💡 둘 다 비워도 진단은 작동하며, 둘 다 입력하면 둘 모두 흡수됩니다.</span>'
        '</div>',
        unsafe_allow_html=True
    )

    # Rewrite Engine JSON 자동 변환 expander
    with st.expander("🔗 **[1차 수정용]** Rewrite Engine 진단·처방 JSON (자동 변환)"):
        st.caption(
            "**📌 언제 사용?** Writer Engine으로 집필한 시나리오의 첫 각색(1차) 시 사용합니다.\n\n"
            "Rewrite Engine에서 다운로드한 진단·처방 JSON을 업로드하거나 텍스트로 붙여넣으면, "
            "**CHRIS 분석 + SHIHO 처방 + MOON 전략(시장·장르·차별성·전체 방향)**이 자동으로 수정 지시문으로 변환됩니다.\n\n"
            "단, MOON이 직접 쓴 각색 원고(`moon_rewrite`)는 흡수되지 않고 Revise Engine이 "
            "자체 집필합니다 (AI ESCAPE A1~A32 + Genre Pack 적용).\n\n"
            "💡 **이 입력은 Round 1에서만 의미 있습니다.** "
            "Round 2 이상은 아래 [N차 수정용] 검증 보고서 흡수를 사용하세요."
        )

        rj_col1, rj_col2 = st.columns([1, 1])
        with rj_col1:
            uploaded_json = st.file_uploader(
                "JSON 파일 업로드",
                type=["json"],
                key="rewrite_json_uploader",
            )
        with rj_col2:
            pasted_json = st.text_area(
                "또는 JSON 직접 붙여넣기",
                height=120,
                key="rewrite_json_pasted",
                placeholder='{\n  "scores": {...},\n  "pros_cons": {...},\n  ...\n}',
            )

        if st.button("📥 Rewrite Engine JSON → 수정 지시문 변환", key="btn_convert_rewrite_json"):
            json_source = None
            if uploaded_json is not None:
                try:
                    json_source = uploaded_json.read().decode("utf-8")
                except Exception as e:
                    st.error(f"파일 읽기 실패: {e}")
            elif pasted_json.strip():
                json_source = pasted_json.strip()

            if json_source:
                try:
                    # v2.0 — 원문도 별도 저장 (메타 흡수용)
                    st.session_state.rewrite_json_text = json_source
                    st.session_state.rewrite_metadata = None  # 캐시 무효화

                    converted = parse_rewrite_engine_json(json_source)
                    if converted:
                        # 기존 지시문이 있으면 추가, 없으면 대체
                        if st.session_state.instruction.strip():
                            st.session_state.instruction = (
                                st.session_state.instruction.rstrip() +
                                "\n\n--- Rewrite Engine 진단·처방 결과 ---\n\n" +
                                converted
                            )
                        else:
                            st.session_state.instruction = converted
                        st.success(f"✅ 변환 완료! 지시문에 추가되었습니다 ({len(converted):,}자). v2.0 진단 시 preserve_notes·weak_zones도 자동 흡수됩니다.")
                        st.rerun()
                    else:
                        st.warning("⚠️ JSON에서 변환 가능한 진단·처방 내용을 찾지 못했습니다.")
                except Exception as e:
                    st.error(f"변환 실패: {e}")
            else:
                st.warning("⚠️ JSON 파일을 업로드하거나 텍스트로 붙여넣어 주세요.")

    # ★★★ v3.3 — Round N 검증 보고서 흡수 ★★★
    with st.expander("📊 **[N차 수정용]** Revise Engine 검증 보고서 흡수 (Round 2 이상)"):
        st.caption(
            "**📌 언제 사용?** Revise Engine으로 한 번 이상 각색한 시나리오를 "
            "추가로 다듬어 점수를 올릴 때 사용합니다 (Round 2, 3, ...).\n\n"
            "이전 라운드의 검증 보고서(.docx) 또는 검증 JSON을 업로드하면, "
            "미반영 처방 + AI 작법 위반 + 원본 유지 진단 씬이 자동으로 수정 지시문에 흡수됩니다.\n\n"
            "💡 **이 입력은 Round 2 이상에서만 사용합니다.** "
            "Writer Engine 결과의 첫 각색은 위 [1차 수정용] Rewrite JSON을 사용하세요.\n\n"
            "🎯 **목표:** Round 1의 6.x점 → Round 2의 7.5~8.0점 → Round 3의 8.0+ 점."
        )

        vr_col1, vr_col2 = st.columns([1, 1])
        with vr_col1:
            uploaded_verify_docx = st.file_uploader(
                "검증 보고서 (.docx) 업로드",
                type=["docx"],
                key="verify_report_docx_uploader",
            )
        with vr_col2:
            uploaded_verify_json = st.file_uploader(
                "또는 검증 JSON 업로드",
                type=["json"],
                key="verify_report_json_uploader",
            )

        col_round, col_a29 = st.columns([1, 1])
        with col_round:
            current_round = st.number_input(
                "현재 라운드 번호",
                min_value=1, max_value=10,
                value=st.session_state.get("round_n", 1),
                step=1,
                help="1=초회 작업, 2=Round 2 (검증 보고서 흡수 후 재집필)",
                key="round_n_input"
            )
            if current_round != st.session_state.get("round_n", 1):
                st.session_state.round_n = current_round
        with col_a29:
            auto_fix = st.checkbox(
                "✨ A29 자동 후처리 (한 박자 → 잠깐, 찰나 → 잠깐 등)",
                value=st.session_state.get("auto_fix_a29_enabled", True),
                key="auto_fix_a29_checkbox",
                help="집필 결과에서 A29 시간 정밀 표기를 자동 치환합니다."
            )
            st.session_state.auto_fix_a29_enabled = auto_fix

        if st.button("📥 검증 보고서 → 수정 지시문 변환", key="btn_convert_verify"):
            verify_meta = None

            # JSON 우선
            if uploaded_verify_json is not None:
                try:
                    import json as _json_temp
                    json_text = uploaded_verify_json.read().decode("utf-8")
                    json_data = _json_temp.loads(json_text)
                    # JSON 형식 검증 보고서를 파싱 보고서 형식으로 정규화
                    verify_meta = {
                        "round_n": json_data.get("round_n", 1),
                        "previous_score": json_data.get("overall_score"),
                        "previous_verdict": json_data.get("verdict", ""),
                        "not_reflected": [],
                        "partial_reflected": [],
                        "ai_escape_violations": json_data.get("ai_escape_check", {}).get("violations", []),
                        "untouched_scenes": json_data.get("normalized_for_next_round", {}).get("untouched_scenes", []),
                        "next_recommendations": json_data.get("next_round_recommendations", [])
                                                  or json_data.get("recommendations", []),
                        "instruction_text": "",
                    }
                    # instruction_text 재구성
                    parts = []
                    if verify_meta["previous_score"]:
                        parts.append(
                            f"[Round 직전 검증 결과]\n"
                            f"이전 점수: {verify_meta['previous_score']}/10 ({verify_meta['previous_verdict']})\n"
                            f"이번 라운드 목표: 8.0/10."
                        )
                    if verify_meta["next_recommendations"]:
                        parts.append("\n[재수정 권고]")
                        for i, r in enumerate(verify_meta["next_recommendations"], 1):
                            r_text = r if isinstance(r, str) else (
                                r.get("recommendation", "") or r.get("text", "") or str(r)
                            )
                            parts.append(f"{i}. {r_text[:300]}")
                    if verify_meta["untouched_scenes"]:
                        parts.append(f"\n[원본 유지 진단 씬 — 우선 검토 대상]\n{', '.join(verify_meta['untouched_scenes'])}")
                    if verify_meta["ai_escape_violations"]:
                        parts.append("\n[AI 작법 위반 — 재집필 시 회피]")
                        for v in verify_meta["ai_escape_violations"][:10]:
                            v_dict = v if isinstance(v, dict) else {}
                            parts.append(f"- {v_dict.get('pattern_id','?')} {v_dict.get('scene_id','')} : {v_dict.get('quote','')[:60]}")
                    verify_meta["instruction_text"] = "\n".join(parts)
                except Exception as e:
                    st.error(f"JSON 파싱 실패: {e}")

            elif uploaded_verify_docx is not None:
                try:
                    verify_meta = parse_verification_docx(uploaded_verify_docx)
                except Exception as e:
                    st.error(f"DOCX 파싱 실패: {e}")

            if verify_meta and not verify_meta.get("error"):
                st.session_state.verify_report_metadata = verify_meta
                st.session_state.verify_report_text = verify_meta.get("instruction_text", "")
                # round_n 자동 증가
                if verify_meta.get("round_n"):
                    st.session_state.round_n = verify_meta["round_n"] + 1
                else:
                    st.session_state.round_n = max(2, st.session_state.get("round_n", 1) + 1)

                # 지시문에 자동 추가
                vt = verify_meta.get("instruction_text", "")
                if vt:
                    if st.session_state.instruction.strip():
                        st.session_state.instruction = (
                            st.session_state.instruction.rstrip() +
                            "\n\n--- v3.3 검증 보고서 흡수 (Round " + str(st.session_state.round_n) + ") ---\n\n" +
                            vt
                        )
                    else:
                        st.session_state.instruction = vt

                # 통계 표시
                stats = []
                if verify_meta.get("not_reflected"):
                    stats.append(f"미반영 처방 {len(verify_meta['not_reflected'])}개")
                if verify_meta.get("partial_reflected"):
                    stats.append(f"부분 반영 {len(verify_meta['partial_reflected'])}개")
                if verify_meta.get("ai_escape_violations"):
                    stats.append(f"AI 작법 위반 {len(verify_meta['ai_escape_violations'])}회")
                if verify_meta.get("untouched_scenes"):
                    stats.append(f"원본 유지 씬 {len(verify_meta['untouched_scenes'])}개")
                if verify_meta.get("next_recommendations"):
                    stats.append(f"재수정 권고 {len(verify_meta['next_recommendations'])}개")

                st.success(
                    f"✅ Round {st.session_state.round_n} 흡수 완료!\n\n" +
                    "\n".join(f"  • {s}" for s in stats) +
                    f"\n\n이전 점수: **{verify_meta.get('previous_score', 'N/A')}/10** "
                    f"({verify_meta.get('previous_verdict', 'N/A')})\n"
                    f"이번 라운드 목표: **8.0/10** 도달"
                )
                st.rerun()
            elif verify_meta and verify_meta.get("error"):
                st.error(f"파싱 실패: {verify_meta['error']}")
            else:
                st.warning("⚠️ 검증 보고서 파일(DOCX 또는 JSON)을 업로드해주세요.")

    st.session_state.instruction = st.text_area(
        "지시문",
        value=st.session_state.instruction,
        height=240,
        placeholder="예:\n• 3막에서 주인공이 특정 남주를 선택하는 결말이 뻔함. 자기 발견 엔딩으로 재집필.\n"
                    "• 유진의 대사가 너무 설명적임. 행동으로 보여주도록.\n"
                    "• S#35 카페 씬의 시작을 '도착 과정' 대신 '이미 진행 중인 상황'으로 변경.\n\n"
                    "또는 Rewrite Engine의 진단·처방 내용을 그대로 붙여넣기도 가능합니다.",
        label_visibility="collapsed",
    )

    # ── LOCKED ──
    st.markdown('<div class="rev-card-title">3. LOCKED — 절대 건드리지 말 요소</div>', unsafe_allow_html=True)
    st.markdown('<div class="rev-caption">수정에서 제외할 요소를 자유롭게 적어주세요. LOCKED는 지시문보다 우선합니다.</div>',
                unsafe_allow_html=True)
    st.session_state.locked = st.text_area(
        "LOCKED",
        value=st.session_state.locked,
        height=120,
        placeholder="예:\n• 주인공의 직업(쇼핑 호스트) 유지\n"
                    "• 엔딩에서 주인공이 혼자 남는 구도 유지\n"
                    "• S#50의 반전은 건드리지 말 것\n"
                    "• 세웅 캐릭터의 대사는 그대로 유지",
        label_visibility="collapsed",
    )

    # ── 직업 전문성 (선택사항) ──
    st.markdown('<div class="rev-card-title">4. 주요 캐릭터 직업 <span style="font-weight:400; color:#8E8E99; font-size:0.85rem;">(선택사항)</span></div>',
                unsafe_allow_html=True)
    st.markdown('<div class="rev-caption">입력 시 해당 직업의 전문 용어·공간 디테일·금지 사항이 수정본 집필에 반영됩니다. '
                '비워두면 원본에서 자동 감지합니다.</div>', unsafe_allow_html=True)
    st.session_state.profession_input = st.text_area(
        "직업",
        value=st.session_state.profession_input,
        height=80,
        placeholder="예: 유진=쇼핑 호스트, 진호=변호사, 세웅=셰프\n"
                    "또는 단순히: 변호사, 셰프, 쇼핑 호스트",
        label_visibility="collapsed",
    )

    # ── 시대 / 역사영화 / 실화 (선택사항) ──
    st.markdown('<div class="rev-card-title">5. 시대 · 실화 정보 <span style="font-weight:400; color:#8E8E99; font-size:0.85rem;">(사극·시대극·실화영화 작업 시)</span></div>',
                unsafe_allow_html=True)
    st.markdown('<div class="rev-caption">현대 배경이면 그대로 두세요. 사극·시대극이면 시대를 선택하고, 실화 기반이면 체크하세요.</div>',
                unsafe_allow_html=True)

    period_keys = get_period_keys_for_ui()
    period_labels = get_period_labels_for_ui()

    cp1, cp2 = st.columns([1, 1])
    with cp1:
        # 시대 드롭다운
        try:
            current_idx = period_keys.index(st.session_state.period_key)
        except ValueError:
            current_idx = 0
        selected_period = st.selectbox(
            "시대 배경",
            options=period_keys,
            index=current_idx,
            format_func=lambda k: period_labels.get(k, k),
        )
        st.session_state.period_key = selected_period

    with cp2:
        # 역사영화 유형 (사극일 때만 활성화)
        is_historical = (st.session_state.period_key != "(현대)")
        ht_options = ["정통", "팩션", "퓨전"]
        try:
            ht_idx = ht_options.index(st.session_state.historical_type)
        except ValueError:
            ht_idx = 0
        selected_ht = st.radio(
            "역사영화 유형",
            options=ht_options,
            index=ht_idx,
            horizontal=True,
            disabled=not is_historical,
            help="정통: 사실 충실 / 팩션: 사실+허구 결합 / 퓨전: 현대 감각 적극 도입",
        )
        st.session_state.historical_type = selected_ht
        if not is_historical:
            st.caption("→ 시대 선택 시 활성화됩니다")

    # 실화 기반 체크박스
    st.session_state.fact_based = st.checkbox(
        "🎬 이 작품은 실화 또는 역사적 사건을 기반으로 합니다 (명예훼손·인격권 가이드 적용)",
        value=st.session_state.fact_based,
        help="실명 사용·특정 가능 디테일·실존 공인 악역화 등의 리스크를 자동 점검합니다",
    )

    # ── 6. 고급 옵션 (작가 톤 학습 + 장르 DNA) — 펼치기 ──
    advanced_label = "🎛️ 고급 옵션 — 작가 톤 학습 + 장르 DNA (선택사항)"
    if st.session_state.work_mode == "continuation":
        advanced_label = "🎛️ 고급 옵션 — 장르 DNA로 톤 강화 (선택사항)"

    with st.expander(advanced_label, expanded=False):
        st.markdown('<div class="rev-caption">필요시 펼쳐서 추가 자료를 업로드하세요. 모두 선택사항입니다.</div>',
                    unsafe_allow_html=True)

        tab_tone, tab_diff, tab_genre = st.tabs([
            "📐 톤 레퍼런스 (작가 톤 1편)",
            "🔬 Diff 학습 (Before vs After)",
            "🎬 장르 DNA (참고작 1~3편)"
        ])

        # 탭 1: 톤 레퍼런스 (작가 손본 1편)
        with tab_tone:
            st.caption("작가가 직접 손본 시나리오 1개를 업로드하면, 작가 고유의 톤 DNA를 자동 추출해 모든 새 각색에 강제 주입합니다.")
            ref_file = st.file_uploader(
                "작가가 손본 시나리오 (DOCX 또는 PDF)",
                type=["docx", "pdf"],
                key="tone_ref_uploader",
                help="예: v2_3 같은 작가가 직접 다듬은 버전"
            )
            if ref_file:
                _text = extract_text_from_uploaded_file(ref_file)
                if _text:
                    st.session_state.tone_ref_text = _text
                    st.session_state.tone_ref_filename = ref_file.name
                    st.success(f"✅ 톤 레퍼런스 로드: {ref_file.name} ({len(_text):,}자)")
                    if st.session_state.tone_dna:
                        st.info("✓ 톤 DNA 추출 완료. **Stage 1 진단 시작** 버튼을 누르면 이 톤이 시나리오 분석·집필에 자동 적용됩니다.")
                    else:
                        st.caption("→ Stage 1 진단 시작 시 톤 DNA가 자동으로 먼저 추출되고, 그 결과로 시나리오를 진단합니다.")
            elif st.session_state.tone_ref_filename:
                st.info(f"📎 등록됨: {st.session_state.tone_ref_filename} ({len(st.session_state.tone_ref_text):,}자)")
                if st.button("🗑️ 톤 레퍼런스 제거", key="btn_clear_tone_ref"):
                    st.session_state.tone_ref_text = ""
                    st.session_state.tone_ref_filename = ""
                    st.session_state.tone_dna = None
                    st.rerun()

        # 탭 2: Diff 학습 (Before + After)
        with tab_diff:
            st.caption("이전 버전(Before) vs 작가 손본 최신(After) 두 개를 비교해 편집 패턴(삭제·압축·통합 기준)을 자동 학습합니다.")

            use_main = st.checkbox(
                "✅ Before로 1번 원본 시나리오를 자동 사용 (권장)",
                value=st.session_state.diff_use_main_as_before,
                key="diff_use_main_chk",
                help="대부분의 경우 1번 원본을 Before로 쓰는 것이 자연스럽습니다."
            )
            st.session_state.diff_use_main_as_before = use_main

            ref2_file = st.file_uploader(
                "손본 최신 버전 (After · DOCX 또는 PDF)",
                type=["docx", "pdf"],
                key="diff_refined_uploader",
                help="예: v2_3"
            )
            if ref2_file:
                _text = extract_text_from_uploaded_file(ref2_file)
                if _text:
                    st.session_state.diff_refined_text = _text
                    st.session_state.diff_refined_filename = ref2_file.name
                    st.success(f"✅ After 등록: {ref2_file.name} ({len(_text):,}자)")
            elif st.session_state.diff_refined_filename:
                st.info(f"📎 After: {st.session_state.diff_refined_filename}")

            if not use_main:
                st.markdown("**고급 옵션 — Before 별도 업로드:**")
                orig_file = st.file_uploader(
                    "Before (별도 지정 · DOCX 또는 PDF)",
                    type=["docx", "pdf"],
                    key="diff_orig_uploader",
                )
                if orig_file:
                    _text = extract_text_from_uploaded_file(orig_file)
                    if _text:
                        st.session_state.diff_orig_text = _text
                        st.session_state.diff_orig_filename = orig_file.name
                        st.success(f"✅ Before: {orig_file.name}")

            if st.session_state.diff_refined_text:
                if st.session_state.diff_analysis:
                    st.info("✓ Diff 학습 완료. **Stage 1 진단 시작** 버튼을 누르면 학습된 편집 패턴이 시나리오 분석·집필에 자동 적용됩니다.")
                else:
                    st.caption("→ Stage 1 진단 시작 시 편집 패턴이 자동으로 먼저 학습되고, 그 결과로 시나리오를 진단합니다.")
                if st.button("🗑️ Diff 자료 제거", key="btn_clear_diff"):
                    st.session_state.diff_refined_text = ""
                    st.session_state.diff_refined_filename = ""
                    st.session_state.diff_orig_text = ""
                    st.session_state.diff_orig_filename = ""
                    st.session_state.diff_analysis = None
                    st.rerun()

        # 탭 3: 장르 DNA (참고작 1~3편)
        with tab_genre:
            # 가장 위에 '선택사항' 강조
            st.markdown(
                '<div style="background:#EEF3FB; padding:10px 14px; border-radius:6px; '
                'border-left:3px solid #4A6CF7; margin:8px 0; font-size:0.88rem;">'
                '<b>💡 장르 DNA는 100% 선택사항입니다.</b><br>'
                '없어도 엔진은 정상 작동합니다 (일반 장르 룰 적용).<br>'
                '있으면 같은 장르 명작들의 본질(코믹 폭발·정보 비대칭 등)을 정량적으로 강제할 수 있습니다.'
                '</div>',
                unsafe_allow_html=True
            )

            st.caption("📝 한국어·영문 시나리오 모두 가능. 영문 명작도 자유롭게 사용하세요 (출력은 항상 한국어).")

            # ─────────────────────────────────────────
            # STEP 1. 참고작 업로드 (메인 동선)
            # ─────────────────────────────────────────
            st.markdown('<div style="background:#FFF8DD; padding:10px 14px; border-radius:6px; '
                        'border-left:3px solid #FFCB05; margin:12px 0 8px; font-size:0.88rem;">'
                        '<b>① 처음 사용 — 참고작이 없으면?</b><br>'
                        '아래 사이트에서 무료로 다운로드 가능합니다:<br>'
                        '• <b>IMSDb</b> (imsdb.com) — Hollywood 명작 다수<br>'
                        '• <b>SimplyScripts</b> (simplyscripts.com) — 장르별 정리<br>'
                        '• <b>Script Slug</b> (scriptslug.com) — 최신작<br>'
                        '같은 장르 명작 1~3편 다운 → 아래 업로드 → 분석 시 장르 DNA 자동 추출'
                        '</div>', unsafe_allow_html=True)

            genre_files = st.file_uploader(
                "참고작 시나리오 (1~3편 · 같은 장르 · DOCX 또는 PDF)",
                type=["docx", "pdf"],
                key="genre_ref_uploader",
                accept_multiple_files=True,
                help="한국·영문 모두 가능. 외부 시나리오는 보통 PDF로 유통됩니다 (IMSDb·SimplyScripts 등). "
                     "예: 로코 → 「When Harry Met Sally」+「(500) Days of Summer」 / "
                     "느와르 → 「Drive」+「No Country for Old Men」 / 호러 → 「Hereditary」+「Get Out」"
            )
            if genre_files:
                texts = []
                names = []
                for gf in genre_files[:3]:
                    _text = extract_text_from_uploaded_file(gf)
                    if _text:
                        texts.append(_text)
                        names.append(gf.name)
                if texts:
                    st.session_state.genre_ref_texts = texts
                    st.session_state.genre_ref_filenames = names
                    total_chars = sum(len(t) for t in texts)
                    st.success(f"✅ 참고작 {len(texts)}편 로드 (총 {total_chars:,}자): {', '.join(names)}")
                    if st.session_state.genre_dna:
                        st.info("✓ 장르 DNA 추출 완료. **Stage 1 진단 시작** 버튼을 누르면 이 DNA가 시나리오 분석·집필에 자동 적용됩니다.")
                    else:
                        st.caption("→ Stage 1 진단 시작 시 장르 DNA가 자동으로 먼저 추출되고, 그 결과로 시나리오를 진단합니다.")
            elif st.session_state.genre_ref_filenames:
                st.info(f"📎 등록된 참고작: {', '.join(st.session_state.genre_ref_filenames)}")

            # ─────────────────────────────────────────
            # STEP 2. 추출된 DNA 다운로드 (라이브러리 보관)
            # ─────────────────────────────────────────
            if st.session_state.genre_dna:
                st.markdown('<div style="background:#EAF3DE; padding:8px 12px; border-radius:6px; '
                            'border-left:3px solid #2EC484; margin:14px 0 8px; font-size:0.88rem;">'
                            '<b>② 추출 완료 — 다음 작품에서 재사용하려면 JSON으로 보관하세요</b>'
                            '</div>', unsafe_allow_html=True)
                import json as _json
                dna_json = _json.dumps(st.session_state.genre_dna, ensure_ascii=False, indent=2)
                cdl1, cdl2 = st.columns([3, 1])
                with cdl1:
                    st.download_button(
                        "💾 장르 DNA JSON 다운로드 (라이브러리 보관)",
                        data=dna_json.encode("utf-8"),
                        file_name=f"genre_dna_{st.session_state.genre.replace(' ','_')}.json",
                        mime="application/json",
                        key="dl_genre_dna",
                        help="다음 프로젝트에서 재사용하려면 이 파일을 보관하세요",
                        use_container_width=True,
                    )
                with cdl2:
                    if st.button("🗑️ 제거", key="btn_clear_genre_dna", use_container_width=True):
                        st.session_state.genre_ref_texts = []
                        st.session_state.genre_ref_filenames = []
                        st.session_state.genre_dna = None
                        st.rerun()

            # ─────────────────────────────────────────
            # STEP 3. 라이브러리 — 이전에 보관한 JSON 재사용 (반복 사용자)
            # ─────────────────────────────────────────
            st.markdown("---")
            with st.expander("📚 장르 DNA 라이브러리 — 이전에 추출해둔 JSON 불러오기"):
                st.caption("같은 장르로 작업한 적이 있다면, 이전에 다운로드해둔 JSON을 업로드해서 "
                           "참고작 다시 안 올리고 바로 적용할 수 있습니다.")
                dna_json_file = st.file_uploader(
                    "장르 DNA JSON 업로드",
                    type=["json"],
                    key="genre_dna_json_uploader"
                )
                if dna_json_file:
                    try:
                        import json as _json
                        raw = dna_json_file.read().decode("utf-8")
                        loaded = _json.loads(raw)
                        st.session_state.genre_dna = loaded
                        summary = loaded.get("genre_dna", {}).get("summary", "장르 DNA 로드됨")
                        st.success(f"✅ 라이브러리에서 로드: {summary[:120]}")
                    except Exception as e:
                        st.error(f"JSON 로드 실패: {e}")

    # ── 실행 버튼 ──
    st.markdown("---")

    # 모드별 입력 검증
    has_scenario = bool(st.session_state.raw_text)
    work_mode = st.session_state.work_mode
    ready = False
    error_msg = None
    info_msg = None

    if not has_scenario:
        error_msg = "⚠️ 1번 원본 시나리오 업로드는 필수입니다."
    elif work_mode == "continuation":
        if not st.session_state.diff_refined_text:
            error_msg = "⚠️ 이어쓰기 모드는 2번 손본 시나리오 업로드가 필수입니다."
        else:
            ready = True
    elif work_mode == "partial":
        if not st.session_state.revision_ranges:
            error_msg = ("⚠️ 부분 수정 모드는 재집필 구간 선택이 필수입니다. "
                         "1막/2막/3막 버튼을 누르거나 직접 입력하세요.")
        else:
            ready = True
    elif work_mode == "expansion":
        # ★ v2.9 비트 보강 확장 모드
        if st.session_state.target_added_scenes <= 0:
            error_msg = "⚠️ 비트 보강 확장 모드는 추가할 씬 수가 1 이상이어야 합니다."
        elif not st.session_state.protected_ranges:
            error_msg = "⚠️ 비트 보강 확장 모드는 보호 구간(LOCKED) 설정이 필수입니다."
        else:
            ready = True
            current_scenes = _detect_scene_count(st.session_state.raw_text)
            target = st.session_state.target_added_scenes
            info_msg = (
                f"🎯 비트 보강 확장: {current_scenes}씬 → "
                f"{current_scenes + target}씬 (+{target}씬)"
            )
    else:
        # 전체 각색 모드
        ready = True
        aux_inputs = [
            bool(st.session_state.instruction.strip()),
            bool(st.session_state.rewrite_json_text.strip()),
            bool(st.session_state.genre_dna or st.session_state.genre_ref_texts),
            bool(st.session_state.tone_ref_text),
        ]
        if not any(aux_inputs):
            info_msg = "ℹ️ 보조 자료 없이 진행합니다. 시나리오 자체를 자동 진단해 약점을 찾습니다."

    if error_msg:
        st.warning(error_msg)
    elif info_msg:
        st.info(info_msg)
    else:
        # 등록된 자료 한눈에
        active = []
        if st.session_state.instruction.strip(): active.append("📝 피드백")
        if st.session_state.diff_refined_text: active.append("🔬 Diff")
        if st.session_state.rewrite_json_text.strip(): active.append("🔗 Rewrite JSON")
        if st.session_state.tone_ref_text: active.append("📐 톤 레퍼런스")
        if st.session_state.genre_dna or st.session_state.genre_ref_texts: active.append("🎬 장르 DNA")
        if st.session_state.section_mode: active.append("✂️ 구간 모드")
        st.caption(f"등록된 자료: {' · '.join(active)}")

    # 진단 버튼 직전 — 무슨 일이 일어날지 미리 안내
    if ready:
        pending = []
        if st.session_state.tone_ref_text and not st.session_state.tone_dna:
            pending.append("톤 DNA 추출")
        if st.session_state.diff_refined_text and not st.session_state.diff_analysis:
            pending.append("Diff 학습")
        if st.session_state.genre_ref_texts and not st.session_state.genre_dna:
            pending.append("장르 DNA 추출")
        if not st.session_state.distribution_diagnostic:
            pending.append("분포 진단")
        if st.session_state.rewrite_json_text.strip() and not st.session_state.rewrite_metadata:
            pending.append("Rewrite 메타 흡수")

        if pending:
            st.markdown(
                f'<div style="background:#EEF3FB; padding:10px 14px; border-radius:6px; '
                f'border-left:3px solid #4A6CF7; margin:8px 0; font-size:0.88rem;">'
                f'<b>🔬 Stage 1 진단 시작 시 자동 진행:</b><br>'
                f'1단계 사전 분석 ({", ".join(pending)}) → '
                f'2단계 시나리오 진단 (수정 플랜 생성)'
                f'</div>',
                unsafe_allow_html=True
            )
        else:
            st.markdown(
                '<div style="background:#EEF3FB; padding:10px 14px; border-radius:6px; '
                'border-left:3px solid #4A6CF7; margin:8px 0; font-size:0.88rem;">'
                '<b>🔬 Stage 1 진단 시작 시 자동 진행:</b><br>'
                '캐시된 사전 분석 결과를 활용해 바로 시나리오 진단(수정 플랜 생성)을 시작합니다.'
                '</div>',
                unsafe_allow_html=True
            )

    # ★ v2.7 자동 배치 분할 정보 박스 + 슬라이더
    if st.session_state.raw_text:
        _detected_scene_count = _detect_scene_count(st.session_state.raw_text)
        if _detected_scene_count > 0:
            _current_bs = st.session_state.get("diagnose_batch_size", 12)
            _expected_batches = max(1, (_detected_scene_count + _current_bs - 1) // _current_bs)
            _batch_color = "#10B981" if _expected_batches > 1 else "#6B7280"
            _batch_msg = (
                f"감지된 씬 수 <b>{_detected_scene_count}씬</b> → "
                f"진단 시 <b>{_expected_batches}배치</b>로 자동 분할 처리됩니다."
                if _expected_batches > 1 else
                f"감지된 씬 수 <b>{_detected_scene_count}씬</b> → 단일 배치 처리 (분할 불필요)."
            )
            st.markdown(
                f'<div style="background:#F0FDF4; padding:10px 14px; border-radius:6px; '
                f'border-left:3px solid {_batch_color}; margin:8px 0; font-size:0.88rem;">'
                f'<b>📦 v2.7 자동 배치 분할:</b> {_batch_msg}<br>'
                f'<span style="color:#666; font-size:0.82rem;">'
                f'토큰 잘림 방지를 위해 시나리오를 자동으로 N씬 단위로 쪼개 진단합니다. '
                f'사용자는 별도로 자를 필요 없이 버튼 한 번만 누르면 됩니다.</span>'
                f'</div>',
                unsafe_allow_html=True
            )
            with st.expander("⚙️ 자동 배치 분할 — 고급 옵션", expanded=False):
                new_bs = st.slider(
                    "DIAGNOSE 배치당 씬 개수",
                    min_value=8, max_value=15,
                    value=_current_bs,
                    step=1,
                    help=(
                        "한 배치에 포함할 씬 개수. 숫자가 작을수록 안전 마진은 커지지만 "
                        "API 호출 횟수가 늘어납니다. 기본 12 권장. "
                        "장르 룰이 무거운 경우(헐리우드 작법, 직업 Pack, 시대 Pack 동시) 10 권장."
                    ),
                    key="diagnose_batch_size_slider"
                )
                if new_bs != _current_bs:
                    st.session_state.diagnose_batch_size = new_bs
                    st.rerun()
                st.caption(
                    f"💡 토큰 안전 마진: 배치당 약 {new_bs}씬 → 입력 18~24K + 출력 6~8K = 안전 영역. "
                    f"71씬 시나리오 → {(_detected_scene_count + new_bs - 1) // new_bs}배치 (예상)."
                )

            # ★ v2.8 Beat-Aware Diagnose 입력 박스 (expansion 모드 외에서만 표시)
            if st.session_state.work_mode != "expansion":
                _current_target = st.session_state.get("target_added_scenes", 0)
                st.markdown(
                    f'<div style="background:#FEF3C7; padding:10px 14px; border-radius:6px; '
                    f'border-left:3px solid #F59E0B; margin:8px 0; font-size:0.88rem;">'
                    f'<b>🎯 v2.8 Beat-Aware Diagnose:</b> '
                    f'시나리오 확장(예: 71씬 → 100씬) 시 사용. '
                    f'15-Beat 구조 매핑 → 약점 비트 진단 → 자동 분배.<br>'
                    f'<span style="color:#666; font-size:0.82rem;">'
                    f'추가할 씬 수를 0보다 크게 설정하면 비트 인식 모드가 활성화됩니다. '
                    f'0이면 v2.7 일반 진단 모드.<br>'
                    f'※ 더 정밀한 보호 구간 LOCKED가 필요하면 작업 모드를 '
                    f'<b>🎯 비트 보강 확장</b>으로 선택하세요.</span>'
                    f'</div>',
                    unsafe_allow_html=True
                )
                new_target = st.number_input(
                    "🎯 추가할 씬 수 (Target Added Scenes)",
                    min_value=0,
                    max_value=200,
                    value=_current_target,
                    step=1,
                    help=(
                        "0: 일반 진단 (v2.7 모드)\n"
                        "1~200: Beat-Aware Diagnose 활성화 — 15-Beat 구조 진단 후 약점 비트에 ADD 자동 분배.\n\n"
                        "예시: 71씬 → 100씬 확장 시 29 입력."
                    ),
                    key="target_added_scenes_input"
                )
                if new_target != _current_target:
                    st.session_state.target_added_scenes = new_target
                    st.rerun()

                if new_target > 0:
                    _expected_total = _detected_scene_count + new_target
                    st.caption(
                        f"📈 확장 목표: {_detected_scene_count}씬 → **{_expected_total}씬** "
                        f"(+{new_target}씬 추가). "
                        f"진단 시 Phase 1(비트 매핑) → Phase 2(비트 인식 배치 진단) 순으로 진행됩니다."
                    )
                else:
                    st.caption(
                        f"💡 추가 씬 0 = 일반 진단 모드 (v2.7). 비트 매핑 단계 생략."
                    )

    # ★★★ v3.0: 진단 직전 라우팅 미리보기 — 어느 모드로 진단할지 명확히 ★★★
    work_mode_v30 = st.session_state.get("work_mode", "")
    target_added_v30 = st.session_state.get("target_added_scenes", 0)

    if work_mode_v30 == "expansion" and target_added_v30 > 0:
        st.markdown(
            f'<div style="background:#FCE7F3; padding:14px 18px; border-radius:8px; '
            f'border-left:4px solid #EC4899; margin:12px 0;">'
            f'<b style="font-size:1rem;">🎯 진단 모드: 비트 보강 확장 (v3.0)</b><br>'
            f'<span style="font-size:0.88rem; color:#374151;">'
            f'• 추가할 씬 수: <b>+{target_added_v30}씬</b><br>'
            f'• 보호 구간(LOCKED): <b>'
            + (", ".join(f"{r.get('from','')}~{r.get('to','')}" for r in st.session_state.get("protected_ranges", [])) or "(미설정)")
            + f'</b><br>'
            f'• 작업 영역: <b>'
            + (", ".join(f"{r.get('from','')}~{r.get('to','')}" for r in st.session_state.get("revision_ranges", [])) or "(자동 산출)")
            + f'</b><br>'
            f'<br>'
            f'진단 시 Phase 1(15-Beat 매핑) → Phase 2(비트 인식 진단) → Phase 5(LOCKED 차단)이 자동 진행됩니다.'
            f'</span></div>',
            unsafe_allow_html=True
        )
    elif work_mode_v30 == "continuation":
        st.markdown(
            '<div style="background:#DBEAFE; padding:12px 16px; border-radius:8px; '
            'border-left:4px solid #3B82F6; margin:12px 0;">'
            '<b>📝 진단 모드: 이어쓰기 (Fast Path 1)</b> — '
            '손본 시나리오의 톤을 학습해 재집필 구간을 같은 결로 다시 씁니다. '
            '<span style="color:#666;">(AI 진단 호출 없이 즉시 완료)</span>'
            '</div>',
            unsafe_allow_html=True
        )
    elif work_mode_v30 == "partial":
        st.markdown(
            '<div style="background:#DCFCE7; padding:12px 16px; border-radius:8px; '
            'border-left:4px solid #22C55E; margin:12px 0;">'
            '<b>✂️ 진단 모드: 부분 수정 (Fast Path 1)</b> — '
            '지정 구간만 재집필. 나머지는 보호.'
            '</div>',
            unsafe_allow_html=True
        )
    elif work_mode_v30 == "full":
        if target_added_v30 > 0:
            st.markdown(
                f'<div style="background:#FEF3C7; padding:12px 16px; border-radius:8px; '
                f'border-left:4px solid #F59E0B; margin:12px 0;">'
                f'<b>🎯 진단 모드: 전체 각색 + Beat-Aware (v2.8)</b> — '
                f'+{target_added_v30}씬 추가 모드'
                f'</div>',
                unsafe_allow_html=True
            )
        else:
            st.markdown(
                '<div style="background:#F3F4F6; padding:12px 16px; border-radius:8px; '
                'border-left:4px solid #6B7280; margin:12px 0;">'
                '<b>📝 진단 모드: 전체 각색 (v2.7)</b> — 자동 배치 분할 진단'
                '</div>',
                unsafe_allow_html=True
            )

    c1, c2 = st.columns([1, 1])
    with c1:
        # ★ v3.0: 버튼 라벨에 현재 모드 명시
        btn_label_map = {
            "expansion": "🎯 비트 보강 확장 진단 시작",
            "continuation": "📝 이어쓰기 진단 시작",
            "partial": "✂️ 부분 수정 진단 시작",
            "full": "📝 전체 각색 진단 시작",
        }
        btn_label = btn_label_map.get(work_mode_v30, "🔬 Stage 1: 진단 시작 (DIAGNOSE)")

        if st.button(btn_label,
                     disabled=not ready, use_container_width=True):
            client = get_client()
            if client:
                # 자동 배치 분할이 진행률을 직접 표시하므로 spinner는 짧게만
                with st.spinner("🔬 사전 분석 중... (Sonnet 4.6)"):
                    result = run_diagnose(client)
                    if result:
                        st.session_state.diagnose_result = result
                        st.session_state.step = 1
                        st.rerun()
                    else:
                        st.error("진단 실패. 다시 시도해주세요.")
    with c2:
        if st.button("🔄 초기화", use_container_width=True):
            reset_workflow()
            st.rerun()


def show_step_1_diagnose():
    """Step 1: DIAGNOSE 결과 확인 + REVISE 실행."""
    dr = st.session_state.diagnose_result.get("revision_plan", {})

    st.markdown('<div class="rev-card-title">🔬 Stage 1: 진단 결과 (Revision Plan)</div>',
                unsafe_allow_html=True)

    # v2.2 — 구간 모드 상태 표시 (가장 위에)
    if st.session_state.section_mode:
        prot_str = ", ".join(
            f"{r.get('from','')}~{r.get('to','')}"
            for r in st.session_state.protected_ranges
        ) or "(없음)"
        rev_str = ", ".join(
            f"{r.get('from','')}~{r.get('to','')}"
            for r in st.session_state.revision_ranges
        ) or "(없음)"

        # 보호 구간이 비어 있으면 강한 경고
        if not st.session_state.protected_ranges:
            st.error(
                "⚠️ **구간 모드 ON이지만 보호 구간이 등록되지 않았습니다.**\n\n"
                "- 자동 감지 모드라면: 6번 Diff 학습 탭에 손본본을 올렸는지 확인하세요.\n"
                "- 수동 모드라면: 7번 처리 모드에서 재집필 구간을 직접 입력하세요.\n\n"
                "→ 입력으로 돌아가 7번 처리 모드를 다시 점검해주세요. "
                "현재 진단은 전체 각색처럼 동작했습니다."
            )
        else:
            st.markdown(
                f'<div style="background:#EAF3DE; padding:10px 14px; border-radius:6px; '
                f'border-left:3px solid #2EC484; margin:8px 0 16px; font-size:0.9rem;">'
                f'<b>✂️ 구간 지정 모드 ON</b><br>'
                f'🔒 보호 구간 (안 건드림): <code>{prot_str}</code><br>'
                f'✏️ 재집필 구간 (다시 씀): <code>{rev_str}</code>'
                f'</div>',
                unsafe_allow_html=True
            )
    else:
        st.markdown(
            '<div style="background:#FFF8DD; padding:8px 14px; border-radius:6px; '
            'border-left:3px solid #FFCB05; margin:8px 0 16px; font-size:0.85rem;">'
            '<b>📝 전체 각색 모드</b> — 시나리오 전체가 진단·수정 대상입니다.'
            '</div>',
            unsafe_allow_html=True
        )

    # 요약
    summary = dr.get("summary", "")
    if summary:
        st.markdown('<div class="rev-card"><b style="color:#191970;">전체 수정 방향</b></div>',
                    unsafe_allow_html=True)
        st.write(summary)

    # 예상 씬 수
    cnt = dr.get("estimated_scene_count", "")
    conf = dr.get("confidence", "")
    if cnt or conf:
        st.markdown(f'<span class="rev-badge">수정 대상 씬: {cnt}</span>'
                    f'<span class="rev-badge y">진단 신뢰도: {conf}/10</span>',
                    unsafe_allow_html=True)

    st.markdown("---")

    # LOCKED 요약
    locked_summary = dr.get("locked_summary", "")
    if locked_summary:
        with st.expander("🔒 LOCKED로 인식된 요소"):
            st.write(locked_summary)

    # 충돌
    conflicts = dr.get("conflicts", [])
    if conflicts:
        st.markdown('<div class="rev-card-title" style="color:#E14444;">⚠️ 지시문 vs LOCKED 충돌</div>',
                    unsafe_allow_html=True)
        for c in conflicts:
            st.warning(f"• **지시:** {c.get('instruction_item','')}\n\n"
                       f"**충돌:** {c.get('locked_conflict','')}\n\n"
                       f"**해결:** {c.get('resolution','')}")

    # 수정 대상 씬 목록
    scenes = dr.get("target_scenes", [])
    if scenes:
        st.markdown(f'<div class="rev-card-title">📋 수정 대상 씬 ({len(scenes)}개)</div>',
                    unsafe_allow_html=True)
        for idx, sc in enumerate(scenes, 1):
            with st.expander(f"Scene {idx}: {sc.get('scene_id', '')}  —  {sc.get('scene_position', '')}"):
                st.markdown(f"**플롯상 기능:** {sc.get('original_function','')}")
                st.markdown(f"**보존 요소:** {sc.get('preservation_notes','')}")
                items = sc.get("revision_items", [])
                for i, it in enumerate(items, 1):
                    src = it.get("source", "")
                    badge = "📝" if src == "user_instruction" else "🔍"
                    st.markdown(f"{badge} **[{i}] {it.get('target_element','')}**")
                    st.markdown(f"- 이슈: {it.get('issue','')}")
                    st.markdown(f"- 방향: {it.get('proposed_direction','')}")

    # Out of Scope
    oos = dr.get("out_of_scope", [])
    if oos:
        with st.expander("⏩ 처리 불가 항목 (Out of Scope)"):
            for item in oos:
                st.markdown(f"- {item}")

    st.markdown("---")

    # 실행 버튼
    c1, c2 = st.columns([1, 1])
    with c1:
        if st.button("✍️ Stage 2: 집필 시작 (REVISE)", use_container_width=True):
            # 배치 분할 실행 (DIAGNOSE 결과 기반)
            batches = split_into_batches(
                st.session_state.diagnose_result,
                batch_size=st.session_state.batch_size
            )
            if not batches:
                st.error("수정 대상 씬이 없습니다. 진단 결과를 확인해주세요.")
            else:
                st.session_state.revise_batches = batches
                st.session_state.batch_results = {}
                st.session_state.current_batch = 0
                st.session_state.step = 2
                st.success(f"✅ {len(batches)}차 각색으로 나누어집니다. 차례대로 집필합니다.")
                st.rerun()
    with c2:
        if st.button("◀ 입력으로 돌아가기", use_container_width=True):
            st.session_state.step = 0
            st.session_state.diagnose_result = None
            st.rerun()


def _priority_badge(priority: str) -> str:
    """우선순위 배지 HTML."""
    colors = {
        "HIGH":   ("#E14444", "#FFFFFF"),
        "MEDIUM": ("#FFCB05", "#191970"),
        "LOW":    ("#8E8E99", "#FFFFFF"),
    }
    bg, fg = colors.get(priority, ("#8E8E99", "#FFFFFF"))
    return f'<span style="display:inline-block; padding:2px 8px; background:{bg}; color:{fg}; border-radius:4px; font-size:0.72rem; font-weight:800; font-family:Paperlogy,sans-serif;">{priority}</span>'


def _type_badge(t: str) -> str:
    """작업 종류 배지 HTML."""
    icons = {
        "REWRITE": "✏️",
        "ADD":     "➕",
        "DELETE":  "🗑️",
        "MERGE":   "🔗",
        "SPLIT":   "✂️",
    }
    icon = icons.get(t, "📝")
    return f'<span style="display:inline-block; padding:2px 8px; background:#EEEEF6; color:#191970; border-radius:4px; font-size:0.72rem; font-weight:700; margin-left:4px;">{icon} {t}</span>'


def show_step_2_revise():
    """Step 2: 배치 단위 순차 집필 UI."""
    batches = st.session_state.revise_batches or []
    batch_results = st.session_state.batch_results or {}

    if not batches:
        st.error("배치 정보가 없습니다. 진단 단계로 돌아가주세요.")
        if st.button("◀ 진단으로 돌아가기"):
            st.session_state.step = 1
            st.rerun()
        return

    total_batches = len(batches)
    completed_count = sum(1 for i in range(1, total_batches + 1) if i in batch_results)

    # 전체 씬 개수 / 완료 씬 개수
    total_scenes = sum(len(b["scenes"]) for b in batches)
    completed_scenes = sum(
        len(batch_results[i].get("revision_result", {}).get("revised_scenes", []))
        for i in range(1, total_batches + 1)
        if i in batch_results
    )

    # ── 헤더 ──
    st.markdown('<div class="rev-card-title">✍️ Stage 2: 배치 단위 집필 (REVISE)</div>',
                unsafe_allow_html=True)

    # 진행률 바
    progress = completed_count / total_batches if total_batches else 0
    st.progress(progress, text=f"배치 {completed_count} / {total_batches} 완료  ({completed_scenes} / {total_scenes} 씬)")

    # ── 💾 진행 상황 자동 백업 / 복구 패널 ──
    backup_col1, backup_col2 = st.columns([1, 1])

    with backup_col1:
        # 백업 다운로드 — 완료된 배치가 1개 이상 있을 때만
        if batch_results:
            import json as _json
            from datetime import datetime as _dt
            backup_data = {
                "version": "v2.2",
                "saved_at": _dt.now().isoformat(),
                "title": st.session_state.title,
                "genre": st.session_state.genre,
                "total_batches": total_batches,
                "completed_batches": completed_count,
                "batch_results": batch_results,
                "diagnose_result": st.session_state.diagnose_result,
                # 구간 모드 정보 보존
                "section_mode": st.session_state.section_mode,
                "work_mode": st.session_state.work_mode,
                "protected_ranges": st.session_state.protected_ranges,
                "revision_ranges": st.session_state.revision_ranges,
            }
            backup_json = _json.dumps(backup_data, ensure_ascii=False, indent=2)
            # ★ v3.3.2: 백업 파일명에 라운드 번호 + 안전한 제목 처리
            import re as _re_backup
            _safe_title_backup = _re_backup.sub(r'[/*?:"<>|]', '_',
                                                  (st.session_state.title or "제목없음").strip())
            _round_n_backup = st.session_state.get("round_n", 1)
            backup_filename = (
                f"{_safe_title_backup}_백업_R{_round_n_backup}_"
                f"{completed_count}of{total_batches}_"
                f"{_dt.now().strftime('%Y%m%d_%H%M')}_Blue.json"
            )

            st.download_button(
                f"💾 진행 상황 백업 ({completed_count}/{total_batches} 각색)",
                data=backup_json.encode("utf-8"),
                file_name=backup_filename,
                mime="application/json",
                key="backup_progress",
                help="에러 대비. 다음 각색 시작 전에 받아두세요. 파일명에 라운드·진행도·시·분 표시.",
                use_container_width=True,
            )

    with backup_col2:
        # 백업 복구 업로드
        with st.popover("📂 백업에서 복구", use_container_width=True):
            st.caption("이전 작업 백업 JSON을 올리면 그 지점부터 이어서 작업할 수 있습니다.")
            restore_file = st.file_uploader(
                "백업 JSON 업로드",
                type=["json"],
                key="restore_uploader",
            )
            if restore_file:
                try:
                    import json as _json
                    raw = restore_file.read().decode("utf-8")
                    loaded = _json.loads(raw)

                    if st.button("✅ 이 백업으로 복구", key="confirm_restore", type="primary"):
                        # 핵심 상태 복구
                        st.session_state.batch_results = {
                            int(k): v for k, v in loaded.get("batch_results", {}).items()
                        }
                        if loaded.get("diagnose_result"):
                            st.session_state.diagnose_result = loaded["diagnose_result"]
                        if loaded.get("section_mode") is not None:
                            st.session_state.section_mode = loaded["section_mode"]
                        if loaded.get("work_mode"):
                            st.session_state.work_mode = loaded["work_mode"]
                        if loaded.get("protected_ranges"):
                            st.session_state.protected_ranges = loaded["protected_ranges"]
                        if loaded.get("revision_ranges"):
                            st.session_state.revision_ranges = loaded["revision_ranges"]

                        restored_count = len(st.session_state.batch_results)
                        st.success(f"✅ 복구 완료: {restored_count}차 각색 결과 복원됨")
                        st.rerun()
                    else:
                        st.info(
                            f"📅 백업 시각: {loaded.get('saved_at', '?')[:16]}\n\n"
                            f"📌 작품: {loaded.get('title', '?')}\n\n"
                            f"📊 진행: {loaded.get('completed_batches', 0)} / {loaded.get('total_batches', 0)} 배치"
                        )
                except Exception as e:
                    st.error(f"백업 파일 로드 실패: {e}")

    # 각색 묶음 안내
    plan = st.session_state.diagnose_result.get("revision_plan", {})
    strategy = plan.get("batch_strategy", "")
    if strategy:
        st.info(f"📋 **각색 진행 전략:** {strategy}")

    st.markdown("---")

    # ── 각색 카드 목록 ──
    for batch in batches:
        bidx = batch["batch_index"]
        scenes = batch["scenes"]
        is_done = bidx in batch_results
        is_current = (bidx == completed_count + 1) and not is_done

        # 카드 헤더
        if is_done:
            status_icon = "✅"
            status_color = "#2EC484"
            status_text = "완료"
        elif is_current:
            status_icon = "▶️"
            status_color = "#FFCB05"
            status_text = "다음 차례"
        else:
            status_icon = "⏸️"
            status_color = "#8E8E99"
            status_text = "대기 중"

        # ★ v3.1 한국어 통계 (priority_summary, type_summary 한글화)
        # priority_summary는 "HIGH 6개" 식 → 그대로 사용
        # type_summary는 "REWRITE 6개" 식 → "수정 6 / 추가 0" 형태로 변환
        type_summary_kr = batch.get("type_summary", "")
        # type_summary 파싱 (예: "REWRITE 6개" 또는 "REWRITE 4 · ADD 2")
        import re as _re_v31
        rewrite_n = 0
        add_n = 0
        delete_n = 0
        merge_n = 0
        split_n = 0
        for sc in scenes:
            t = sc.get("type", "REWRITE")
            if t == "ADD": add_n += 1
            elif t == "DELETE": delete_n += 1
            elif t == "MERGE": merge_n += 1
            elif t == "SPLIT": split_n += 1
            else: rewrite_n += 1
        type_parts_kr = []
        if rewrite_n > 0: type_parts_kr.append(f"수정 {rewrite_n}")
        if add_n > 0: type_parts_kr.append(f"✨추가 {add_n}")
        if delete_n > 0: type_parts_kr.append(f"삭제 {delete_n}")
        if merge_n > 0: type_parts_kr.append(f"합치기 {merge_n}")
        if split_n > 0: type_parts_kr.append(f"쪼개기 {split_n}")
        type_summary_kr_final = " · ".join(type_parts_kr) if type_parts_kr else f"씬 {len(scenes)}개"

        with st.container():
            st.markdown(
                f'<div style="background:#FFFFFF; border:2px solid {status_color}; '
                f'border-radius:10px; padding:16px; margin-bottom:12px;">'
                f'<div style="display:flex; justify-content:space-between; align-items:center; margin-bottom:8px;">'
                f'<div style="font-family:Paperlogy,sans-serif; font-weight:800; font-size:1.05rem; color:#191970;">'
                f'{status_icon} 각색 {bidx}차 ({bidx} / {total_batches})'
                f'</div>'
                f'<div style="color:{status_color}; font-weight:700; font-size:0.85rem;">{status_text}</div>'
                f'</div>'
                f'<div style="color:#8E8E99; font-size:0.85rem; margin-bottom:8px;">'
                f'우선순위: {batch["priority_summary"]}  ·  {type_summary_kr_final}  ·  총 {len(scenes)}씬'
                f'</div>'
                f'</div>',
                unsafe_allow_html=True
            )

            # 씬 목록 (작은 expander)
            with st.expander(f"📝 수정 대상 씬 ({len(scenes)}개) — 클릭하면 무엇을 고치는지 보입니다", expanded=is_current and not is_done):
                for sc in scenes:
                    sid = sc.get("scene_id", "")
                    pri = sc.get("priority", "MEDIUM")
                    typ = sc.get("type", "REWRITE")
                    pos = sc.get("scene_position", "")
                    func = sc.get("original_function", "")

                    # ★ v3.1: 핵심 수정 방향 한 줄 추출
                    # ★ v3.2.1 핫픽스: None 안전 처리
                    items = sc.get("revision_items", []) or []
                    headline = ""
                    if items:
                        first_item = items[0] if items[0] else {}
                        proposed = (first_item.get("proposed_direction") or
                                    first_item.get("issue") or "")
                        # 명시적 str 변환 (혹시 dict나 다른 타입이 들어와도 안전)
                        proposed = str(proposed) if proposed else ""
                        if len(proposed) > 60:
                            headline = proposed[:60] + "..."
                        else:
                            headline = proposed

                    if not headline:
                        # preservation_notes 폴백 (None 안전)
                        pn = sc.get("preservation_notes") or ""
                        pn = str(pn) if pn else ""
                        if pn:
                            headline = pn[:60] + ("..." if len(pn) > 60 else "")

                    # ADD/REWRITE 한국어 라벨
                    type_kr = {
                        "ADD": "✨추가",
                        "REWRITE": "✏️수정",
                        "DELETE": "🗑️삭제",
                        "MERGE": "🔗합치기",
                        "SPLIT": "✂️쪼개기",
                    }.get(typ, typ)

                    # 메인 라인
                    main_line = (
                        f'{_priority_badge(pri)} {_type_badge(typ)} '
                        f'<b>{sid}</b>'
                    )
                    if pos:
                        main_line += f'  <span style="color:#8E8E99; font-size:0.85rem;">— {pos}</span>'
                    st.markdown(main_line, unsafe_allow_html=True)

                    # ★ v3.1: 핵심 수정 방향 한 줄 (가장 중요한 정보)
                    if headline:
                        st.markdown(
                            f'<div style="margin:4px 0 8px 8px; padding:6px 10px; '
                            f'background:#FFF8E1; border-left:3px solid #FFCB05; '
                            f'border-radius:4px; font-size:0.9rem;">'
                            f'<b style="color:#191970;">→ 핵심 방향:</b> {headline}'
                            f'</div>',
                            unsafe_allow_html=True
                        )

                    if func:
                        st.caption(f"플롯상 기능: {func}")

                    # 추가 수정 항목 (2번째 이후)
                    if len(items) > 1:
                        with st.expander(f"📌 세부 수정 항목 {len(items)}개 보기", expanded=False):
                            for it in items:
                                st.markdown(f"  • *{it.get('issue','')}* → {it.get('proposed_direction','')}")
                    st.markdown("")

            # 각색 액션 버튼
            bc1, bc2, bc3 = st.columns([2, 1, 1])
            with bc1:
                if is_done:
                    # 완료된 각색 — 결과 미리보기 + 재집필
                    btn_label = f"🔄 {bidx}차 각색 다시 하기"
                    if st.button(btn_label, key=f"rewrite_batch_{bidx}", use_container_width=True):
                        client = get_client()
                        if client:
                            with st.spinner(f"✍️ {bidx}차 각색 다시 진행 중... (Opus 4.6)"):
                                result = run_revise_batch(client, bidx, scenes, total_batches)
                                if result:
                                    st.session_state.batch_results[bidx] = result
                                    st.success(f"✅ {bidx}차 각색 완료")
                                    st.rerun()
                                else:
                                    st.error(f"{bidx}차 각색 실패")
                elif is_current:
                    # 다음 차례 — 집필 시작
                    btn_label = f"▶️ {bidx}차 각색 시작 ({len(scenes)}씬)"
                    if st.button(btn_label, key=f"run_batch_{bidx}", type="primary", use_container_width=True):
                        client = get_client()
                        if client:
                            with st.spinner(f"✍️ {bidx}차 각색 진행 중... ({len(scenes)}씬, Opus 4.6, 1~3분 소요)"):
                                result = run_revise_batch(client, bidx, scenes, total_batches)
                                if result:
                                    st.session_state.batch_results[bidx] = result
                                    st.success(f"✅ {bidx}차 각색 완료")
                                    # 다음 묶음 들어가기 전 백업 권고
                                    if bidx < total_batches:
                                        st.info(
                                            f"💡 **다음 각색 진행 전, 위쪽 [💾 진행 상황 백업] 버튼을 눌러 "
                                            f"현재까지 결과({bidx}/{total_batches}차)를 다운로드해두세요. "
                                            f"에러 발생 시 백업에서 복구할 수 있습니다.**"
                                        )
                                    st.rerun()
                                else:
                                    st.error(
                                        f"❌ {bidx}차 각색 실패. "
                                        f"위쪽 [💾 진행 상황 백업]으로 지금까지 결과를 먼저 저장하세요. "
                                        f"그 다음 다시 시도해주세요."
                                    )
                else:
                    # 대기 중 (이전 묶음 미완료)
                    st.button(f"⏸️ {bidx}차 각색 (대기 중)", disabled=True, use_container_width=True,
                              key=f"wait_batch_{bidx}")

            with bc2:
                if is_done:
                    # 결과 미리보기 토글
                    if st.button("👁️ 결과 보기", key=f"preview_batch_{bidx}", use_container_width=True):
                        st.session_state[f"show_preview_{bidx}"] = not st.session_state.get(f"show_preview_{bidx}", False)
                        st.rerun()

            with bc3:
                if is_done:
                    if st.button("🗑️ 결과 삭제", key=f"delete_batch_{bidx}", use_container_width=True):
                        del st.session_state.batch_results[bidx]
                        # 이후 각색 결과도 삭제 (순차 의존성)
                        for later_idx in range(bidx + 1, total_batches + 1):
                            st.session_state.batch_results.pop(later_idx, None)
                        st.rerun()

            # 결과 미리보기 영역
            if is_done and st.session_state.get(f"show_preview_{bidx}", False):
                br = batch_results[bidx]
                rr = br.get("revision_result", {})
                with st.container():
                    st.markdown(
                        f'<div style="background:#F7F7F5; border-left:4px solid #2EC484; '
                        f'padding:12px 16px; border-radius:0 8px 8px 0; margin-top:8px; margin-bottom:8px;">'
                        f'<b style="color:#191970;">{bidx}차 각색 결과 요약</b><br/>'
                        f'<span style="color:#1A1A2E;">{rr.get("summary","(요약 없음)")}</span>'
                        f'</div>',
                        unsafe_allow_html=True
                    )

                    revised_scenes = rr.get("revised_scenes", [])
                    for sc in revised_scenes:
                        header = sc.get("scene_header", "")
                        pri = sc.get("priority", "MEDIUM")
                        typ = sc.get("type", "REWRITE")
                        with st.expander(f"{header}  [{pri} · {typ}]"):
                            col_o, col_r = st.columns([1, 1])
                            with col_o:
                                st.markdown("**📄 원본 발췌**")
                                excerpt = sc.get("original_excerpt", "")
                                st.text(excerpt if excerpt else "(원본 없음 — 신규 추가/삭제)")
                            with col_r:
                                st.markdown("**✏️ 수정본**")
                                content = sc.get("revised_content", "")
                                st.text(content if content else "(삭제됨)")

                            notes = sc.get("revision_notes", {})
                            if notes:
                                st.markdown("---")
                                st.markdown(f"**변경:** {notes.get('what_changed','')}")
                                st.markdown(f"**보존:** {notes.get('what_preserved','')}")

    st.markdown("---")

    # ── 모든 각색 완료 시 다음 단계 ──
    all_done = (completed_count == total_batches)

    if all_done:
        st.success(f"🎉 모든 각색 완료! 총 {completed_scenes}개 씬이 수정되었습니다.")

    c1, c2, c3 = st.columns([2, 1, 1])
    with c1:
        if st.button("✅ Stage 3: 검증 시작 (VERIFY)",
                     disabled=not all_done, use_container_width=True, type="primary"):
            # 모든 각색 결과 통합
            all_results = [batch_results[i] for i in range(1, total_batches + 1) if i in batch_results]
            merged = merge_batch_results(all_results)
            st.session_state.revise_result = merged

            client = get_client()
            if client:
                with st.spinner("✅ 통합 수정본 검증 중... (Sonnet 4.6)"):
                    result = run_verify(client)
                    if result:
                        st.session_state.verify_result = result
                        st.session_state.step = 3
                        st.rerun()
                    else:
                        st.error("검증 실패. 다시 시도해주세요.")
    with c2:
        # 모두 한 번에 (남은 배치 자동 순차 실행)
        remaining = total_batches - completed_count
        if remaining > 0:
            if st.button(f"⏩ 남은 {remaining}개 자동 진행", use_container_width=True):
                client = get_client()
                if client:
                    progress_bar = st.progress(0, text="배치 자동 실행 중...")
                    for i, batch in enumerate(batches):
                        bidx = batch["batch_index"]
                        if bidx in batch_results:
                            continue
                        progress_bar.progress(
                            (completed_count + (i - completed_count + 1)) / total_batches,
                            text=f"{bidx}차 각색 진행 중... ({len(batch['scenes'])}개 씬)"
                        )
                        result = run_revise_batch(client, bidx, batch["scenes"], total_batches)
                        if result:
                            st.session_state.batch_results[bidx] = result
                        else:
                            st.error(f"{bidx}차 각색에서 실패. 중단합니다.")
                            break
                    progress_bar.empty()
                    st.rerun()
    with c3:
        if st.button("◀ 진단으로", use_container_width=True):
            st.session_state.step = 1
            st.session_state.revise_batches = None
            st.session_state.batch_results = {}
            st.rerun()


def show_step_3_verify():
    """Step 3: VERIFY 결과 확인 + 완료 이동."""
    vr = st.session_state.verify_result.get("verify_report", {})

    verdict = vr.get("overall_verdict", "")
    score = vr.get("overall_score", "")
    reason = vr.get("verdict_reason", "")

    verdict_class = "approved" if verdict == "APPROVED" else \
                    "needs" if verdict == "NEEDS_REVISION" else "rejected"

    st.markdown('<div class="rev-card-title">✅ Stage 3: 검증 결과</div>',
                unsafe_allow_html=True)

    # 판정
    st.markdown(f'<div style="text-align:center; padding: 20px 0;">'
                f'<span class="rev-verdict {verdict_class}">{verdict}</span>'
                f'<div style="font-size:1.4rem; font-weight:900; color:#191970; margin-top:12px;">'
                f'{score} / 10.0</div></div>', unsafe_allow_html=True)

    if reason:
        st.info(reason)

    st.markdown("---")

    # 4축 점수
    def _show_section(title_text, key, icon):
        block = vr.get(key, {})
        sc = block.get("score", "")
        st.markdown(f'<div class="rev-card-title">{icon} {title_text}  —  {sc}/10</div>',
                    unsafe_allow_html=True)
        return block

    # 1. 지시사항 반영
    block = _show_section("1. 지시사항 반영도", "instruction_compliance", "📝")
    for item in block.get("items", []):
        status = item.get("status", "")
        inst = item.get("instruction_item", "")
        ev = item.get("evidence", "")
        if status == "Y":
            st.success(f"✓ [Y] {inst}\n\n→ {ev}")
        elif status == "Partial":
            st.warning(f"△ [Partial] {inst}\n\n→ {ev}")
        else:
            st.error(f"✗ [N] {inst}\n\n→ {ev}")

    # 2. LOCKED 보존
    block = _show_section("2. LOCKED 보존도", "locked_preservation", "🔒")
    for item in block.get("items", []):
        status = item.get("status", "")
        li = item.get("locked_item", "")
        ev = item.get("evidence", "")
        if status == "Preserved":
            st.success(f"✓ [Preserved] {li}\n\n→ {ev}")
        elif status == "N/A":
            st.caption(f"— [N/A] {li}")
        else:
            st.error(f"✗ [Violated] {li}\n\n→ {ev}")

    # 3. AI ESCAPE
    block = _show_section("3. AI SCREENPLAY ESCAPE", "ai_escape_check", "🤖")
    clean = block.get("clean_patterns", "")
    if clean:
        st.success(f"✓ 위반 없는 패턴: {clean} / 20")
    violations = block.get("violations", [])
    for v in violations:
        pid = v.get("pattern_id", "")
        pname = v.get("pattern_name", "")
        scene = v.get("scene_id", "")
        quote = v.get("quote", "")
        sev = v.get("severity", "")
        emoji = "🔴" if sev == "High" else "🟠" if sev == "Medium" else "🟡"
        st.markdown(f"{emoji} **[{pid}] {pname}** — {scene} ({sev})")
        if quote:
            st.caption(f'→ "{quote}"')

    # 4. 장르 준수도
    block = _show_section("4. 장르 RULE PACK 준수도", "genre_compliance", "🎬")
    st.markdown("**Must Have**")
    for item in block.get("must_have_check", []):
        stt = item.get("status", "")
        emoji = "✓" if stt == "Met" else "△" if stt == "Partial" else "✗"
        st.markdown(f"{emoji} [{stt}] {item.get('item','')}  —  *{item.get('note','')}*")
    st.markdown("**Fails to Avoid**")
    for item in block.get("fails_check", []):
        stt = item.get("status", "")
        emoji = "✓" if stt == "Avoided" else "△" if stt == "Improved" else "✗"
        st.markdown(f"{emoji} [{stt}] {item.get('item','')}  —  *{item.get('note','')}*")

    # 하이라이트
    highlights = vr.get("side_by_side_highlights", [])
    if highlights:
        st.markdown("---")
        st.markdown('<div class="rev-card-title">🌟 핵심 변화 요약</div>', unsafe_allow_html=True)
        for h in highlights:
            st.markdown(f"• **{h.get('scene_id','')}**: {h.get('key_change','')}")
            st.caption(h.get("improvement_note", ""))

    # 재수정 권고
    recs = vr.get("recommendations", [])
    if recs:
        st.markdown("---")
        st.markdown('<div class="rev-card-title" style="color:#E14444;">🔁 재수정 권고</div>',
                    unsafe_allow_html=True)
        for r in recs:
            st.markdown(f"• {r}")

    st.markdown("---")

    # 완료로 이동
    c1, c2 = st.columns([1, 1])
    with c1:
        if st.button("🎉 완료 페이지로 (다운로드)", use_container_width=True):
            st.session_state.step = 4
            st.rerun()
    with c2:
        if st.button("◀ 집필 결과로 돌아가기", use_container_width=True):
            st.session_state.step = 2
            st.session_state.verify_result = None
            st.rerun()


def show_step_4_complete():
    """Step 4: 다운로드 + 초기화."""
    st.markdown('<div style="text-align:center; padding: 30px 0;">'
                '<span style="font-size: 2.4rem; font-weight: 950; color: #2EC484;">'
                '🎉 각색 완료!</span></div>', unsafe_allow_html=True)

    st.markdown('<div class="rev-card-title">📥 다운로드</div>', unsafe_allow_html=True)

    title = st.session_state.title or "수정본"
    genre = st.session_state.genre

    # ★ v3.3.2 — 파일명 가이드 안내 박스
    _round_v332 = st.session_state.get("round_n", 1)
    _score_v332 = st.session_state.get("verify_result", {}).get("overall_score", "?")
    st.markdown(
        f'<div style="background:#F8F9FB; padding:10px 14px; border-radius:6px; '
        f'border-left:3px solid #6B7280; margin:8px 0; font-size:0.85rem;">'
        f'<b>📁 v3.3.2 — 파일명 규칙</b><br>'
        f'모든 출력 파일은 <code>제목_종류_R(라운드)_score(점수)_YYYYMMDD_HHMM_Blue.확장자</code> 형식.<br>'
        f'<span style="color:#666;">'
        f'예시: <code>{title}_수정본_R{_round_v332}_score{int(_score_v332*10) if isinstance(_score_v332, (int,float)) else "??"}_'
        f'{datetime.now().strftime("%Y%m%d_%H%M")}_Blue.docx</code><br>'
        f'💡 같은 작품을 여러 라운드 작업해도 파일이 자동으로 구분됩니다 (시·분·라운드·점수 모두 포함).'
        f'</span></div>',
        unsafe_allow_html=True
    )

    c1, c2 = st.columns(2)

    # 수정본 DOCX
    with c1:
        try:
            is_historical = (st.session_state.period_key != "(현대)")
            # v2.2 — 구간 모드 + 손본본 있으면 손본본을 베이스로 사용
            base_text = st.session_state.raw_text
            if (st.session_state.section_mode
                    and st.session_state.diff_refined_text
                    and st.session_state.section_input_method in ("auto", "hybrid")):
                base_text = st.session_state.diff_refined_text

            docx_bytes = create_revised_docx(
                st.session_state.revise_result,
                title=title,
                genre=genre,
                original_text=base_text,
                fact_based=st.session_state.fact_based,
                historical=is_historical,
                historical_type=st.session_state.historical_type if is_historical else "",
            )
            st.download_button(
                "📄 수정본 (DOCX)",
                data=docx_bytes,
                file_name=_get_round_aware_filename(title, "revised"),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_revised",
                help="원본 + 수정된 씬을 통합한 최종 시나리오 (한국 표준 서식). 파일명에 라운드·시·분 표시.",
                use_container_width=True,
            )
        except Exception as e:
            st.error(f"수정본 DOCX 생성 오류: {e}")

    # 검증 보고서 DOCX
    with c2:
        try:
            verify_bytes = create_verify_docx(
                st.session_state.verify_result,
                title=title,
            )
            st.download_button(
                "✅ 검증 보고서 (DOCX)",
                data=verify_bytes,
                file_name=_get_round_aware_filename(title, "verify"),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_verify",
                help="지시사항 반영 + LOCKED 보존 + AI ESCAPE + 장르 준수도 검증. 파일명에 라운드·시·분 표시.",
                use_container_width=True,
            )
        except Exception as e:
            st.error(f"검증 보고서 DOCX 생성 오류: {e}")

    # ★ v3.3 — 검증 보고서 JSON 다운로드 (다음 라운드 자동 흡수용)
    st.markdown("---")
    st.markdown(
        '<div style="background:#FCE7F3; padding:10px 14px; border-radius:6px; '
        'border-left:3px solid #EC4899; margin:8px 0; font-size:0.88rem;">'
        '<b>🆕 v3.3 신규:</b> 검증 보고서 JSON 다운로드 — '
        '다음 라운드 작업 시 자동 흡수 가능. '
        '<span style="color:#666;">(이번 라운드 결과를 다음 라운드에 자동 반영해 8점 도달 사이클 구축)</span>'
        '</div>',
        unsafe_allow_html=True
    )
    cj1, cj2 = st.columns([1, 1])
    with cj1:
        try:
            current_round = st.session_state.get("round_n", 1)
            verify_json_bytes = export_verify_json(
                st.session_state.verify_result,
                title=title,
                round_n=current_round,
            )

            st.download_button(
                "📊 검증 보고서 (JSON) — 다음 라운드용",
                data=verify_json_bytes,
                file_name=_get_round_aware_filename(title, "verify", "json"),
                mime="application/json",
                key="dl_verify_json",
                help=(
                    "다음 라운드 작업 시 'v3.3 검증 보고서 흡수' 영역에 업로드하면 자동 변환.\n"
                    "파일명 형식: 제목_검증보고서_R(라운드)_score(점수)_YYYYMMDD_HHMM_Blue.json"
                ),
                use_container_width=True,
            )
        except Exception as e:
            st.error(f"검증 보고서 JSON 생성 오류: {e}")
    with cj2:
        st.caption(
            f"💡 현재 라운드: **Round {st.session_state.get('round_n', 1)}**. "
            f"다운로드한 JSON을 다음 작업의 입력 화면 → 'v3.3 검증 보고서 흡수' 영역에 업로드."
        )

    # JSON 백업 다운로드
    st.markdown("---")
    with st.expander("🗃️ 원본 JSON 백업 다운로드 (고급)"):
        full_state = {
            "meta": {
                "title": title,
                "genre": genre,
                "intensity": st.session_state.intensity,
                "generated_at": datetime.now().isoformat(),
                "engine": "BLUE JEANS REVISE ENGINE v3.3.5",
            },
            "input": {
                "instruction": st.session_state.instruction,
                "locked": st.session_state.locked,
            },
            "diagnose": st.session_state.diagnose_result,
            "revise":   st.session_state.revise_result,
            "verify":   st.session_state.verify_result,
        }
        json_bytes = json.dumps(full_state, ensure_ascii=False, indent=2).encode("utf-8")
        safe_title = re.sub(r'[/*?:"<>|]', '_', title)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M")  # ★ v3.3.2: 시·분 추가
        st.download_button(
            "📋 전체 JSON 다운로드",
            data=json_bytes,
            file_name=f"{safe_title}_revise_full_{timestamp}.json",
            mime="application/json",
            key="dl_json_full",
            use_container_width=True,
        )

    st.markdown("---")

    c1, c2 = st.columns([1, 1])
    with c1:
        if st.button("🔄 새 시나리오 각색 시작", use_container_width=True):
            reset_workflow()
            st.rerun()
    with c2:
        if st.button("◀ 검증 결과로 돌아가기", use_container_width=True):
            st.session_state.step = 3
            st.rerun()


# =================================================================
# [9] 메인 라우터
# =================================================================
render_hero()
render_stepbar()

step = st.session_state.step
if step == 0:
    show_step_0_input()
elif step == 1:
    show_step_1_diagnose()
elif step == 2:
    show_step_2_revise()
elif step == 3:
    show_step_3_verify()
elif step == 4:
    show_step_4_complete()

# ── 푸터 ──
st.markdown("---")
st.markdown(
    '<div style="text-align:center; color:#8E8E99; font-size:0.75rem; padding:20px 0;">'
    'BLUE JEANS PICTURES · REVISE ENGINE v3.3.5  ·  '
    'Powered by Claude Opus 4.6 + Sonnet 4.6  ·  '
    '<span style="color:#10B981;">Auto Batch Split</span>  ·  '
    '<span style="color:#F59E0B;">Beat-Aware Diagnose</span>  ·  '
    '<span style="color:#EC4899;">Beat Expansion Mode</span>  ·  '
    '<span style="color:#6366F1;">작가 친화 UI</span>  ·  '
    '<span style="color:#0EA5E9;">Writer v3.5.1 Sync</span>  ·  '
    '<span style="color:#8B5CF6;">Round N+1 Cycle</span>'
    '</div>',
    unsafe_allow_html=True,
)
